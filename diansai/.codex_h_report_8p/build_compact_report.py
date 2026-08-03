from __future__ import annotations

import copy
import hashlib
import math
from pathlib import Path
from zipfile import ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


ROOT = Path(r"D:\my_code\my_code\diansai")
WORK = ROOT / ".codex_h_report_8p"
TEMPLATE = WORK / "template.docx"
OUTPUT = ROOT / "H题模板" / "2026电赛报告H题_车载平衡滚球运动控制系统_MaixCAM识别_K230图传_阶段模板.docx"
FIGURE = WORK / "system_architecture_compact.png"

EXPECTED_TEMPLATE_HASH = "A5F3D12F406534731CD67480569F84A29EDF36E269486DA5F5DF2634A2EE35B9"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def archive_media_hashes(path: Path) -> dict[str, str]:
    with ZipFile(path) as package:
        return {
            name: hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        }


def section_signature(doc: Document) -> list[tuple]:
    return [
        (
            section.start_type,
            section.page_width,
            section.page_height,
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
            section.header_distance,
            section.footer_distance,
            section.different_first_page_header_footer,
            section.header.is_linked_to_previous,
            section.footer.is_linked_to_previous,
        )
        for section in doc.sections
    ]


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_text_run(paragraph: Paragraph, text: str, bold: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")
    return run


def set_paragraph(paragraph: Paragraph, text: str, label: str | None = None) -> None:
    clear_paragraph(paragraph)
    if label and text.startswith(label):
        add_text_run(paragraph, label, bold=True)
        add_text_run(paragraph, text[len(label) :])
    else:
        add_text_run(paragraph, text)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 8.5) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = add_text_run(paragraph, text, bold=bold)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_matrix(table: Table, rows: list[list[str]]) -> None:
    if len(rows) != len(table.rows):
        raise ValueError(f"row count mismatch: {len(rows)} != {len(table.rows)}")
    for row_index, values in enumerate(rows):
        if len(values) != len(table.columns):
            raise ValueError(f"column count mismatch in row {row_index}")
        for col_index, value in enumerate(values):
            set_cell(
                table.cell(row_index, col_index),
                value,
                bold=(row_index == 0),
                size=8.2 if len(table.columns) >= 5 else 8.5,
            )


def set_image_cell(cell, path: Path, width_inches: float, alt_text: str) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    inline = paragraph._p.xpath(".//wp:inline")[-1]
    doc_pr = inline.find(qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("descr", alt_text)


def image_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "msyhbd.ttc" if bold else "msyh.ttc"
    return ImageFont.truetype(str(Path(r"C:\Windows\Fonts") / name), size)


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, width=5, color="#536577") -> None:
    draw.line((start, end), fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max(math.hypot(dx, dy), 1.0)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    left = (x2 - 17 * ux + 8 * px, y2 - 17 * uy + 8 * py)
    right = (x2 - 17 * ux - 8 * px, y2 - 17 * uy - 8 * py)
    draw.polygon((end, left, right), fill=color)


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy,
    title: str,
    detail: str,
    fill: str,
) -> None:
    draw.rounded_rectangle(xy, radius=6, fill=fill, outline="#425466", width=3)
    x0, y0, x1, y1 = xy
    title_font = image_font(27, True)
    detail_font = image_font(20)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    detail_box = draw.textbbox((0, 0), detail, font=detail_font)
    draw.text(((x0 + x1 - title_box[2]) / 2, y0 + 20), title, font=title_font, fill="#16263A")
    draw.text(
        ((x0 + x1 - detail_box[2]) / 2, y1 - 52),
        detail,
        font=detail_font,
        fill="#26384D",
    )


def make_system_figure() -> None:
    image = Image.new("RGB", (1760, 920), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 25), "车载平衡滚球系统结构", font=image_font(40, True), fill="#173B63")
    draw.line((55, 90, 1705, 90), fill="#CBD5DF", width=3)

    draw.text((60, 115), "底盘循迹与运动控制", font=image_font(28, True), fill="#173B63")
    row1 = [
        ((60, 175, 330, 330), "8路红外模块", "8路独立IO并行输入", "#EDF2F6"),
        ((410, 175, 700, 330), "MSPM0G3507", "5 ms串级控制", "#DDEBF7"),
        ((790, 175, 1040, 330), "AT8236", "10 kHz电机驱动", "#FCE9D9"),
        ((1130, 175, 1390, 330), "MG513 x 2", "编码器差速驱动", "#E5F0DA"),
        ((1480, 175, 1700, 330), "底盘", "双轮+万向轮", "#F0F0F0"),
    ]
    for xy, title, detail, fill in row1:
        draw_box(draw, xy, title, detail, fill)
    for start, end in [((330, 252), (410, 252)), ((700, 252), (790, 252)), ((1040, 252), (1130, 252)), ((1390, 252), (1480, 252))]:
        draw_arrow(draw, start, end)
    draw_box(draw, (410, 365, 700, 490), "LSM6DSV16XTR", "SPI角速度/偏航反馈", "#E5F0DA")
    draw_arrow(draw, (555, 365), (555, 330))
    draw_arrow(draw, (1250, 350), (720, 350), color="#17857E")

    draw.text((60, 520), "滚球识别、执行与独立图传", font=image_font(28, True), fill="#173B63")
    row2 = [
        ((60, 580, 330, 745), "MaixCAM Pro", "钢球识别与连续追踪", "#DDEBF7"),
        ((410, 580, 700, 745), "滚球控制", "位置/速度控制(独立章节)", "#E5F0DA"),
        ((790, 580, 1080, 745), "闭环42步进", "UART绝对位置指令", "#FCE9D9"),
        ((1170, 580, 1430, 745), "摇杆连杆", "改变摆杆倾角", "#F0F0F0"),
    ]
    for xy, title, detail, fill in row2:
        draw_box(draw, xy, title, detail, fill)
    for start, end in [((330, 662), (410, 662)), ((700, 662), (790, 662)), ((1080, 662), (1170, 662))]:
        draw_arrow(draw, start, end)
    draw_arrow(draw, (1300, 770), (200, 770), color="#17857E")
    draw.text((650, 778), "摄像头获取钢球位置反馈", font=image_font(20), fill="#17857E")

    draw_box(draw, (1480, 560, 1700, 675), "K230", "独立图传发送", "#E9E4F3")
    draw_box(draw, (1480, 730, 1700, 845), "接收端", "显示/录像回放", "#F0F0F0")
    draw_arrow(draw, (1590, 675), (1590, 730))
    draw.text((65, 865), "注：MaixCAM Pro负责识别与追踪，K230仅负责图传，两条链路相互独立。", font=image_font(19), fill="#4B5563")
    image.save(FIGURE)


def clone_paragraph_xml(paragraph: Paragraph):
    clone = copy.deepcopy(paragraph._p)
    ppr = clone.find(qn("w:pPr"))
    if ppr is not None:
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is not None:
            ppr.remove(sect_pr)
    for child in list(clone):
        if child.tag != qn("w:pPr"):
            clone.remove(child)
    return clone


class BodyBuilder:
    def __init__(self, doc: Document, paragraph_templates, table_templates):
        self.doc = doc
        self.paragraph_templates = paragraph_templates
        self.table_templates = table_templates
        self.break_before_next = False

    @property
    def body(self):
        return self.doc._element.body

    def append_paragraph(self, template_index: int, text: str, label: str | None = None) -> Paragraph:
        paragraph_xml = copy.deepcopy(self.paragraph_templates[template_index])
        self.body.insert(self.body.index(self.body.sectPr), paragraph_xml)
        paragraph = Paragraph(paragraph_xml, self.doc._body)
        set_paragraph(paragraph, text, label)
        if self.break_before_next:
            paragraph.paragraph_format.page_break_before = True
            self.break_before_next = False
        return paragraph

    def page_break(self) -> None:
        self.break_before_next = True

    def append_table(self, template_index: int, rows: list[list[str]]) -> Table:
        table_xml = copy.deepcopy(self.table_templates[template_index])
        self.body.insert(self.body.index(self.body.sectPr), table_xml)
        table = Table(table_xml, self.doc._body)
        set_table_matrix(table, rows)
        return table

    def append_figure(self, path: Path, width_inches: float, alt_text: str) -> None:
        table_xml = copy.deepcopy(self.table_templates[0])
        self.body.insert(self.body.index(self.body.sectPr), table_xml)
        table = Table(table_xml, self.doc._body)
        set_image_cell(table.cell(0, 0), path, width_inches, alt_text)

    def append_visual_placeholder(self) -> None:
        table_xml = copy.deepcopy(self.table_templates[0])
        self.body.insert(self.body.index(self.body.sectPr), table_xml)
        table = Table(table_xml, self.doc._body)
        text = (
            "【视觉追踪图示预留】\n"
            "建议补入：相机画面、ROI范围、YOLO候选框与跨帧轨迹示意。\n"
            "建议标注：钢球中心坐标、置信度、有效标志、时间戳及坐标原点。\n"
            "遮挡或短时丢失时的预测保持、搜索与重捕获路径可在图中一并说明。\n"
            "\n"
            "本图示位仅服务于MaixCAM Pro的识别与连续追踪，不放置视觉输出之后的位置/速度PID。"
        )
        set_cell(table.cell(0, 0), text, size=9.5)
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def clear_old_body(doc: Document, anchor_paragraph_index: int = 34) -> None:
    anchor = doc.paragraphs[anchor_paragraph_index]._p
    body = doc._element.body
    current = anchor.getnext()
    while current is not None and current.tag != qn("w:sectPr"):
        next_node = current.getnext()
        body.remove(current)
        current = next_node


def update_front_matter(doc: Document) -> None:
    set_paragraph(
        doc.paragraphs[5],
        "摘要：本文设计车载平衡滚球运动控制系统。底盘以MSPM0G3507为主控，八路自校准红外模块通过八路独立IO并行输入，LSM6DSV16XTR提供角速度信息，AT8236驱动两台MG513电机完成差速循迹。滚球装置采用MaixCAM Pro进行钢球识别与连续追踪，并向自闭环42步进电机发送位置指令，经摇杆连杆改变摆杆倾角；另一块K230仅承担图传发送与录像链路。底盘软件已形成5 ms位置、偏航和轮速串级控制框架，视觉追踪参数、图示及全部性能数据留待联调和实测补充。本稿按官方模板形成可编辑阶段报告，不虚构测试结论。",
        "摘要：",
    )
    set_paragraph(doc.paragraphs[6], "关键词：八路红外循迹；MaixCAM Pro；滚球控制；闭环步进电机；K230图传", "关键词：")
    set_paragraph(
        doc.paragraphs[7],
        "阶段说明：当前处于初级联调阶段。文中程序结构来自最新版工程；视觉识别与连续追踪章节已按视觉方案形成阶段稿，图示与复测数据仍待补入。",
        "阶段说明：",
    )
    set_paragraph(doc.paragraphs[8], "")
    set_paragraph(doc.paragraphs[9], "")
    set_paragraph(doc.paragraphs[10], "目 录")

    entries = [
        ("一、系统方案", "1", "toc 1"),
        ("1、总体方案与机械布局", "1", "toc 2"),
        ("2、方案选择", "1", "toc 2"),
        ("二、底盘循迹理论与控制", "2", "toc 1"),
        ("1、八路IO偏差计算", "2", "toc 2"),
        ("2、差速模型与串级控制", "2", "toc 2"),
        ("三、滚球机构与控制", "3", "toc 1"),
        ("1、球杆模型与机械执行", "3", "toc 2"),
        ("2、位置-速度控制与步进指令", "3", "toc 2"),
        ("四、电路与程序设计", "4", "toc 1"),
        ("五、视觉识别与连续追踪", "5", "toc 1"),
        ("六、测试方案与待测记录", "6", "toc 1"),
        ("七、结论与改进", "7", "toc 1"),
        ("八、参考文献", "7", "toc 1"),
    ]
    for offset, (title, page, style_name) in enumerate(entries, start=11):
        set_paragraph(doc.paragraphs[offset], f"{title}\t{page}")
        doc.paragraphs[offset].style = doc.styles[style_name]
    for index in range(11 + len(entries), 35):
        set_paragraph(doc.paragraphs[index], "")


def build_body(doc: Document, paragraph_templates, table_templates) -> None:
    b = BodyBuilder(doc, paragraph_templates, table_templates)

    # Body page 1
    b.append_paragraph(35, "车载平衡滚球运动控制系统（H题）")
    b.append_paragraph(36, "【本科组】")
    b.append_paragraph(37, "一、系统方案")
    b.append_paragraph(38, "1、总体方案与机械布局")
    b.append_paragraph(
        140,
        "系统分为底盘循迹、滚球控制和图传三条链路。MSPM0G3507读取八路红外数字状态、编码器和LSM6DSV16XTR，计算左右轮目标并由AT8236驱动两台MG513；车体采用两主动轮加一万向轮。MaixCAM Pro负责钢球视觉识别与追踪，输出球位后在本机完成滚球控制并通过UART向自闭环42步进电机发送位置指令。K230作为独立图传发送端，将摆杆画面传到场外接收设备显示和录像，不参与识别与控制。",
    )
    b.append_paragraph(
        140,
        "机械部分采用一端铰接的25 cm凹槽PPR摆杆，另一端由42步进电机经摇杆连杆抬升。摄像头固定在摆杆上方，视场覆盖全行程；铰点高度、连杆孔距、允许摆角和整车外廓以装配实测为准，并按题目35 cm x 25 cm及h不小于5 cm约束复核。",
    )
    b.append_figure(FIGURE, 5.55, "车载平衡滚球系统结构图")
    b.append_paragraph(41, "图1 车载平衡滚球系统结构图")
    b.append_paragraph(38, "2、方案选择")
    b.append_paragraph(
        140,
        "巡线选用八路自校准红外模块，其板内完成黑白阈值调节，八路OUT分别接入主控的独立GPIO并同步采样。底盘采用直流减速电机和编码器，适合轮速闭环；滚球端采用自闭环步进电机，便于执行绝对位置命令并检测回零/到位状态。视觉识别与图传使用两块独立板卡，使识别控制负载和无线录像互不干扰。",
    )
    b.page_break()

    # Body page 2
    b.append_paragraph(37, "二、底盘循迹理论与控制")
    b.append_paragraph(38, "1、八路IO偏差计算")
    b.append_paragraph(
        140,
        "最新版底盘程序将OUT1~OUT8接至PA27、PA26、PA25、PA24、PB25、PB24、PB20、PA22，每5 ms并行读取一次。默认按低电平有效把原始电平转换为bits[8]。K1与K2模式使用两组位置权重，均按从左到右的负到正顺序排列；对所有有效通道求权重平均并归一化，得到e_line属于[-1,1]。全同电平保持上一帧，可抑制瞬时干扰。",
    )
    b.append_paragraph(60, "e_line = [Σ(b_i w_i) / Σb_i] / w_max，b_i∈{0,1}")
    b.append_paragraph(
        140,
        "终点横线采用滑动窗口判断：在连续20帧（约100 ms）内，若至少4路曾检测到黑线，则置位停车标志并调用底盘停止函数。该逻辑只说明当前程序判据，窗口长度、有效电平和触发路数仍需在A点横线实测后确认。",
    )
    b.append_paragraph(38, "2、差速模型与串级控制")
    b.append_paragraph(
        140,
        "设车体线速度为v，轮距为B，目标偏航角速度为ω，则左右轮目标速度可写为v_L=v-Bω/2、v_R=v+Bω/2。程序采用三级串联结构：位置PD把e_line转换为目标角速度；偏航环比较目标角速度与LSM6DSV16XTR的Z轴角速度，生成左右轮速度差；双轮增量式PID再将速度误差转换为AT8236的PWM输出。控制周期为5 ms，电机PWM初始化频率为10 kHz。",
    )
    b.append_paragraph(60, "ω* = PD(e_line)，Δv = PID_yaw(ω* - ω_z)，u_L/u_R = PID_speed(v_L* - v_L, v_R* - v_R)")
    b.append_paragraph(
        140,
        "LSM6DSV16XTR通过SPI初始化并校验器件ID，静止采样用于估计角速度零偏。上电默认关闭控制环，按键启动时加载对应模式参数并复位历史量；急停、失线、参数切换均应清除积分并令PWM回零。当前代码参数属于调试初值，报告最终值须以实车记录替换。",
    )
    b.page_break()

    # Body page 3
    b.append_paragraph(37, "三、滚球机构与控制")
    b.append_paragraph(38, "1、球杆模型与机械执行")
    b.append_paragraph(
        140,
        "以摆杆中心O为x=0，沿凹槽向驱动端为正方向，摆杆小角度为θ。把钢球近似为纯滚动实心球，I=2mr^2/5，则在忽略滑动并把车体纵向加速度记为a_x时，可得x¨约等于-(5/7)(gθ+a_x)。该模型表明摆角直接改变球的加速度，车体起步、弯道和制动会作为外扰进入滚球系统。摩擦、回差和杆面不水平造成的偏差由限幅、死区和后续标定修正。",
    )
    b.append_paragraph(60, "(m + I/r^2)x¨ = -m(g sinθ + a_x cosθ)，小角度下 x¨ ≈ -(5/7)(gθ + a_x)")
    b.append_paragraph(
        140,
        "自闭环42步进电机安装在驱动端，经摇臂和连杆把电机转角转换为摆杆端点位移。MaixCAM Pro通过115200 bit/s UART发送使能、回零、绝对位置和位置查询指令；正常上电应先读取位置并完成回零，再进入控制。电机角度到摆角的比例、机械零位、软限位和回差补偿均需由装配标定表确定。",
    )
    b.append_paragraph(38, "2、位置-速度控制与步进指令")
    b.append_paragraph(
        140,
        "本节描述视觉输出之后的滚球控制，与第5节视觉追踪算法分开。视觉模块输出球位置x_k和时间戳，控制侧先进行坐标标定与速度估计。位置环根据目标位置x_r生成受限目标速度，速度环再生成目标摆角θ_d；摆角经机构映射为电机绝对位置。积分限幅、输出限幅、角度死区和丢球回水平用于降低饱和与机械冲击。",
    )
    b.append_paragraph(60, "v_r = sat[K_px(x_r-x)+K_ixΣe_x]；θ_d = sat[K_pv(v_r-v)+K_ivΣe_v]")
    b.append_paragraph(
        140,
        "任务3的目标序列为O到+5 cm，再折返至-5 cm并稳定；动态任务将x_r设为O或任意指定位置。到位判据、保持时间、预测补偿和开环辅助段在最新版视觉目录中仍处于多版本调试状态，最终报告只保留实际采用的状态机和实测参数。",
    )
    b.page_break()

    # Body page 4
    b.append_paragraph(37, "四、电路与程序设计")
    b.append_paragraph(38, "1、主要硬件与接口")
    b.append_paragraph(42, "表1 主要硬件与接口")
    b.append_table(
        3,
        [
            ["模块", "型号/数量", "接口", "主要作用", "当前说明"],
            ["底盘主控", "MSPM0G3507 / 1", "GPIO、PWM、SPI、UART", "循迹、姿态、轮速与状态机", "最新版工程"],
            ["循迹模块", "8路自校准红外 / 1", "8路独立数字IO", "输出黑线0/1状态", "独立GPIO同步采样"],
            ["惯性传感", "LSM6DSV16XTR / 1", "SPI", "Z轴角速度与偏航信息", "零偏待实测"],
            ["底盘驱动", "AT8236 / 1", "10 kHz PWM/方向", "驱动两台MG513", "引脚按实物复核"],
            ["底盘电机", "MG513 / 2", "编码器反馈", "左右轮差速运动", "轮径/计数待测"],
            ["识别主控", "MaixCAM Pro / 1", "Camera、UART 115200", "钢球识别、追踪、步进指令", "算法段预留"],
            ["滚球执行", "自闭环42步进 / 1", "TTL串口", "经摇杆连杆改变摆角", "限位/回差待测"],
            ["图传发送", "K230 / 1", "摄像/无线链路待确认", "独立实时图传与录像", "不参与识别"],
        ],
    )
    b.append_paragraph(38, "2、软件结构与时序")
    b.append_paragraph(
        140,
        "底盘以1 ms SysTick计时，每5 ms读取八路GPIO、LSM6DSV16XTR和编码器，依次更新循迹位置、偏航及双轮速度环与PWM。串口、OLED和按键采用轮询/分片处理；K1/K2参数、启动斜坡和横线停车逻辑以最新版深层工程为准。",
    )
    b.append_paragraph(
        140,
        "MaixCAM Pro运行相机、检测/追踪、坐标估计、滚球控制和步进通信；K230独立图传，不向控制链发送检测结果。启动时完成传感器自检和步进回零，异常时停止电机命令并保留录像/日志。",
    )
    b.append_paragraph(38, "3、电源与抗干扰")
    b.append_paragraph(
        140,
        "电池分支向底盘、步进和数字系统供电并共地，电机与逻辑/视觉支路分开稳压布线，AT8236及步进驱动就近配置去耦和储能电容；峰值电流、稳压余量和图传时长待实测。",
    )
    b.page_break()

    # Body page 5, with room for the vision teammate's final figures and data.
    b.append_paragraph(37, "五、视觉识别与连续追踪")
    b.append_paragraph(
        140,
        "MaixCAM Pro负责钢球图像采集和视觉识别，K230仅承担独立图传。相机画面先限定摆杆有效区域ROI，再在ROI内使用定制YOLOv5模型完成钢球候选检测，并结合候选筛选和轮廓精修得到钢球中心坐标、置信度和时间戳。方案文档记录检测帧率约30 fps、模型验证集准确率约84%，该数值仍需在最终安装姿态、光照和摆杆背景下复测。",
    )
    b.append_paragraph(38, "1、连续追踪与丢失处理")
    b.append_paragraph(
        140,
        "连续追踪以当前检测结果和上一帧轨迹预测位置进行跨帧关联，并用EMA对中心坐标作平滑，按时间戳计算速度供后级控制使用。当出现短时遮挡或检测置信度下降时，保持有限帧的预测轨迹并标记有效状态；超过丢失阈值后进入搜索/重捕获，重新获得满足置信度条件的目标后恢复正常输出。追踪输出字段包括x、y（或标定后的一维球位）、速度、时间戳、置信度和valid标志，坐标原点及正方向与第3节控制接口统一。",
    )
    b.append_visual_placeholder()
    b.append_paragraph(38, "2、追踪验证数据（待复测）")
    b.append_paragraph(
        140,
        "待补入不同球位、车体静止/运动、反光和短时遮挡工况下的检测成功率、位置重复误差、端到端延迟、丢失持续时间及重捕获时间。视觉控制方案中的延迟补偿、微分先行和状态机属于追踪输出后的控制环节，放在第3节描述，本节只保留识别、追踪和接口数据。",
    )
    b.page_break()

    # Body page 6
    b.append_paragraph(37, "六、测试方案与待测记录")
    b.append_paragraph(38, "1、测试原则")
    b.append_paragraph(
        140,
        "所有性能结论至少基于5次完整测试，失败试验也保留。循迹测试记录总时间、停车偏差和脱轨情况；滚球测试同步保存视频、球位时间序列、目标值和步进命令。位置误差按max|x-x_r|统计，并同时给出平均值；图传需验证覆盖全摆杆、连续显示、录像和回放。",
    )
    b.append_paragraph(42, "表2 分项测试记录（当前待实测）")
    b.append_table(
        15,
        [
            ["测试项目", "题目限值", "记录内容", "当前结果", "证据"],
            ["K230图传", "稳定显示并完整录像", "连续性/覆盖/回放", "待实测", "视频待编号"],
            ["循迹一圈停车", "不大于20 s；偏差不大于2 cm", "5次时间与停车偏差", "待实测", "日志/视频"],
            ["静态O到正负5 cm", "不大于5 s；误差不大于1 cm", "运行时间/两端最大误差", "待实测", "球位序列"],
            ["A-B动态中心", "不大于8 s；误差不大于1 cm", "A-B时间/中心误差", "待实测", "同步记录"],
            ["一圈动态中心", "不大于30 s；误差不大于1 cm", "总时间/中心误差", "待实测", "同步记录"],
            ["任意指定位置", "不大于30 s；误差不大于1 cm", "目标/时间/最大误差", "待实测", "同步记录"],
        ],
    )
    b.append_paragraph(38, "2、联调顺序与误差分析")
    b.append_paragraph(
        140,
        "建议依次完成八路IO方向检查、左右轮编码器符号、单轮速度环、LSM6DSV16XTR零偏、低速循迹、横线停车、步进回零和角度标定、静态滚球控制，最后进行整车动态联合测试。每次只改变一组参数，并保存对应日志。最新版底盘源码晚于现有部分构建日志，最终提交前需重新全量构建并记录0错误、0警告。",
    )
    b.append_paragraph(
        140,
        "主要误差来源包括红外安装高度与横向位置、轮胎打滑和左右轮差异、IMU零偏、连杆回差、摆杆不水平、视觉标定与延迟、电机供电压降。当前阶段仅给出验证方法，不给出未测量的达标结论。",
    )
    b.page_break()

    # Body page 7
    b.append_paragraph(37, "七、结论与改进")
    b.append_paragraph(
        140,
        "本阶段已确定以MSPM0G3507、八路独立IO红外、LSM6DSV16XTR、AT8236和两台MG513构成底盘，以MaixCAM Pro、自闭环42步进电机和摇杆连杆构成滚球闭环，并使用另一块K230完成独立图传。底盘最新版程序已具备5 ms串级控制、K1/K2模式、启动斜坡和横线停车框架；滚球端具备步进电机回零、绝对位置命令和控制程序备份。视觉识别与连续追踪方案已补入，最终版本固定、整车联合调参和表2数据仍待完成，因此本文不对任务2至任务6作达标判定。",
        "本阶段已确定",
    )
    b.append_paragraph(
        140,
        "下一步首先固定视觉追踪版本及输出协议，完成像素-厘米、步进角-摆角和左右轮速度标定；随后按由内到外的顺序闭合轮速、偏航、循迹和滚球环路；最后在K230持续图传录像条件下完成5次以上整车测试，并用原始数据替换摘要、结论和测试表中的“待实测”。",
    )
    b.append_paragraph(37, "八、参考文献")
    references = [
        "[1] 2026年全国大学生电子设计竞赛组委会. 车载平衡滚球运动控制系统（H题）[Z]. 2026.",
        "[2] Texas Instruments. MSPM0G3507 Mixed-Signal Microcontrollers Data Sheet[EB/OL].",
        "[3] STMicroelectronics. LSM6DSV16X 6-axis IMU Data Sheet[EB/OL].",
        "[4] AT8236双通道直流电机驱动芯片数据手册[Z].",
        "[5] Sipeed. MaixCAM Pro Documentation[EB/OL].",
        "[6] Canaan. K230 Technical Documentation[EB/OL].",
    ]
    for reference in references:
        b.append_paragraph(140, reference)
    b.append_paragraph(
        140,
        "注：最终提交时应补齐实际使用数据手册的版本、发布日期或访问日期，并删除未在正文引用的条目。",
    )


def main() -> None:
    if sha256_file(TEMPLATE) != EXPECTED_TEMPLATE_HASH:
        raise RuntimeError("template SHA-256 mismatch; fresh distillation required")

    make_system_figure()
    doc = Document(TEMPLATE)
    cover_xml = [etree.tostring(p._p, with_tail=False) for p in doc.paragraphs[:5]]
    cover_text = [p.text for p in doc.paragraphs[:5]]
    section_before = section_signature(doc)
    media_before = archive_media_hashes(TEMPLATE)
    paragraph_templates = {index: clone_paragraph_xml(doc.paragraphs[index]) for index in [35, 36, 37, 38, 41, 42, 60, 140]}
    table_templates = {index: copy.deepcopy(doc.tables[index]._tbl) for index in [0, 3, 15]}

    update_front_matter(doc)
    clear_old_body(doc)
    build_body(doc, paragraph_templates, table_templates)
    set_update_fields_on_open(doc)
    doc.core_properties.title = "车载平衡滚球运动控制系统（H题）阶段模板"
    doc.core_properties.subject = "MaixCAM Pro识别、K230独立图传、八路IO循迹"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    result = Document(OUTPUT)
    assert [p.text for p in result.paragraphs[:5]] == cover_text
    assert [etree.tostring(p._p, with_tail=False) for p in result.paragraphs[:5]] == cover_xml
    assert section_signature(result) == section_before
    assert len(result.sections) == 3
    media_after = archive_media_hashes(OUTPUT)
    for name, digest in media_before.items():
        assert media_after.get(name) == digest, f"template media changed: {name}"

    # The official cover is preserved byte-for-byte and intentionally retains
    # the template date placeholder; audit editable report content separately.
    body_text = "\n".join(p.text for p in result.paragraphs[5:])
    body_text += "\n" + "\n".join(cell.text for table in result.tables for row in table.rows for cell in row.cells)
    forbidden = ["S0/S1/S2", "分时复用", "TB6612", "MPU6050", "六路红外", "K230负责识别"]
    leaked = [token for token in forbidden if token in body_text]
    assert not leaked, f"forbidden legacy terms leaked: {leaked}"
    assert "【请填写" not in body_text
    assert "____" not in body_text
    assert "K230仅" in body_text or "K230作为独立图传" in body_text
    assert "八路独立IO" in body_text

    print(f"created={OUTPUT}")
    print(f"paragraphs={len(result.paragraphs)} tables={len(result.tables)} sections={len(result.sections)}")
    print(f"inline_shapes={len(result.inline_shapes)} media={len(media_after)}")
    print("cover_xml=IDENTICAL")
    print("sections=IDENTICAL")
    print("template_media=IDENTICAL")


if __name__ == "__main__":
    main()
