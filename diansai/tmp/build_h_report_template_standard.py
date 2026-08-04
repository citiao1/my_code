# -*- coding: utf-8 -*-
from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\my_code\my_code\diansai")
SOURCE = ROOT / "H题模板" / "2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_封面目录分页版.docx"
OUTPUT = ROOT / "H题模板" / "2026电赛报告H题_MaixCAM方案_通用模板规范版.docx"
ASSET_DIR = ROOT / "tmp" / "h_report_assets"
ARCH_IMAGE = ASSET_DIR / "system_architecture_template_standard.png"
MECH_IMAGE = ASSET_DIR / "rocker_mechanism.png"

BLACK = RGBColor(0, 0, 0)
SONGTI = "宋体"
LATIN = "Times New Roman"
MATH = "Cambria Math"


def set_run_font(run, size=12, bold=None, italic=None, east_asia=SONGTI, latin=LATIN):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, size, bold=False, before=0, after=0, keep_next=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    style.font.name = LATIN
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN)
    fonts.set(qn("w:hAnsi"), LATIN)
    fonts.set(qn("w:eastAsia"), SONGTI)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(22)
    fmt.keep_with_next = keep_next


def configure_document_styles(doc):
    configure_style(doc.styles["Normal"], 12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    configure_style(doc.styles["Heading 1"], 14, bold=True, before=3, keep_next=True)
    configure_style(doc.styles["Heading 2"], 12, bold=True, before=1, keep_next=True)
    configure_style(doc.styles["Heading 3"], 12, bold=True, before=1, keep_next=True)
    if "Caption" in [s.name for s in doc.styles]:
        configure_style(doc.styles["Caption"], 10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name, left in (("TOC 1", 0), ("TOC 2", 21)):
        if name in [s.name for s in doc.styles]:
            configure_style(doc.styles[name], 10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            doc.styles[name].paragraph_format.left_indent = Pt(left)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_text_run(paragraph, text, size=12, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


def add_math_run(paragraph, base, sub=None, sup=None, size=12):
    run = paragraph.add_run(base)
    set_run_font(run, size=size, italic=True, east_asia=SONGTI, latin=MATH)
    if sub is not None:
        sub_run = paragraph.add_run(str(sub))
        set_run_font(sub_run, size=size, italic=True, east_asia=SONGTI, latin=MATH)
        sub_run.font.subscript = True
    if sup is not None:
        sup_run = paragraph.add_run(str(sup))
        set_run_font(sup_run, size=size, italic=True, east_asia=SONGTI, latin=MATH)
        sup_run.font.superscript = True


def add_rich_paragraph(doc, parts, bold_label=None, keep_with_next=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.keep_with_next = keep_with_next
    if bold_label:
        add_text_run(p, bold_label, bold=True)
    for part in parts:
        if isinstance(part, str):
            add_text_run(p, part)
        elif part[0] == "math":
            add_math_run(p, part[1], part[2] if len(part) > 2 else None, part[3] if len(part) > 3 else None)
        elif part[0] == "code":
            run = p.add_run(part[1])
            set_run_font(run, size=10.5, latin="Consolas")
        elif part[0] == "bold":
            add_text_run(p, part[1], bold=True)
    return p


def add_body(doc, text, bold_label=None):
    return add_rich_paragraph(doc, [text], bold_label=bold_label)


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def m_run(text):
    node = OxmlElement("m:r")
    rpr = OxmlElement("m:rPr")
    sty = OxmlElement("m:sty")
    sty.set(qn("m:val"), "i")
    rpr.append(sty)
    node.append(rpr)
    value = OxmlElement("m:t")
    value.text = text
    node.append(value)
    return node


def m_sub(base, sub):
    node = OxmlElement("m:sSub")
    node.append(OxmlElement("m:sSubPr"))
    e = OxmlElement("m:e")
    e.append(m_run(base))
    s = OxmlElement("m:sub")
    s.append(m_run(sub))
    node.extend([e, s])
    return node


def m_sup(base, sup):
    node = OxmlElement("m:sSup")
    node.append(OxmlElement("m:sSupPr"))
    e = OxmlElement("m:e")
    e.append(m_run(base))
    s = OxmlElement("m:sup")
    s.append(m_run(sup))
    node.extend([e, s])
    return node


def m_frac(numerator, denominator):
    node = OxmlElement("m:f")
    node.append(OxmlElement("m:fPr"))
    num = OxmlElement("m:num")
    den = OxmlElement("m:den")
    for item in numerator:
        num.append(m_part(item))
    for item in denominator:
        den.append(m_part(item))
    node.extend([num, den])
    return node


def m_part(item):
    if isinstance(item, str):
        return m_run(item)
    if item[0] == "sub":
        return m_sub(item[1], item[2])
    if item[0] == "sup":
        return m_sup(item[1], item[2])
    if item[0] == "frac":
        return m_frac(item[1], item[2])
    raise ValueError(item)


def add_formula(doc, parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    math_para = OxmlElement("m:oMathPara")
    math_para_pr = OxmlElement("m:oMathParaPr")
    justification = OxmlElement("m:jc")
    justification.set(qn("m:val"), "center")
    math_para_pr.append(justification)
    math_para.append(math_para_pr)
    math = OxmlElement("m:oMath")
    for part in parts:
        math.append(m_part(part))
    math_para.append(math)
    p._p.append(math_para)
    return p


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fill_cell(cell, content, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(16)
    parts = content if isinstance(content, list) else [content]
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run_font(run, size=9, bold=header)
        elif part[0] == "math":
            add_math_run(p, part[1], part[2] if len(part) > 2 else None, part[3] if len(part) > 3 else None, size=9)
        elif part[0] == "code":
            run = p.add_run(part[1])
            set_run_font(run, size=8.5, latin="Consolas")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if header:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E7E6E6")
        tc_pr.append(shd)


def add_table(doc, caption, headers, rows, widths_cm=None):
    cap = doc.add_paragraph(style="Caption")
    add_text_run(cap, caption, size=10.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for j, header in enumerate(headers):
        fill_cell(table.rows[0].cells[j], header, header=True)
    for row_data in rows:
        cells = table.add_row().cells
        for j, content in enumerate(row_data):
            fill_cell(cells[j], content)
    if widths_cm:
        for row in table.rows:
            for index, width in enumerate(widths_cm):
                row.cells[index].width = Cm(width)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def image_font(size, bold=False):
    font_dir = Path(r"C:\Windows\Fonts")
    for name in (("simsun.ttc", "simhei.ttf") if not bold else ("simhei.ttf", "simsun.ttc")):
        path = font_dir / name
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_arrow(draw, start, end, width=5):
    draw.line((start, end), fill="#4A4A4A", width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    left = (x2 - 18 * ux + 9 * px, y2 - 18 * uy + 9 * py)
    right = (x2 - 18 * ux - 9 * px, y2 - 18 * uy - 9 * py)
    draw.polygon((end, left, right), fill="#4A4A4A")


def draw_box(draw, xy, title, lines, fill):
    draw.rounded_rectangle(xy, radius=8, fill=fill, outline="#555555", width=3)
    x0, y0, x1, y1 = xy
    tf = image_font(30, bold=True)
    df = image_font(23)
    title_box = draw.textbbox((0, 0), title, font=tf)
    draw.text(((x0 + x1 - title_box[2]) / 2, y0 + 18), title, font=tf, fill="#000000")
    y = y0 + 70
    for line in lines:
        box = draw.textbbox((0, 0), line, font=df)
        draw.text(((x0 + x1 - box[2]) / 2, y), line, font=df, fill="#000000")
        y += 34


def make_architecture_image():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(image)
    title = "系统功能边界与信号流"
    draw.text((60, 35), title, font=image_font(40, bold=True), fill="#000000")
    draw_box(draw, (70, 145, 390, 340), "8路红外阵列", ["黑线数字状态", "5 ms采样"], "#F2F2F2")
    draw_box(draw, (520, 120, 880, 365), "底盘控制器", ["横向偏差PD", "偏航角速度PI", "双轮速度PID"], "#E2F0D9")
    draw_box(draw, (1030, 145, 1360, 340), "差速底盘", ["左右轮电机", "编码器反馈"], "#FCE4D6")
    draw_arrow(draw, (390, 245), (520, 245))
    draw_arrow(draw, (880, 245), (1030, 245))
    draw_arrow(draw, (1030, 315), (880, 315), width=4)
    draw_box(draw, (70, 515, 390, 730), "MaixCAM视觉", ["仅识别与跟踪钢珠", "输出x、valid、时间戳"], "#DDEBF7")
    draw_box(draw, (520, 490, 880, 755), "滚球控制与映射", ["位置PD生成球杆角", "限幅与斜率限制", "球杆角到电机角"], "#E2F0D9")
    draw_box(draw, (1030, 515, 1360, 730), "步进电机机构", ["UART绝对位置", "摇臂连杆抬升杆端"], "#FCE4D6")
    draw_arrow(draw, (390, 625), (520, 625))
    draw_arrow(draw, (880, 625), (1030, 625))
    draw_box(draw, (1490, 285, 1725, 600), "公共约束", ["统一供电", "共地", "急停", "状态显示", "日志记录"], "#F2F2F2")
    draw_arrow(draw, (1490, 365), (1360, 265), width=3)
    draw_arrow(draw, (1490, 535), (1360, 625), width=3)
    image.save(ARCH_IMAGE)


def add_picture(doc, path, caption, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    add_text_run(cap, caption, size=10.5)


def replace_front_matter(doc):
    paragraphs = doc.paragraphs
    # Cover and front-matter titles are black Songti; original font sizes are retained.
    for p in paragraphs[:31]:
        for run in p.runs:
            size = run.font.size.pt if run.font.size else 12
            set_run_font(run, size=size, bold=run.bold)

    subtitle = paragraphs[4]
    clear_paragraph(subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(subtitle, "MaixCAM方案：通用模板规范版", size=15, bold=True)

    abstract = paragraphs[5]
    clear_paragraph(abstract)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_run(abstract, "摘要：", bold=True)
    add_text_run(
        abstract,
        "系统面向车载平衡滚球任务，采用底盘循迹与球杆调节相对独立、统一供电和协同状态管理的总体结构。底盘以MSPM0G3507为控制器，利用8路数字红外阵列形成黑线横向偏差，经位置PD、偏航角速度PI和双轮速度PID三级闭环实现循迹、分段调速、计时与终点停车。MaixCAM上的视觉模块仅负责钢珠识别、连续跟踪、像素到物理坐标标定以及有效性判断，并向控制层输出球位、有效标志和时间戳；滚球控制层根据目标位置误差通过位置PD生成球杆目标角，经安全限幅和机构标定换算为步进电机绝对角度，再由UART驱动摇臂连杆机构完成球杆一端升降。本文依据现有底盘与步进电机代码给出方案论证、理论模型、程序接口、机械结构及测试方法，视觉算法细节和所有实测结果保留为后续填写项，不虚构尚未取得的数据。",
    )
    abstract.paragraph_format.line_spacing = Pt(22)

    keywords = paragraphs[6]
    clear_paragraph(keywords)
    add_text_run(keywords, "关键词：", bold=True)
    add_text_run(keywords, "MaixCAM；钢珠视觉跟踪；8路红外循迹；位置PD；步进电机；摇臂连杆机构")

    stage_table = doc.tables[0]
    fill_cell(
        stage_table.cell(0, 0),
        "编写范围说明：视觉章节只保留钢珠识别、跟踪、坐标标定、数据有效性和性能测试位置；位置PD、球杆角度生成、机构映射、UART通信及步进电机控制均归入控制与执行机构部分。视觉算法和实测数据由队友完成后据实补入。",
    )

    body_header = doc.sections[-1].header.paragraphs[0]
    clear_paragraph(body_header)
    body_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text_run(body_header, "H题  MaixCAM方案：底盘、机械机构与滚球控制", size=9)

    doc.core_properties.title = "H题 MaixCAM方案：底盘、机械机构与滚球控制"
    doc.core_properties.subject = "8路红外循迹、钢珠视觉跟踪、位置PD与步进电机摇臂机构"
    doc.core_properties.author = "参赛队"


def remove_old_body(doc):
    first = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            first = p._p
            break
    if first is None:
        raise RuntimeError("Body heading not found")
    body = doc.element.body
    children = list(body)
    start = children.index(first)
    for child in children[start:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_body(doc):
    add_heading(doc, "一、系统方案", 1)
    add_heading(doc, "1、总体方案", 2)
    add_body(
        doc,
        "系统由循迹底盘、钢珠视觉检测、滚球控制、步进电机执行机构、电源与状态管理五部分组成。底盘控制器独立完成高频循迹闭环，MaixCAM完成图像采集并直接在本机得到钢珠位置；控制层读取有效球位后计算位置误差，通过位置PD得到球杆目标角，再经摇臂机构标定和串口协议转换为步进电机绝对位置命令。该分工不把底盘实时性依赖于视觉帧率，同时保证视觉测量到机械动作的数据链清晰可追溯。",
        bold_label="总体思路：",
    )
    add_body(
        doc,
        "球杆沿车体纵向中心线布置，一端固定铰接，另一端由步进电机、摇臂和连杆抬升。MaixCAM安装在球杆上方，视场覆盖钢珠全部有效行程，并与球杆支座分别加强，避免执行机构振动改变相机外参。底盘动力、逻辑、视觉和步进驱动采用分支供电并统一共地；启动、停车、回零和故障状态由软件统一管理。杆长、铰点高度、相机高度及整车外廓在机械定型后填入尺寸表。",
        bold_label="机械布局：",
    )
    make_architecture_image()
    add_picture(doc, ARCH_IMAGE, "图1  系统功能边界与信号流")
    add_table(
        doc,
        "表1  系统模块与接口分工",
        ["模块", "输入", "处理内容", "输出"],
        [
            ["底盘感知", "8路红外、IMU、编码器", "形成横向偏差、偏航角速度和轮速", "底盘状态量"],
            ["底盘控制", "底盘状态量、发车模式", "三级闭环、调速、里程和终点判定", "左右电机PWM"],
            ["钢珠视觉", "MaixCAM图像", "仅做钢珠识别、跟踪、标定与有效性判断", ["球位", ("math", "x"), "、", ("math", "valid"), "、时间戳"]],
            ["滚球控制", "目标位置与有效球位", "位置PD、限幅、斜率限制、机构映射", "球杆角和电机角命令"],
            ["机械执行", "电机绝对位置命令", "步进电机带动摇臂连杆升降杆端", "球杆倾角变化"],
        ],
        [2.4, 3.3, 5.2, 3.1],
    )

    add_heading(doc, "2、小车循迹方案论证", 2)
    add_body(
        doc,
        "少量开关式红外配合分段转向逻辑实现简单，但在连续弯道上输出离散、可用横向分辨率低；灰度模拟阵列能提供更连续的反射强度，但需要模拟采样、逐通道标定和环境光补偿。现有底盘已完成8路数字红外、IMU与双编码器接口，采用位置加权能够在保留较低计算量的同时获得近似连续偏差，因此选择8路数字红外阵列作为循迹传感器。",
        bold_label="方案比较：",
    )
    add_body(
        doc,
        "最终方案把8路黑线状态按横向位置加权，先由位置PD得到目标偏航角速度，再由陀螺仪偏航PI生成左右轮差速，最后由双轮编码器速度PID输出PWM。相比直接由红外状态切换左右电机，该结构能把路径误差、车体旋转和轮速扰动分别限制在对应环路中，且与现有工程的5 ms调度周期一致。正式比赛前仍需在实际黑线宽度、传感器高度和照度下复核有效极性、权重和脱线恢复方向。",
        bold_label="选择结论：",
    )
    add_table(
        doc,
        "表2  循迹方案比较",
        ["比较项", "分段转向", "模拟灰度阵列", "8路数字红外加权"],
        [
            ["偏差连续性", "低", "高", "中等，可由多点加权改善"],
            ["硬件与调试量", "低", "较高", "现有硬件可直接使用"],
            ["环境适应性", "依赖阈值", "需逐路标定", "需核对数字阈值和安装高度"],
            ["与现有代码一致性", "低", "低", "高"],
        ],
        [3.0, 3.8, 3.8, 5.0],
    )

    add_heading(doc, "3、钢珠视觉检测与跟踪方案论证", 2)
    add_body(
        doc,
        "本节只定义纯视觉识别与跟踪任务，不包含位置PD或电机控制。MaixCAM应固定曝光或设置受控曝光范围，在覆盖球杆有效行程的ROI内检测钢珠，输出钢珠中心像素、检测置信或质量指标、有效标志和采样时间。相机分辨率、帧率、曝光、安装高度、镜头焦距及ROI边界由视觉组完成后填入表3。",
        bold_label="成像边界：",
    )
    add_body(
        doc,
        "候选方法可比较灰度/颜色阈值加连通域质心、边缘或圆检测，以及轻量学习型检测。阈值质心计算量小、延迟低，但钢珠高光和槽壁反光会改变分割结果；圆检测具有几何约束，但运动模糊或局部遮挡时轮廓不完整；学习型方法对复杂背景更稳健，但需要数据集、训练和更高推理开销。最终算法必须由视觉组以实拍数据比较定位重复性、单帧耗时和丢帧率后确定，报告当前不预设结论。",
        bold_label="方法比较：",
    )
    add_body(
        doc,
        "建议的模块流程为：采集图像、裁剪ROI、光照预处理、候选区域提取、几何/面积筛选、中心定位、时间连续性检查、像素到杆上坐标映射、有效性输出。短时漏检时只报告无效或沿用带年龄信息的跟踪状态，不能把上一帧坐标伪装成当前测量；连续丢帧阈值由控制层根据实测帧率设置。此处预留算法流程图、标定曲线和跟踪误差图位置。",
        bold_label="待补内容：",
    )
    add_table(
        doc,
        "表3  视觉识别与跟踪参数预留表",
        ["项目", "最终填写内容", "验证方法"],
        [
            ["成像配置", "分辨率：待填；帧率：待填；曝光：待填；ROI：待填", "保存原始帧并核对全行程覆盖"],
            ["检测流程", "最终算法与关键阈值：待视觉组填写", "用反光、阴影和运动模糊样本比较"],
            ["坐标输出", ["中心像素、杆上位置", ("math", "x"), "、", ("math", "valid"), "、时间戳"], "记录连续输出和时间间隔"],
            ["性能指标", "重复定位误差、平均/最大耗时、丢帧率：待实测", "静态多点与动态往返测试"],
            ["异常处理", "面积/速度门限、遮挡和连续丢帧策略：待填", "人工遮挡与反光干扰复测"],
        ],
        [3.0, 7.5, 5.0],
    )

    add_heading(doc, "4、摆杆执行机构与控制方案论证", 2)
    add_body(
        doc,
        "舵机集成位置闭环，接口简单，但齿轮回差和内部控制参数不可直接调整；直流减速电机需要额外位置传感器与驱动闭环；步进电机便于按脉冲实现绝对位置运动，配合现有串口驱动器可读取位置、回零、停止和失能。结合本队已有驱动代码及小角度精细调节需求，选择步进电机加摇臂连杆机构。失步风险通过回零、位置查询、限幅和故障失能处理。",
        bold_label="执行机构比较：",
    )
    add_body(
        doc,
        "控制结构采用位置PD直接生成球杆目标角。P项根据钢珠与目标位置的偏差给出恢复作用，D项利用误差变化率提供阻尼；随后对角度、角速度和命令变化量限幅，再用实测标定表把球杆角映射为电机绝对角。该控制链属于滚球控制与执行机构，不属于视觉识别章节。",
        bold_label="控制结构：",
    )
    add_picture(doc, MECH_IMAGE, "图2  步进电机摇臂带动球杆单端升降示意")

    add_heading(doc, "二、理论分析与计算", 1)
    add_heading(doc, "1、小车循迹运动学与终点控制", 2)
    add_rich_paragraph(
        doc,
        [
            "设第", ("math", "i"), "路数字红外有效状态为", ("math", "b", "i"), "，对应横向位置权重为", ("math", "w", "i"), "。检测到黑线时", ("math", "b", "i"), "=1，否则为0。多路同时有效时采用加权平均并归一化，得到横向偏差", ("math", "e", "y"), "：",
        ],
    )
    add_formula(doc, [("sub", "e", "y"), "(k)=", ("frac", ["Σ", ("sub", "w", "i"), ("sub", "b", "i")], ["Σ", ("sub", "b", "i")]), "，  |", ("sub", "e", "y"), "|≤1"])
    add_body(doc, "当前代码权重为{-3，-2，-1，-0.1，0.1，1，2，3}。当8路原始状态全同，程序保持上一帧识别结果，避免全白脱线、强光或全黑横线使偏差突变；该策略必须结合实际场地复核，尤其要确认无有效通道时的恢复方向。")
    add_rich_paragraph(doc, ["外层位置PD把", ("math", "e", "y"), "转换为目标偏航角速度", ("math", "ω", "d"), "，离散控制律为："])
    add_formula(doc, [("sub", "ω", "d"), "(k)=", ("sub", "K", "p"), ("sub", "e", "y"), "(k)+", ("sub", "K", "d"), ("frac", [("sub", "e", "y"), "(k)-", ("sub", "e", "y"), "(k-1)"], [("sub", "T", "s")])])
    add_rich_paragraph(doc, ["偏航环以IMU的Z轴角速度", ("math", "ω", "z"), "为反馈，经PI产生差速修正", ("math", "Δv"), "；左右轮目标速度按基准速度", ("math", "v", "b"), "叠加差速后，再由编码器速度PID输出PWM。"])
    add_formula(doc, [("sup", "v", "*"), "L=", ("sub", "v", "b"), "-0.4Δv，  ", ("sup", "v", "*"), "R=", ("sub", "v", "b"), "+0.4Δv"])
    add_body(doc, "现有工程每5 ms读取红外、IMU和编码器。编码器换算初值为3726计数/m，并对轮速做一阶滤波。K1模式用于快速循迹与终点停车，K2模式用于较低速度联调。终点线由最近20个控制帧的滑动窗口判断，即窗口约100 ms；正式版本应加入起步屏蔽或允许终点识别的里程门限，并重新测量预减速点和停车偏差。")
    add_table(
        doc,
        "表4  现有底盘控制初值与待校准项",
        ["控制环节", "代码初值", "报告使用说明"],
        [
            ["循迹位置PD（K1）", [("math", "K", "p"), "=57.0，", ("math", "K", "d"), "=2.1，输出限幅500"], "仅作当前工程起点，需实车复调"],
            ["偏航角速度PI", [("math", "K", "p"), "=17.5，", ("math", "K", "i"), "=15.0，积分限幅80"], "启动前校准陀螺仪零偏"],
            ["双轮速度PID", [("math", "K", "p"), "=1.0，", ("math", "K", "i"), "=80.0，", ("math", "K", "d"), "=0"], "PWM限幅9000，单项增量限幅1000"],
            ["调度周期", "5 ms（200 Hz）", "红外、IMU、编码器及三级控制同帧更新"],
            ["终点窗口", "20帧，约100 ms", "需结合起步屏蔽与里程窗口验证"],
        ],
        [4.2, 5.5, 5.8],
    )

    add_heading(doc, "2、球杆系统动力学模型", 2)
    add_rich_paragraph(doc, ["以球杆水平中心为原点，沿杆方向为", ("math", "x"), "轴，球杆小角度为", ("math", "θ"), "，车体沿杆方向加速度为", ("math", "a", "c"), "。在钢珠纯滚动、无滑动且", ("math", "θ"), "较小时，均匀实心球沿杆方向的近似动力学为："])
    add_formula(doc, [("sup", "x", "¨"), "=", ("frac", ["5"], ["7"]), "(gθ-", ("sub", "a", "c"), ")"])
    add_body(doc, "该式说明球杆角度决定钢珠加速度，而车辆起步、弯道速度变化和制动会通过车体加速度直接形成扰动。位置P控制只能提供恢复作用，若缺少速度相关阻尼，钢珠会越过目标点并往复振荡，因此位置PD中的D项是必要的。实际系统还存在滚动摩擦、槽壁接触、杆面倾斜非线性、步进电机回差和视觉延迟，最终参数必须通过实物阶跃记录整定。")
    add_body(doc, "待机械定型后应测量钢珠质量、半径、球杆有效长度、槽宽、球杆质量与最大允许角，并通过静态倾角试验检查模型符号。上述参数当前均不填写数值，避免把设计估计写成实测结果。")

    add_heading(doc, "3、视觉坐标映射与状态估计", 2)
    add_body(doc, "本节仍属于视觉测量，不执行位置PD。视觉组应在球杆上选取不少于5个已知物理位置，记录对应中心像素，先检验一维线性模型；若残差随位置有系统变化，再采用分段线性或透视校正。设图像沿杆方向像素坐标为u，杆上物理坐标为x，线性标定可写为：")
    add_formula(doc, ["x=", ("sub", "a", "0"), "+", ("sub", "a", "1"), "u"])
    add_rich_paragraph(doc, ["每帧输出建议包含", ("math", "x"), "、", ("math", "valid"), "和采样时刻", ("math", "t", "k"), "。位置平滑或速度估计必须使用真实时间间隔，不能默认视觉帧率恒定。可采用中值滤波抑制孤立异常点，再以一阶低通或", ("math", "α"), "-", ("math", "β"), "滤波估计位置与速度；滤波阶数、门限和延迟由实测决定。"])
    add_table(
        doc,
        "表5  视觉坐标标定与跟踪数据预留",
        ["物理位置/cm", "像素坐标/pixel", "重复测量标准差/cm", "备注"],
        [["待填", "待填", "待填", "左端附近"], ["待填", "待填", "待填", "中间区域"], ["待填", "待填", "待填", "右端附近"]],
        [3.6, 3.8, 4.6, 3.5],
    )

    add_heading(doc, "4、滚球闭环、角度映射与轨迹规划", 2)
    add_rich_paragraph(doc, ["设目标球位为", ("math", "x", "r"), "，视觉测得球位为", ("math", "x"), "，位置误差定义为："])
    add_formula(doc, ["e(k)=", ("sub", "x", "r"), "(k)-x(k)"])
    add_rich_paragraph(doc, ["位置PD根据当前误差和滤波后的误差变化率生成目标球杆角", ("math", "θ", "d"), "："])
    add_formula(doc, [("sub", "θ", "d"), "(k)=sat[ s(", ("sub", "K", "p"), "e(k)+", ("sub", "K", "d"), ("sup", "ė", "f"), "(k)), ±", ("sub", "θ", "max"), "]"])
    add_rich_paragraph(doc, ["其中", ("math", "s"), "为由实物正方向确定的符号，", ("math", "θ", "max"), "为安全角度限幅。输出还应经过位置死区、角度斜率限制和电机命令死区。目标位置切换采用斜坡或S形轨迹，避免阶跃指令直接产生最大倾角。视觉短时无效时保持带超时的安全命令，超过门限后清除PD历史并缓慢回到水平。"])
    add_rich_paragraph(doc, ["设电机相对水平零点角为", ("math", "φ"), "，摇臂有效半径为", ("math", "r"), "，安装相位为", ("math", "φ", "0"), "，固定铰点到驱动端距离为", ("math", "L"), "。忽略连杆偏置时："])
    add_formula(doc, ["Δh=r[sin(φ+", ("sub", "φ", "0"), ")-sin", ("sub", "φ", "0"), "]，  θ≈arctan", ("frac", ["Δh"], ["L"])])
    add_body(doc, "实际机构存在连杆长度、连接孔偏置和回差，控制程序不直接使用理论比例，而是用电子水平仪记录至少5组“球杆实测角－电机绝对角”，正负方向分别测量，再以分段线性插值求电机目标角。当前{-2°对应-6°，0°对应0°，+2°对应+6°}只作为占位，不能写成最终标定结果。")
    add_formula(doc, [("sub", "T", "m"), "≥", ("frac", ["SFr"], ["η"])])
    add_body(doc, "转矩校核中，F为驱动端等效竖向载荷，r为摇臂半径，η为传动效率，S为安全系数。F应包含球杆自重、钢珠最不利位置重力分量、连杆惯性和车辆振动附加载荷。待机构尺寸和电机型号确定后再代入数值。")

    add_heading(doc, "三、电路与程序设计", 1)
    add_heading(doc, "1、控制电路与硬件选择", 2)
    add_body(doc, "底盘以MSPM0G3507为控制器，8路红外OUT1至OUT8依次接PA27、PA26、PA25、PA24、PB25、PB24、PB20和PA22；IMU用于Z轴角速度反馈，左右编码器提供轮速和里程。MaixCAM负责相机采集、钢珠位置输出、滚球PD和步进电机串口命令。两控制器任务独立，不要求底盘等待视觉帧。")
    add_body(doc, "步进电机驱动器通过UART与MaixCAM连接，当前代码使用/dev/ttyS0、115200 bit/s、电机地址0x01和校验字节0x6B。必须确认MaixCAM与驱动器共地、串口电平兼容和收发方向；电机电源与逻辑/视觉电源分支供电，并在驱动器附近配置足够的去耦和浪涌余量。最终原理图应补充电池电压、稳压器型号、各支路峰值电流、保险或反接保护。")
    add_table(
        doc,
        "表6  步进电机串口协议与安全约束",
        ["功能", "当前命令/参数", "进入下一状态的条件"],
        [
            ["通信检测", "0x36读取单圈位置", "收到完整且校验正确的响应帧"],
            ["保存水平零点", "0x93，仅人工调平后执行一次", "确认球杆真实水平，正常上电禁止重写"],
            ["回到已存零点", "0x9A，最长等待10 s", "连续2次进入1.5°容差"],
            ["绝对位置运动", "0xFD，当前3200脉冲/圈", "球杆角和电机角双重限幅均通过"],
            ["停止/失能", "0xFE / 0xF3", "异常退出、回零失败或越限立即执行"],
        ],
        [3.3, 5.3, 6.0],
    )

    add_heading(doc, "2、程序结构与任务状态机", 2)
    add_body(doc, "底盘5 ms任务顺序为：读取8路红外并生成偏差、位置PD计算目标偏航角速度、滑动窗口判断终点线、偏航PI计算差速、更新编码器并执行左右轮速度PID，最后按分频周期处理发车加速或里程减速。统一停车函数负责清除控制使能、目标速度和PID历史，并锁存运行时间。")
    add_body(doc, "MaixCAM程序划分为视觉采集、球位数据检查、滚球控制、机构映射和电机通信五个逻辑模块。视觉任务只写入带时间戳的测量结构；控制任务先检查valid和数据年龄，再计算位置PD；执行任务完成球杆角限幅、角度斜率限制、电机角插值和绝对位置发送。任何UART、回零、超时或越限故障都转入FAULT状态，停止并失能电机。")
    add_table(
        doc,
        "表7  整机状态机概要",
        ["状态", "主要动作", "正常转移条件", "异常处理"],
        [
            ["INIT", "初始化底盘、相机、UART和参数", "各接口初始化完成", "记录故障并保持电机停止"],
            ["HOME", "查询位置、使能并回已存水平零点", "连续两次进入零位容差", "超时则停止并失能"],
            ["READY", "等待模式与发车命令", "启动命令且传感器有效", "禁止输出运动命令"],
            ["RUN", "底盘循迹、视觉测量、位置PD和步进执行", "终点判定或人工停止", "丢帧超时回水平；通信故障转FAULT"],
            ["STOP", "底盘停车、球杆缓慢回水平、锁存时间", "重新确认后回READY", "保持安全输出"],
            ["FAULT", "停止并失能，显示故障原因", "人工排故和重新初始化", "禁止自动恢复运动"],
        ],
        [2.2, 5.3, 4.5, 4.0],
    )

    add_heading(doc, "3、标定、限幅与故障处理", 2)
    add_body(doc, "联调顺序为：机械空载检查、确认电机正负方向、人工建立水平零点、完成球杆角与电机角标定、验证视觉坐标正方向和标定、在静态底盘上整定位置PD、低速循迹、最后进行整车动态滚球。每一步只改变一类参数并保存日志，上一环节未通过零位、限位和重复性检查前不得进入下一环节。")
    add_table(
        doc,
        "表8  标定、限幅与故障处理清单",
        ["项目", "需要记录的量", "作用", "当前状态"],
        [
            ["8路红外", "横向位置、离地高度、黑白电平、环境光", "确定权重、极性和脱线策略", "代码初值已有，待实测"],
            ["编码器", "实跑1 m左右轮计数", "修正3726计数/m", "待实测"],
            ["视觉标定", "像素－厘米标定点、重复性和帧间隔", "生成有效球位和超时门限", "预留给视觉组"],
            ["机构映射", "正负方向至少5组球杆角－电机角", "替换占位插值表", "待实测"],
            ["软件限幅", "最大球杆角、电机角、角速度和命令死区", "防止碰撞、冲击和频繁动作", "保守初值，待验证"],
            ["故障策略", "丢帧、UART超时、回零失败、越限", "停止、回水平或失能", "流程已定义"],
        ],
        [3.2, 5.2, 4.8, 2.8],
    )

    add_heading(doc, "四、测试方案与测试结果", 1, page_break=True)
    add_heading(doc, "1、测试条件与仪器", 2)
    add_body(doc, "测试应记录场地尺寸与黑线宽度、环境照度、电池电压、整车质量、球杆尺寸、钢珠规格、MaixCAM成像参数、底盘和滚球控制周期。建议使用卷尺或钢直尺、电子水平仪、秒表、万用表、上位机串口日志及完整录像。当前尚未取得的条件均在表中标为待测。")
    add_table(
        doc,
        "表9  测试条件记录表",
        ["项目", "记录值", "仪器/来源"],
        [
            ["场地、黑线与环境照度", "待测", "卷尺、照度计或现场记录"],
            ["电池空载/带载电压", "待测", "万用表"],
            ["球杆有效长度、摇臂半径、最大行程", "待测", "钢直尺、电子水平仪"],
            ["MaixCAM分辨率、帧率、曝光、ROI", "待视觉组填写", "程序配置与日志"],
            ["控制参数与软件版本", "待最终烧录后填写", "Git提交号或文件校验值"],
        ],
        [5.0, 5.0, 6.0],
    )

    add_heading(doc, "2、测试步骤与数据处理", 2)
    add_body(doc, "每项测试均先确认零位、限位和传感器有效，再从规定初始球位或车辆姿态开始。记录模式、按键动作、计时起止事件、目标位置、球位、球杆目标角、电机角、8路红外状态、左右轮目标/实测速度和PWM。失败试验不得删除，应注明脱轨、掉球、丢帧或通信异常发生时刻。")
    add_body(doc, "同一工况建议重复不少于5次，分别给出均值、最大值、标准差和成功次数。位置误差按杆上统一正方向计算，停车误差按终点线与车体指定基准点的距离计算；异常值只有在能给出明确硬件或记录故障证据时才可剔除，并保留原始数据索引。")

    add_heading(doc, "3、分项测试记录", 2)
    add_table(
        doc,
        "表10  底盘循迹与停车测试预留表",
        ["工况", "次数", "运行时间/s", "停车误差/cm", "是否脱轨", "备注"],
        [["低速联调", "1～5", "待测", "待测", "待记录", ""], ["快速模式", "1～5", "待测", "待测", "待记录", ""], ["电池低电压工况", "1～5", "待测", "待测", "待记录", ""]],
        [3.0, 2.0, 3.0, 3.2, 2.8, 3.0],
    )
    add_table(
        doc,
        "表11  纯视觉识别与跟踪测试预留表",
        ["测试项", "位置/工况", "定位误差/cm", "单帧耗时/ms", "丢帧率/%", "结论"],
        [["静态重复性", "左/中/右各点", "待视觉组实测", "待测", "待测", ""], ["动态往返", "不同速度", "待视觉组实测", "待测", "待测", ""], ["反光/遮挡", "典型干扰", "待视觉组实测", "待测", "待测", ""]],
        [2.8, 3.2, 3.2, 3.2, 2.8, 2.8],
    )
    add_table(
        doc,
        "表12  滚球控制与机械机构测试预留表",
        ["目标位置", "初始位置", "到位时间/s", "稳态最大误差/cm", "峰值球杆角/°", "是否触发保护"],
        [["中心", "待填", "待测", "待测", "待测", "待记录"], ["正向指定点", "待填", "待测", "待测", "待测", "待记录"], ["负向指定点", "待填", "待测", "待测", "待测", "待记录"], ["整车循迹扰动", "待填", "待测", "待测", "待测", "待记录"]],
        [2.8, 2.8, 3.0, 4.0, 3.6, 3.4],
    )

    add_heading(doc, "4、结果判定与误差分析", 2)
    add_body(doc, "正式测试完成后，先逐项判断循迹、停车、钢珠识别和指定位置控制是否满足题目要求，再分析最大误差对应的原始日志。误差来源应按“来源－影响路径－证据－改进”展开：例如视觉反光造成中心偏移，可由原始帧和候选区域面积变化验证；步进机构回差造成正反向映射不一致，可由同一角度双向逼近试验验证；车辆起步和制动造成球位偏移，可由编码器加速度、目标球杆角和球位时间序列对齐验证。")
    add_body(doc, "本稿不填写任何成功率、误差或时间结果。待测试后应把表10至表12中的“待测”替换为原始记录统计值，并补充对应录像、日志文件名及时间段。")

    add_heading(doc, "五、结论与改进", 1)
    add_body(doc, "本文完成了与本队MaixCAM方案一致的底盘与机械执行部分设计。底盘采用8路数字红外位置加权、偏差PD、IMU偏航PI和双轮速度PID实现循迹，并包含发车、分段调速、里程与终点停车逻辑；滚球机构采用一端铰接球杆、步进电机摇臂和连杆完成杆端升降，MaixCAM在获得有效球位后由位置PD生成球杆角，通过标定映射和UART绝对位置命令驱动机构。视觉章节已明确只承担钢珠识别、跟踪、标定与数据有效性输出。")
    add_body(doc, "下一阶段应优先完成球杆尺寸和转矩校核、机械零位与双向角度标定、视觉多点坐标标定及丢帧测试，再在静态底盘上整定位置PD，最后开展低速和快速整车联调。最终报告需用实测数据替换全部待填项，并同步代码中的红外权重、调度周期、终点窗口与控制参数。")

    add_heading(doc, "六、参考文献", 1)
    add_body(doc, "[1] 全国大学生电子设计竞赛组织委员会. 2026年全国大学生电子设计竞赛H题：车载平衡滚球运动控制系统[Z]. 2026.")
    add_body(doc, "[2] Texas Instruments. MSPM0G3507 Mixed-Signal Microcontrollers Data Sheet[EB/OL].")
    add_body(doc, "[3] Sipeed. MaixCAM Documentation[EB/OL].")
    add_body(doc, "[4] 本队所用步进电机驱动器串口协议说明书[Z]. 最终提交前按实物型号补齐版本信息.")

    add_heading(doc, "附录：提交材料与核心程序索引", 1, page_break=True)
    add_table(
        doc,
        "表13  核心程序与正文对应关系",
        ["功能", "文件/符号", "与正文对应"],
        [
            ["8路红外与横线检测", [("code", "Code/Middle/mid_line.c"), "：", ("code", "MID_IR_ReadRaw"), "、", ("code", "MID_Line_CalcError"), "、", ("code", "MID_Line_CheckEmergency")], "二-1、三-2"],
            ["底盘三级闭环", [("code", "Code/Middle/mid_chassis.c"), "：", ("code", "MID_Chassis_PosStep"), "、", ("code", "YawStep"), "、", ("code", "SpeedStep")], "二-1"],
            ["PID实现与初值", [("code", "Code/Middle/mid_pid.c")], "表4"],
            ["5 ms任务调度", [("code", "Code/App/app_scheduler.c")], "三-2"],
            ["发车、里程与停车", [("code", "Code/Middle/mid_key.c"), "、", ("code", "mid_encoder.c")], "二-1、三-2"],
            ["步进电机串口", [("code", "main.py"), "：", ("code", "MotorController"), "、", ("code", "pipe_angle_to_motor_angle")], "二-4、三-1"],
            ["视觉识别与跟踪", "由视觉组补充最终文件名、函数名与版本", "一-3、二-3"],
            ["位置PD", "按最终实现补充函数名；不得归入视觉识别函数", "二-4、三-2"],
        ],
        [3.4, 8.2, 3.0],
    )
    add_body(doc, "提交前自检：更新Word目录和页码；确认标题均为黑色宋体；检查公式变量的斜体与上下标；搜索并清除所有“待填/待测”或用真实数据替换；核对图表编号、单位、代码路径、驱动器型号和参考文献版本；确认视觉识别章节不包含位置PD、电机角映射和串口执行内容。")


def final_format_pass(doc):
    for paragraph in doc.paragraphs:
        has_drawing = bool(paragraph._p.findall(".//" + qn("w:drawing")))
        if has_drawing:
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        elif paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            level = int(paragraph.style.name[-1])
            for run in paragraph.runs:
                set_run_font(run, size=14 if level == 1 else 12, bold=True)
        elif paragraph.style.name == "Normal":
            paragraph.paragraph_format.alignment = paragraph.alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = Pt(22)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def main():
    doc = Document(SOURCE)
    configure_document_styles(doc)
    replace_front_matter(doc)
    remove_old_body(doc)
    build_body(doc)
    final_format_pass(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"OUTPUT={OUTPUT}")
    print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} IMAGES={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
