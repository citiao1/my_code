from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(
    r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_公式优化版.docx"
)
OFFICIAL = Path(
    r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_车载平衡滚球运动控制系统_K230-YOLO版.docx"
)
OUTPUT = Path(
    r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_封面目录分页版.docx"
)


def set_run_font(run, name, size, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_cover_paragraph(doc, before, after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def first_cover_image_blob(path):
    official_doc = Document(path)
    for paragraph in official_doc.paragraphs[:5]:
        blips = paragraph._p.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        )
        if not blips:
            continue
        blip = blips[0]
        relationship_id = blip.get(qn("r:embed"))
        if relationship_id:
            return official_doc.part.related_parts[relationship_id].blob
    raise RuntimeError("Could not locate the official contest logo on the cover")


doc = Document(SOURCE)
body = doc.element.body
children = list(body)
if len(children) < 7:
    raise RuntimeError("Unexpectedly short source document")

# The current draft starts with title, subtitle, stage-scope table, spacer,
# abstract, and keywords. Move the reusable pieces into the new front matter.
old_title, old_subtitle, stage_table, old_spacer, abstract_p, keywords_p = children[:6]
for element in (old_title, old_subtitle, stage_table, old_spacer, abstract_p, keywords_p):
    body.remove(element)

# Build the official-style cover using real paragraphs and the logo from the
# supplied H-problem template.
competition = add_cover_paragraph(doc, before=92, after=38)
set_run_font(competition.add_run("2026年全国大学生电子设计竞赛"), "宋体", 24)

title = add_cover_paragraph(doc, before=0, after=0)
set_run_font(
    title.add_run("车载平衡滚球运动控制系统"),
    "宋体",
    18,
    bold=True,
    color=RGBColor(255, 0, 0),
)
set_run_font(title.add_run("（H题）"), "宋体", 18, bold=True, color=RGBColor(0, 0, 0))

logo = add_cover_paragraph(doc, before=105, after=0)
logo_run = logo.add_run()
logo_run.add_picture(BytesIO(first_cover_image_blob(OFFICIAL)), width=Inches(1.82))

date = add_cover_paragraph(doc, before=142, after=0)
set_run_font(date.add_run("2026年8月1日"), "宋体", 16, bold=True)
date.add_run().add_break(WD_BREAK.PAGE)

# Abstract page. Keep the existing report text and scope callout.
front_subtitle = doc.add_paragraph()
front_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
front_subtitle.paragraph_format.space_before = Pt(44)
front_subtitle.paragraph_format.space_after = Pt(24)
set_run_font(
    front_subtitle.add_run("MaixCAM方案：底盘控制与机械机构阶段初稿"),
    "微软雅黑",
    15,
    bold=True,
    color=RGBColor(38, 92, 116),
)

abstract_wrapper = doc.add_paragraph()
abstract_wrapper._element.addnext(abstract_p)
abstract_wrapper._element.getparent().remove(abstract_wrapper._element)
abstract_paragraph = next(p for p in doc.paragraphs if p._p is abstract_p)
abstract_paragraph.paragraph_format.space_before = Pt(0)
abstract_paragraph.paragraph_format.space_after = Pt(8)
abstract_paragraph.paragraph_format.line_spacing = 1.5

keywords_wrapper = doc.add_paragraph()
keywords_wrapper._element.addnext(keywords_p)
keywords_wrapper._element.getparent().remove(keywords_wrapper._element)
keywords_paragraph = next(p for p in doc.paragraphs if p._p is keywords_p)
keywords_paragraph.paragraph_format.space_before = Pt(0)
keywords_paragraph.paragraph_format.space_after = Pt(12)
keywords_paragraph.paragraph_format.line_spacing = 1.5

stage_spacer = doc.add_paragraph()
stage_spacer.paragraph_format.space_before = Pt(2)
stage_spacer.paragraph_format.space_after = Pt(0)
stage_spacer._element.addnext(stage_table)

toc_page_break = doc.add_paragraph()
toc_page_break.paragraph_format.space_after = Pt(0)
toc_page_break.add_run().add_break(WD_BREAK.PAGE)

toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_title.paragraph_format.space_before = Pt(54)
toc_title.paragraph_format.space_after = Pt(20)
set_run_font(toc_title.add_run("目  录"), "宋体", 24)

toc_placeholder = doc.add_paragraph("[[TOC]]")
toc_placeholder.paragraph_format.space_before = Pt(0)
toc_placeholder.paragraph_format.space_after = Pt(0)

# Insert a next-page section break. The break paragraph's section properties
# describe the front matter; the final body sectPr remains the body section.
section_break = doc.add_paragraph()
section_ppr = section_break._p.get_or_add_pPr()
body_sect_pr = body.sectPr
front_sect_pr = deepcopy(body_sect_pr)
for child in list(front_sect_pr):
    if child.tag in {
        qn("w:headerReference"),
        qn("w:footerReference"),
        qn("w:pgNumType"),
        qn("w:titlePg"),
    }:
        front_sect_pr.remove(child)
section_type = front_sect_pr.find(qn("w:type"))
if section_type is None:
    section_type = OxmlElement("w:type")
    front_sect_pr.insert(0, section_type)
section_type.set(qn("w:val"), "nextPage")
section_ppr.append(front_sect_pr)

# Body pages restart at 1 and retain the original running header/footer.
page_number_type = body_sect_pr.find(qn("w:pgNumType"))
if page_number_type is None:
    page_number_type = OxmlElement("w:pgNumType")
    body_sect_pr.append(page_number_type)
page_number_type.set(qn("w:start"), "1")

# Move all newly built front-matter elements before the first body heading.
front_elements = [
    competition._p,
    title._p,
    logo._p,
    date._p,
    front_subtitle._p,
    abstract_p,
    keywords_p,
    stage_spacer._p,
    stage_table,
    toc_page_break._p,
    toc_title._p,
    toc_placeholder._p,
    section_break._p,
]
for element in front_elements:
    if element.getparent() is not None:
        element.getparent().remove(element)
for index, element in enumerate(front_elements):
    body.insert(index, element)

# Keep the appendix on its own page.
for paragraph in doc.paragraphs:
    if paragraph.text.strip().startswith("附录A"):
        paragraph.paragraph_format.page_break_before = True
        break
else:
    raise RuntimeError("Appendix heading not found")

# TOC styles are explicit so Word field updates render consistently.
for name, left_indent in (("TOC 1", 0.0), ("TOC 2", 0.28)):
    if name in [style.name for style in doc.styles]:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    fmt = style.paragraph_format
    fmt.left_indent = Inches(left_indent)
    fmt.first_line_indent = Inches(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.2
    fmt.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(f"OUTPUT={OUTPUT}")
