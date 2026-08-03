from pathlib import Path

from docx import Document


ROOT = Path(r"D:\my_code\my_code\diansai\H题模板")


for path in sorted(ROOT.glob("*.docx")):
    print(f"\n===== {path.name} =====")
    doc = Document(path)
    print(
        f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} "
        f"sections={len(doc.sections)}"
    )
    shown = 0
    for index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        print(f"P{index:03d} [{paragraph.style.name}] {text[:260]}")
        shown += 1
        if shown >= 220:
            break

    for table_index, table in enumerate(doc.tables[:10]):
        print(f"TABLE {table_index}: {len(table.rows)}x{len(table.columns)}")
        for row in table.rows[:10]:
            cells = [" ".join(cell.text.split())[:120] for cell in row.cells]
            print(" | ".join(cells))
