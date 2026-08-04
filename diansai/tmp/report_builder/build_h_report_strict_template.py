from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "电赛报告"
SOURCE_REPORT = REPORT_DIR / "2026年H题_车载平衡滚球运动控制系统_设计报告初稿.docx"
TEMPLATE = Path(__file__).resolve().parent / "strict_template_source.docx"
INTERMEDIATE = Path(__file__).resolve().parent / "strict_template_intermediate.docx"
OUTPUT = REPORT_DIR / "2026年H题_车载平衡滚球运动控制系统_设计报告（严格模板版）.docx"

FONT_BODY = "宋体"
FONT_H1 = "幼圆"
FONT_H2 = "楷体"
FONT_H3 = "黑体"
FONT_LATIN = "Times New Roman"
BLACK = RGBColor(0, 0, 0)

BODY_WIDTH_DXA = 8307
TABLE_INDENT_DXA = 120
TABLE_WIDTH_DXA = BODY_WIDTH_DXA - TABLE_INDENT_DXA


def paragraph_text(element) -> str:
    return "".join(element.xpath(".//w:t/text()"))


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink"), qn("w:fldSimple")}:
            paragraph._p.remove(child)


def clone_run_with_text(sample_run, text: str):
    run = deepcopy(sample_run._r)
    for child in list(run):
        if child.tag != qn("w:rPr"):
            run.remove(child)
    text_node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    return run


def first_run_with_font(paragraph, font_name: str):
    for run in paragraph.runs:
        r_pr = run._element.rPr
        if r_pr is None or r_pr.rFonts is None:
            continue
        if r_pr.rFonts.get(qn("w:eastAsia")) == font_name:
            return run
    return paragraph.runs[-1]


def set_cover_line(paragraph, text: str, sample_run=None) -> None:
    if sample_run is None:
        sample_run = paragraph.runs[0]
    sample = deepcopy(sample_run._r)
    remove_runs(paragraph)
    class RunProxy:
        _r = sample
    paragraph._p.append(clone_run_with_text(RunProxy(), text))


def set_cover_field(paragraph, label: str, value: str) -> None:
    label_run = paragraph.runs[0]
    value_run = first_run_with_font(paragraph, "楷体_GB2312")
    label_xml = deepcopy(label_run._r)
    value_xml = deepcopy(value_run._r)
    remove_runs(paragraph)

    class LabelProxy:
        _r = label_xml

    class ValueProxy:
        _r = value_xml

    paragraph._p.append(clone_run_with_text(LabelProxy(), label))
    paragraph._p.append(clone_run_with_text(ValueProxy(), value))


def prepare_cover(template_doc: Document):
    paragraphs = template_doc.paragraphs[:17]
    set_cover_line(paragraphs[4], "2026年全国大学生电子设计竞赛赛区赛（TI杯）")
    set_cover_field(paragraphs[12], "题        目：", "      车载平衡滚球运动控制系统（H题）")
    set_cover_field(paragraphs[13], "学        校：", "                                      ")
    set_cover_field(paragraphs[14], "指 导  老 师：", "                                     ")
    set_cover_field(paragraphs[15], "参赛队员姓名：", "                                    ")
    set_cover_field(paragraphs[16], "日        期：", "          2026年08月  日")
    return paragraphs


def clean_section_references(sect_pr) -> None:
    for child in list(sect_pr):
        if child.tag in {qn("w:headerReference"), qn("w:footerReference"), qn("w:titlePg")}:
            sect_pr.remove(child)
    type_node = sect_pr.find(qn("w:type"))
    if type_node is None:
        type_node = OxmlElement("w:type")
        sect_pr.insert(0, type_node)
    type_node.set(qn("w:val"), "nextPage")


def install_cover(report_doc: Document, template_doc: Document) -> None:
    body = report_doc._element.body
    abstract_element = None
    for child in list(body):
        if child.tag == qn("w:p") and paragraph_text(child).strip() == "摘  要":
            abstract_element = child
            break
    if abstract_element is None:
        raise RuntimeError("Could not find the abstract heading in the source report")

    for child in list(body):
        if child is abstract_element:
            break
        body.remove(child)

    cover_paragraphs = prepare_cover(template_doc)
    cover_elements = [deepcopy(paragraph._p) for paragraph in cover_paragraphs]

    cover_sect_pr = deepcopy(template_doc.sections[0]._sectPr)
    clean_section_references(cover_sect_pr)
    last_p_pr = cover_elements[-1].get_or_add_pPr()
    old_sect_pr = last_p_pr.find(qn("w:sectPr"))
    if old_sect_pr is not None:
        last_p_pr.remove(old_sect_pr)
    last_p_pr.append(cover_sect_pr)

    for index, element in enumerate(cover_elements):
        body.insert(index, element)


def set_font(run, east_asia: str, size: float, bold=None) -> None:
    run.font.name = FONT_LATIN
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def clear_color_and_shading(element) -> None:
    p_pr = element.find(qn("w:pPr"))
    if p_pr is not None:
        for tag in ("w:shd", "w:pBdr"):
            node = p_pr.find(qn(tag))
            if node is not None:
                p_pr.remove(node)


def set_style_font(style, east_asia: str, size: float, bold: bool, alignment,
                   before: float, after: float, left: float = 0) -> None:
    style.font.name = FONT_LATIN
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    fmt.alignment = alignment
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(left)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.keep_with_next = True
    fmt.keep_together = True
    fmt.widow_control = True


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, FONT_BODY, 10.5, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 0)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(18)
    normal.paragraph_format.keep_with_next = False
    normal.paragraph_format.keep_together = False

    set_style_font(doc.styles["Heading 1"], FONT_H1, 18, True,
                   WD_ALIGN_PARAGRAPH.CENTER, 12, 6)
    set_style_font(doc.styles["Heading 2"], FONT_H2, 15, False,
                   WD_ALIGN_PARAGRAPH.LEFT, 10, 4, 21)
    set_style_font(doc.styles["Heading 3"], FONT_H3, 14, False,
                   WD_ALIGN_PARAGRAPH.LEFT, 8, 3, 21)
    set_style_font(doc.styles["Heading 4"], FONT_H3, 12, False,
                   WD_ALIGN_PARAGRAPH.LEFT, 6, 2, 21)
    set_style_font(doc.styles["Caption"], FONT_BODY, 9, False,
                   WD_ALIGN_PARAGRAPH.CENTER, 2, 4)
    doc.styles["Caption"].paragraph_format.keep_with_next = True

    for name in ("TOC 1", "TOC 2", "TOC 3"):
        if name in doc.styles:
            set_style_font(doc.styles[name], FONT_BODY, 10.5, False,
                           WD_ALIGN_PARAGRAPH.LEFT, 0, 0)
            doc.styles[name].paragraph_format.keep_with_next = False
            doc.styles[name].paragraph_format.keep_together = False


def style_body_paragraphs(doc: Document) -> None:
    body_started = False
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "摘  要":
            body_started = True
        if not body_started:
            continue
        if text.startswith("目录提示"):
            remove_element(paragraph._p)
            continue

        clear_color_and_shading(paragraph._p)
        style_name = paragraph.style.name
        if style_name == "Heading 1":
            east_asia, size, bold = FONT_H1, 18, True
        elif style_name == "Heading 2":
            east_asia, size, bold = FONT_H2, 15, False
        elif style_name == "Heading 3":
            east_asia, size, bold = FONT_H3, 14, False
        elif style_name == "Heading 4":
            east_asia, size, bold = FONT_H3, 12, False
        elif style_name == "Caption":
            east_asia, size, bold = FONT_BODY, 9, False
        elif style_name.startswith("TOC"):
            east_asia, size, bold = FONT_BODY, 10.5, False
        else:
            east_asia, size, bold = FONT_BODY, 10.5, None

        for run in paragraph.runs:
            set_font(run, east_asia, size, bold if style_name.startswith("Heading") else None)

        if text in {"摘  要", "目  录"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in paragraph.runs:
                set_font(run, FONT_H3, 16, True)


def next_id(root, path: str, attribute: str) -> int:
    values = []
    for node in root.findall(qn(path)):
        value = node.get(qn(attribute))
        if value is not None:
            values.append(int(value))
    return max(values, default=0) + 1


def add_multilevel_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level, pattern, left in ((0, "%1", 0), (1, "%1.%2", 420), (2, "%1.%2.%3", 420)):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        for tag, attr, value in (
            ("w:start", "w:val", "1"),
            ("w:numFmt", "w:val", "decimal"),
            ("w:lvlText", "w:val", pattern),
            ("w:suff", "w:val", "space"),
        ):
            node = OxmlElement(tag)
            node.set(qn(attr), value)
            lvl.append(node)
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), "0")
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def fix_heading_numbering(doc: Document) -> None:
    num_id = add_multilevel_numbering(doc)
    body_started = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "系统方案设计":
            body_started = True
        if not body_started:
            continue
        if paragraph.style.name not in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        p_pr = paragraph._p.get_or_add_pPr()
        old = p_pr.find(qn("w:numPr"))
        if old is None:
            continue
        p_pr.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(int(paragraph.style.name[-1]) - 1))
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.append(num_pr)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_width(parent, tag: str, width: int) -> None:
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    node.set(qn("w:w"), str(width))
    node.set(qn("w:type"), "dxa")


def scale_widths(widths, target):
    total = sum(widths) or len(widths)
    result = [max(400, round(width * target / total)) for width in widths]
    result[-1] += target - sum(result)
    return result


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        set_width(tbl_pr, "w:tblW", TABLE_WIDTH_DXA)
        set_width(tbl_pr, "w:tblInd", TABLE_INDENT_DXA)
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "4")
            node.set(qn("w:color"), "000000")

        grid_cols = list(table._tbl.tblGrid)
        original = [int(col.get(qn("w:w"), "1000")) for col in grid_cols]
        widths = scale_widths(original, TABLE_WIDTH_DXA)
        for col, width in zip(grid_cols, widths):
            col.set(qn("w:w"), str(width))

        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            for height in list(tr_pr.findall(qn("w:trHeight"))):
                tr_pr.remove(height)
            if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))
            for cell, width in zip(row.cells, widths):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                set_width(tc_pr, "w:tcW", width)
                shading = tc_pr.find(qn("w:shd"))
                if shading is not None:
                    tc_pr.remove(shading)
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    clear_color_and_shading(paragraph._p)
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    paragraph.paragraph_format.line_spacing = Pt(13)
                    if row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_font(run, FONT_H3 if row_index == 0 else FONT_BODY, 9,
                                 True if row_index == 0 else None)


def resize_images(doc: Document) -> None:
    max_width = Cm(14.0)
    for shape in doc.inline_shapes:
        if shape.width <= max_width:
            continue
        ratio = max_width / shape.width
        shape.width = int(shape.width * ratio)
        shape.height = int(shape.height * ratio)


def clear_header_footer(container) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)
    element.append(OxmlElement("w:p"))


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    set_font(run, FONT_BODY, 9)


def format_section_header_footer(section) -> None:
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    clear_header_footer(section.header)
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.first_line_indent = Pt(0)
    header_run = header_p.add_run("全国大学生电子设计竞赛  2026年赛区赛（TI杯）")
    set_font(header_run, FONT_BODY, 9)

    clear_header_footer(section.footer)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.first_line_indent = Pt(0)
    add_page_field(footer_p)


def configure_sections(doc: Document) -> None:
    if len(doc.sections) != 2:
        raise RuntimeError(f"Expected 2 sections after cover insertion, found {len(doc.sections)}")
    cover, body = doc.sections
    cover.start_type = WD_SECTION_START.NEW_PAGE
    cover.header_distance = Mm(15)
    cover.footer_distance = Mm(17.5)
    format_section_header_footer(cover)

    body.start_type = WD_SECTION_START.NEW_PAGE
    body.page_width = Mm(210)
    body.page_height = Mm(297)
    body.top_margin = Mm(25.4)
    body.bottom_margin = Mm(25.4)
    body.left_margin = Mm(31.75)
    body.right_margin = Mm(31.75)
    body.header_distance = Mm(15)
    body.footer_distance = Mm(17.5)
    format_section_header_footer(body)


def request_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build() -> Path:
    report = Document(SOURCE_REPORT)
    template = Document(TEMPLATE)
    install_cover(report, template)
    report.save(INTERMEDIATE)

    report = Document(INTERMEDIATE)
    configure_sections(report)
    configure_styles(report)
    style_body_paragraphs(report)
    fix_heading_numbering(report)
    format_tables(report)
    resize_images(report)
    request_field_updates(report)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    report.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
