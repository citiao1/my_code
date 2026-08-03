from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSET_DIR = Path(__file__).resolve().parent / "assets"
OUTPUT_DIR = ROOT / "电赛报告"
OUTPUT_PATH = OUTPUT_DIR / "2026年H题_车载平衡滚球运动控制系统_设计报告初稿.docx"
LOGO_PATH = ROOT / "电赛报告" / "Template-for-Electrical-Competition-Report-main" / "pic" / "电赛logo.png"

NAVY = "17365D"
BLUE = "2F75B5"
TEAL = "167D7F"
ORANGE = "C65911"
RED = "C00000"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EAF1F8"
LIGHT_ORANGE = "FFF4E8"
WHITE = "FFFFFF"
BLACK = "000000"

FONT_BODY = "宋体"
FONT_HEADING_1 = "幼圆"
FONT_HEADING_2 = "楷体"
FONT_HEADING_3 = "黑体"
FONT_LATIN = "Times New Roman"
FONT_MATH = "Cambria Math"

# standard_business_brief preset with a named competition-report override:
# A4 page, 25 mm margins, Chinese five-point body type, 18 pt exact line spacing.
PAGE_WIDTH_DXA = 11907
PAGE_HEIGHT_DXA = 16840
CONTENT_WIDTH_DXA = 9072
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, east_asia=FONT_BODY, latin=FONT_LATIN, size=10.5,
                 bold=None, color=BLACK, italic=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=18,
                          first_line=None, keep_with_next=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if isinstance(line, (int, float)):
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line)
    else:
        fmt.line_spacing = line
    if first_line is not None:
        fmt.first_line_indent = Pt(first_line)
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def style_setup(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(18)
    normal.paragraph_format.widow_control = True

    specs = {
        "Heading 1": (FONT_HEADING_1, 18, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, 14, 8),
        "Heading 2": (FONT_HEADING_2, 15, NAVY, False, WD_ALIGN_PARAGRAPH.LEFT, 12, 6),
        "Heading 3": (FONT_HEADING_3, 14, BLACK, True, WD_ALIGN_PARAGRAPH.LEFT, 9, 4),
        "Heading 4": (FONT_HEADING_3, 12, BLACK, True, WD_ALIGN_PARAGRAPH.LEFT, 7, 3),
    }
    for style_name, (font, size, color, bold, align, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = rgb(color)
        style.paragraph_format.alignment = align
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.left_indent = Pt(0 if style_name == "Heading 1" else 21)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for name, font, size, color, bold, align in [
        ("Caption", FONT_BODY, 9, BLACK, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Quote", FONT_BODY, 9.5, GRAY, False, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = rgb(color)
        style.paragraph_format.alignment = align
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def set_cell_margins(cell, margins=CELL_MARGINS_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in margins.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="A6A6A6", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths, caption=None):
    if caption:
        p = doc.add_paragraph(style="Caption")
        p.paragraph_format.keep_with_next = True
        add_runs(p, caption, size=9, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    repeat_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line=14)
        add_runs(p, header, size=9, bold=True, east_asia=FONT_HEADING_3)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, line=14)
            add_runs(p, str(value), size=9)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    set_paragraph_spacing(after, after=2, line=6)
    return table


PLACEHOLDER_RE = re.compile(r"(\[待(?:填|实测|确认|复核|补图|完善)[^\]]*\])")


def add_runs(paragraph, text, size=10.5, bold=False, color=BLACK,
             east_asia=FONT_BODY, italic=False):
    parts = PLACEHOLDER_RE.split(text)
    for part in parts:
        if not part:
            continue
        is_placeholder = bool(PLACEHOLDER_RE.fullmatch(part))
        run = paragraph.add_run(part)
        set_run_font(
            run,
            east_asia=east_asia,
            size=size,
            bold=True if is_placeholder else bold,
            color=ORANGE if is_placeholder else color,
            italic=italic,
        )
    return paragraph


def add_body(doc, text, bold=False, color=BLACK, first_line=21,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, after=after, line=18, first_line=first_line)
    add_runs(p, text, bold=bold, color=color)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=3, line=18, first_line=0)
    run = p.add_run(text)
    set_run_font(run, east_asia=FONT_BODY, latin=FONT_MATH, size=10.5)
    return p


def paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "6")
            node.set(qn("w:space"), "4")
            node.set(qn("w:color"), border_color)
            p_bdr.append(node)


def add_callout(doc, label, text, fill=LIGHT_ORANGE, color=ORANGE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=4, after=6, line=17, first_line=0)
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    paragraph_shading(p, fill, color)
    r = p.add_run(label + " ")
    set_run_font(r, east_asia=FONT_HEADING_3, size=10, bold=True, color=color)
    add_runs(p, text, size=10, color=BLACK)
    return p


def add_field(paragraph, instruction, display_text=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def next_numbering_id(numbering_part, tag):
    ids = []
    for el in numbering_part.element.findall(qn(tag)):
        attr = "w:abstractNumId" if tag == "w:abstractNum" else "w:numId"
        value = el.get(qn(attr))
        if value is not None:
            ids.append(int(value))
    return max(ids, default=0) + 1


def create_numbering(doc, kind="headings"):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(doc.part.numbering_part, "w:abstractNum")
    num_id = next_numbering_id(doc.part.numbering_part, "w:num")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    if kind == "headings":
        levels = [
            (0, "decimal", "%1", 0, 360, 360),
            (1, "decimal", "%1.%2", 300, 720, 420),
            (2, "decimal", "%1.%2.%3", 300, 900, 600),
        ]
    elif kind == "bullet":
        levels = [(0, "bullet", "•", 360, 720, 360)]
    else:
        levels = [(0, "decimal", "%1.", 360, 720, 360)]

    for level, fmt, pattern, left, text_pos, hanging in levels:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), pattern)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(text_pos))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_pos))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_heading(doc, text, level, heading_num_id=None, numbered=True):
    p = doc.add_paragraph(style=f"Heading {level}")
    if numbered and heading_num_id is not None:
        apply_numbering(p, heading_num_id, level - 1)
    add_runs(
        p,
        text,
        size={1: 18, 2: 15, 3: 14, 4: 12}[level],
        bold=level in (1, 3, 4),
        color=NAVY if level in (1, 2) else BLACK,
        east_asia={1: FONT_HEADING_1, 2: FONT_HEADING_2, 3: FONT_HEADING_3, 4: FONT_HEADING_3}[level],
    )
    return p


def add_list_item(doc, text, num_id, level=0, numbered=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_spacing(p, after=2, line=17, first_line=0)
    apply_numbering(p, num_id, level)
    add_runs(p, text)
    return p


def add_picture(doc, path, width_cm, caption, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Cm(width_cm))
    inline._inline.docPr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.keep_with_next = False
    add_runs(cap, caption, size=9)
    return p, cap


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_section_layout(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(23)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)
    section.different_first_page_header_footer = True


def setup_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=2)
    set_paragraph_spacing(p, line=12, first_line=0)
    r = p.add_run("2026年全国大学生电子设计竞赛")
    set_run_font(r, east_asia=FONT_BODY, size=8.5, color=GRAY)
    r = p.add_run("\tH题设计报告 · 初步调试版")
    set_run_font(r, east_asia=FONT_BODY, size=8.5, color=GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=12, first_line=0)
    r = p.add_run("第 ")
    set_run_font(r, east_asia=FONT_BODY, size=8.5, color=GRAY)
    fld = add_field(p, " PAGE ", "1")
    set_run_font(fld, east_asia=FONT_BODY, size=8.5, color=GRAY)
    r = p.add_run(" 页")
    set_run_font(r, east_asia=FONT_BODY, size=8.5, color=GRAY)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def font(size, bold=False):
    candidates = [
        ("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def fit_text(draw, text, box_width, start_size=36, min_size=22, bold=False):
    for size in range(start_size, min_size - 1, -1):
        f = font(size, bold)
        if draw.textbbox((0, 0), text, font=f)[2] <= box_width:
            return f
    return font(min_size, bold)


def draw_centered_text(draw, box, text, fill="#172238", bold=False, size=34,
                       spacing=8):
    x0, y0, x1, y1 = box
    lines = text.split("\n")
    f = fit_text(draw, max(lines, key=len), x1 - x0 - 30, size, 20, bold)
    heights = [draw.textbbox((0, 0), line, font=f)[3] for line in lines]
    total = sum(heights) + spacing * (len(lines) - 1)
    y = y0 + (y1 - y0 - total) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=f)
        x = x0 + (x1 - x0 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=f, fill=fill)
        y += h + spacing


def block(draw, box, text, fill, outline="#38516D", title_color="#172238",
          radius=10, size=30):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, text, fill=title_color, bold=True, size=size)


def arrow(draw, start, end, color="#51687F", width=5, head=18):
    draw.line([start, end], fill=color, width=width)
    x0, y0 = start
    x1, y1 = end
    dx, dy = x1 - x0, y1 - y0
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    p1 = (x1, y1)
    p2 = (x1 - head * ux + head * 0.55 * px, y1 - head * uy + head * 0.55 * py)
    p3 = (x1 - head * ux - head * 0.55 * px, y1 - head * uy - head * 0.55 * py)
    draw.polygon([p1, p2, p3], fill=color)


def make_system_diagram(path):
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "车载平衡滚球系统总体结构", font=font(48, True), fill="#17365D")
    d.line((70, 105, 1730, 105), fill="#D8E0E8", width=3)

    d.text((90, 145), "底盘循迹与运动控制", font=font(34, True), fill="#17365D")
    blocks = [
        ((90, 220, 340, 350), "8路红外模块\n黑白标定"),
        ((430, 220, 700, 350), "MSPM0G3507\n循迹与轮速控制"),
        ((790, 220, 1030, 350), "AT8236\n双H桥驱动"),
        ((1120, 220, 1380, 350), "MG513X × 2\n差速驱动"),
        ((1470, 220, 1700, 350), "两驱动轮\n+ 万向轮"),
    ]
    fills = ["#EAF1F8", "#DDEBF7", "#FFF0E1", "#E2F0D9", "#F2F2F2"]
    for (box, text), fill in zip(blocks, fills):
        block(d, box, text, fill, size=28)
    for i in range(len(blocks) - 1):
        arrow(d, (blocks[i][0][2], 285), (blocks[i + 1][0][0], 285))
    arrow(d, (1250, 365), (565, 365), color="#167D7F")
    d.text((810, 372), "编码器速度反馈（若使用）", font=font(24), fill="#167D7F")

    d.text((90, 500), "滚球视觉与摆杆控制", font=font(34, True), fill="#17365D")
    lower = [
        ((90, 575, 345, 720), "MaixCAM Pro\n摄像头与显示"),
        ((430, 575, 705, 720), "球位识别\n位置/速度估计"),
        ((790, 575, 1050, 720), "位置-速度\n级联控制器"),
        ((1135, 575, 1400, 720), "42步进电机\n及闭环驱动"),
        ((1480, 575, 1710, 720), "连杆摆杆\n钢球对象"),
    ]
    lower_fills = ["#E2F0D9", "#EAF1F8", "#DDEBF7", "#FFF0E1", "#F2F2F2"]
    for (box, text), fill in zip(lower, lower_fills):
        block(d, box, text, fill, size=28)
    for i in range(len(lower) - 1):
        arrow(d, (lower[i][0][2], 648), (lower[i + 1][0][0], 648))
    arrow(d, (1595, 740), (220, 740), color="#167D7F")
    d.text((730, 750), "摄像头采集球位反馈", font=font(24), fill="#167D7F")

    block(d, (520, 875, 805, 990), "车载电池\n分路稳压/滤波", "#FFF4E8", size=27)
    block(d, (1000, 875, 1285, 990), "按键 + ≤2英寸显示\n计时与状态", "#F2F2F2", size=25)
    arrow(d, (805, 930), (1000, 930), color="#C65911")
    d.text((85, 1005), "注：通信接口、电源电压与具体引脚以最终实物为准。", font=font(22), fill="#666666")
    img.save(path, quality=95)


def make_control_diagram(path):
    img = Image.new("RGB", (1800, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "滚球位置-速度级联控制框图", font=font(48, True), fill="#17365D")
    d.line((70, 105, 1730, 105), fill="#D8E0E8", width=3)
    y = 310
    items = [
        ((70, 245, 245, 375), "目标位置\nx_r"),
        ((340, 245, 565, 375), "位置PI\ne_x=x_r-x"),
        ((660, 245, 890, 375), "目标速度\nv_r"),
        ((985, 245, 1215, 375), "速度PI\ne_v=v_r-v"),
        ((1310, 245, 1515, 375), "目标摆角\ntheta_r"),
        ((1600, 245, 1740, 375), "连杆/球\n对象"),
    ]
    colors = ["#F2F2F2", "#DDEBF7", "#F2F2F2", "#DDEBF7", "#FFF0E1", "#E2F0D9"]
    for (box, text), fill in zip(items, colors):
        block(d, box, text, fill, size=27)
    for i in range(len(items) - 1):
        arrow(d, (items[i][0][2], y), (items[i + 1][0][0], y))

    block(d, (710, 530, 1020, 660), "视觉位置 x\n最小二乘估速 v", "#EAF1F8", size=27)
    arrow(d, (1670, 390), (1670, 595), color="#167D7F")
    arrow(d, (1670, 595), (1020, 595), color="#167D7F")
    arrow(d, (710, 575), (455, 575), color="#167D7F")
    arrow(d, (455, 575), (455, 390), color="#167D7F")
    arrow(d, (865, 530), (865, 410), color="#167D7F")
    d.text((1080, 610), "摄像头反馈", font=font(24), fill="#167D7F")
    d.text((80, 690), "控制约束：速度限幅、摆角限幅、积分抗饱和、角度变化率限制、丢球回水平。", font=font(24), fill="#666666")
    img.save(path, quality=95)


def make_mechanical_diagram(path):
    img = Image.new("RGB", (1800, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "H42摆杆与步进电机连杆机构示意", font=font(48, True), fill="#17365D")
    d.line((70, 105, 1730, 105), fill="#D8E0E8", width=3)

    # Chassis and wheels.
    d.rounded_rectangle((140, 720, 1650, 820), radius=12, fill="#D9E2F3", outline="#17365D", width=4)
    for cx in (420, 1320):
        d.ellipse((cx - 100, 770, cx + 100, 970), fill="#474F59", outline="#111111", width=4)
        d.ellipse((cx - 45, 825, cx + 45, 915), fill="#D9E2F3", outline="#111111", width=3)
    d.ellipse((200, 805, 290, 895), fill="#7F8C99", outline="#111111", width=3)
    d.text((175, 900), "万向轮", font=font(23), fill="#333333")

    # Left hinge tower and tilted rod.
    d.rounded_rectangle((280, 470, 480, 720), radius=8, fill="#EAF1F8", outline="#17365D", width=4)
    d.ellipse((430, 430, 500, 500), fill="#FFF0E1", outline="#17365D", width=4)
    rod_start = (465, 465)
    rod_end = (1430, 390)
    d.line([rod_start, rod_end], fill="#8CA6BF", width=42)
    d.line([rod_start, rod_end], fill="#F7FBFF", width=26)
    d.ellipse((1000, 350, 1080, 430), fill="#7F8C99", outline="#333333", width=4)
    d.text((985, 430), "直径约1 cm钢球", font=font(25), fill="#333333")

    # Stepper and linkage.
    d.rounded_rectangle((1320, 590, 1540, 790), radius=8, fill="#FFF0E1", outline="#C65911", width=4)
    d.ellipse((1375, 635, 1485, 745), fill="#F7D7B8", outline="#C65911", width=4)
    pivot = (1430, 690)
    link_top = (1430, 390)
    d.line([pivot, (1500, 520), link_top], fill="#C65911", width=16)
    d.ellipse((1410, 670, 1450, 710), fill="#FFFFFF", outline="#C65911", width=4)
    d.ellipse((1410, 370, 1450, 410), fill="#FFFFFF", outline="#C65911", width=4)
    d.text((1300, 815), "42步进电机 + 连杆", font=font(26, True), fill="#C65911")

    # Camera and field of view.
    d.rounded_rectangle((780, 150, 1080, 260), radius=8, fill="#E2F0D9", outline="#167D7F", width=4)
    d.ellipse((915, 185, 975, 245), fill="#FFFFFF", outline="#167D7F", width=4)
    d.text((800, 110), "MaixCAM Pro", font=font(28, True), fill="#167D7F")
    d.polygon([(930, 260), (600, 500), (1350, 440)], outline="#7FB3A6")

    # Labels and dimensions.
    d.line((465, 325, 1430, 325), fill="#666666", width=3)
    arrow(d, (475, 325), (465, 325), color="#666666", width=3, head=12)
    arrow(d, (1420, 325), (1430, 325), color="#666666", width=3, head=12)
    d.text((790, 270), "25 cm PPR凹槽摆杆", font=font(28, True), fill="#333333")
    d.line((245, 470, 245, 720), fill="#666666", width=3)
    arrow(d, (245, 705), (245, 720), color="#666666", width=3, head=12)
    arrow(d, (245, 485), (245, 470), color="#666666", width=3, head=12)
    d.text((120, 565), "h ≥ 5 cm", font=font(25), fill="#333333")
    d.text((310, 405), "铰链 / Axis / 625ZZ", font=font(24), fill="#17365D")
    d.text((520, 520), "Top_body / Bottom / Sheet_metal支撑件", font=font(24), fill="#17365D")
    d.text((80, 955), "依据“26年H题车载水平小球”H42装配体的部件关系绘制；最终报告应替换为实物照片和CAD等轴测图。", font=font(22), fill="#666666")
    img.save(path, quality=95)


def make_flow_diagram(path):
    img = Image.new("RGB", (1600, 1120), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "系统软件主流程", font=font(48, True), fill="#17365D")
    d.line((70, 105, 1530, 105), fill="#D8E0E8", width=3)

    nodes = [
        ((565, 140, 1035, 240), "上电初始化\nGPIO / PWM / ADC / UART / Camera"),
        ((565, 295, 1035, 395), "黑白标定、球心零点与像素标定\n步进机构回零并置水平"),
        ((565, 450, 1035, 550), "等待启动按键\n计时清零并选择测试任务"),
        ((140, 650, 600, 775), "底盘周期任务\n读取8路红外 → 归一化\n加权偏差 → 循迹PID → 左右轮输出"),
        ((1000, 650, 1460, 775), "滚球周期任务\n采集ROI → 球位/速度估计\n级联PI → 摆角限幅 → 步进指令"),
        ((565, 880, 1035, 980), "安全与状态处理\n丢线停车 / 丢球回水平 / 显示与图传记录"),
    ]
    fills = ["#EAF1F8", "#FFF0E1", "#F2F2F2", "#DDEBF7", "#E2F0D9", "#FFF4E8"]
    for (box, text), fill in zip(nodes, fills):
        block(d, box, text, fill, size=26)
    arrow(d, (800, 240), (800, 295))
    arrow(d, (800, 395), (800, 450))
    arrow(d, (720, 550), (370, 650))
    arrow(d, (880, 550), (1230, 650))
    arrow(d, (370, 775), (680, 880))
    arrow(d, (1230, 775), (920, 880))
    arrow(d, (800, 980), (800, 1060), color="#167D7F")
    arrow(d, (800, 1060), (1510, 1060), color="#167D7F")
    arrow(d, (1510, 1060), (1510, 605), color="#167D7F")
    arrow(d, (1510, 605), (1400, 605), color="#167D7F")
    d.text((830, 1025), "未结束则继续循环", font=font(22), fill="#167D7F")
    img.save(path, quality=95)


def make_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "system": ASSET_DIR / "system_overview.png",
        "control": ASSET_DIR / "balance_control.png",
        "mechanical": ASSET_DIR / "mechanical_structure.png",
        "flow": ASSET_DIR / "software_flow.png",
    }
    make_system_diagram(paths["system"])
    make_control_diagram(paths["control"])
    make_mechanical_diagram(paths["mechanical"])
    make_flow_diagram(paths["flow"])
    return paths


def cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=14, line=12, first_line=0)
    if LOGO_PATH.exists():
        logo = p.add_run().add_picture(str(LOGO_PATH), width=Cm(4.3))
        logo._inline.docPr.set("descr", "全国大学生电子设计竞赛标志")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=6, line=24, first_line=0)
    add_runs(p, "2026年全国大学生电子设计竞赛赛区赛（TI杯）", size=16, bold=True, east_asia=FONT_HEADING_3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=28, after=12, line=36, first_line=0)
    add_runs(p, "H题  车载平衡滚球\n运动控制系统", size=26, bold=True, color=NAVY, east_asia=FONT_HEADING_3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=28, line=24, first_line=0)
    add_runs(p, "设 计 报 告", size=18, bold=True, east_asia=FONT_HEADING_3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=34, line=18, first_line=0)
    add_runs(p, "初步调试版 · 可编辑模板", size=11, bold=True, color=ORANGE, east_asia=FONT_HEADING_3)

    rows = [
        ("参赛学校", "[待填]"),
        ("参赛队号", "[待填]"),
        ("队员姓名", "[待填]"),
        ("指导教师", "[待填]"),
        ("完成日期", "2026年7月 [待填] 日"),
    ]
    table = doc.add_table(rows=1, cols=2)
    header_cells = table.rows[0].cells
    for idx, text in enumerate(("参赛信息", "填写内容")):
        set_cell_shading(header_cells[idx], NAVY)
        p = header_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line=18, first_line=0)
        add_runs(p, text, size=10, bold=True, color=WHITE, east_asia=FONT_HEADING_3)
    repeat_header_row(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        for idx, text in enumerate((label, value)):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, line=20, first_line=0)
            add_runs(p, text, size=11, bold=(idx == 0), east_asia=FONT_HEADING_3 if idx == 0 else FONT_BODY)
    set_table_geometry(table, [2500, 6572], indent_dxa=TABLE_INDENT_DXA)
    set_table_borders(table, color="B7C3D0", size="5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=24, line=16, first_line=0)
    add_runs(p, "文档状态：总体方案与控制程序框架已建立，实测数据待联调补充", size=9.5, color=GRAY)


def add_instruction_page(doc, bullet_id):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=12, line=24, first_line=0)
    add_runs(p, "初稿使用说明", size=18, bold=True, color=NAVY, east_asia=FONT_HEADING_3)
    add_callout(
        doc,
        "编辑提示",
        "本页用于调试阶段内部协作，正式提交前应整页删除；正文中橙色的[待填]、[待实测]、[待确认]标记也应全部处理。",
    )
    items = [
        "补全封面信息，并在系统安装完成后加入整车、摆杆机构、摄像头视野和显示记录装置的实物照片。",
        "将“当前程序初值”替换为最终使用值，尤其是红外阈值、循迹PID、像素比例、连杆标定表、步进电机速度和角度限制。",
        "每项测试至少重复3次，记录原始数据、平均值、最大误差和失败现象；不要只写“满足要求”。",
        "最终摘要应写入实际完成指标；结论须与第6章测试表逐项对应。",
        "在Word中按Ctrl+A后按F9更新目录和页码，检查图表编号，再删除本页与所有编辑提示。",
    ]
    for item in items:
        add_list_item(doc, item, bullet_id)
    add_heading(doc, "当前方案基线", 2, numbered=False)
    rows = [
        ("底盘", "MSPM0G3507 + AT8236 + 两台MG513X + 万向轮"),
        ("循迹", "8路红外模块，黑白标定、归一化、加权误差和PID差速控制"),
        ("滚球机构", "42步进电机 + 连杆机构抬升摆杆活动端"),
        ("视觉控制", "MaixCAM Pro识别钢球并通过串口发送步进电机指令"),
        ("机械资料", "H42装配体：NEMA17/17HS4401、625ZZ、Axis、Top_body、Bottom、Sheet_metal等"),
    ]
    add_table(doc, ["子系统", "当前方案"], rows, [2200, 6872], caption="表0-1  当前方案基线")


def add_abstract(doc):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=14, line=24, first_line=0)
    add_runs(p, "摘  要", size=16, bold=True, east_asia=FONT_HEADING_3)
    abstract = (
        "本系统面向H题车载平衡滚球运动控制任务，采用两轮差速底盘与单万向轮支撑结构。"
        "底盘以MSPM0G3507为主控制器，经AT8236双H桥驱动两台MG513X电机，使用8路红外光电模块检测黑色环形路线；"
        "滚球装置采用42步进电机和连杆机构改变25 cm PPR凹槽摆杆的倾角。MaixCAM Pro完成钢球图像采集、"
        "灰度特征定位、速度估计及位置-速度级联控制，并向步进电机驱动器发送绝对位置指令。"
        "软件中设置了红外黑白归一化、循迹偏差加权、输出限幅、积分抗饱和、摆角变化率限制、丢线停车和丢球回水平等措施。"
        "目前系统处于初步联调阶段，已完成总体方案、机械装配与主要程序框架；静态±5 cm往返、AB段行驶、整圈稳定和任意目标位置的"
        "最终测试数据为[待实测]。"
    )
    add_body(doc, abstract, first_line=21)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=10, line=18, first_line=0)
    r = p.add_run("关键词：")
    set_run_font(r, east_asia=FONT_HEADING_3, size=10.5, bold=True)
    add_runs(p, "MSPM0G3507；红外循迹；MaixCAM Pro；级联PI；步进电机；平衡滚球")


def add_toc(doc):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=16, line=24, first_line=0)
    add_runs(p, "目  录", size=18, bold=True, color=NAVY, east_asia=FONT_HEADING_3)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_spacing(toc_p, line=18, first_line=0)
    add_field(toc_p, ' TOC \\o "1-3" \\h \\z \\u ', "右键单击此处并选择“更新域”生成目录")
    add_callout(doc, "目录提示", "首次打开文档后按Ctrl+A、F9更新全部域；正式提交前核对目录页码。", fill=LIGHT_GRAY, color=GRAY)


def chapter_1(doc, heading_id, assets):
    add_page_break(doc)
    add_heading(doc, "系统方案设计", 1, heading_id)
    add_heading(doc, "任务分析与指标分解", 2, heading_id)
    add_body(
        doc,
        "系统需在车体尺寸不超过35 cm×25 cm、仅使用红外光电模块循迹的条件下，完成环形路线行驶、起点停车、钢球静态往返定位以及移动状态下的中心点或任意位置稳定。"
        "赛题对报告重点考查方案论证、循迹与摆杆控制理论、电路和程序流程、测试数据完整性及图表规范性。因此设计上将整机拆分为底盘循迹、滚球平衡、图传记录和人机交互四个子系统。",
    )
    requirements = [
        ("图传", "实时显示并完整记录钢球运动视频", "画面覆盖完整摆杆，可回放"),
        ("底盘基础", "A点启动，顺时针一圈后停回A点", "总时间≤20 s；停车偏差≤2 cm"),
        ("静态滚球", "O→+5 cm→-5 cm并稳定", "总时间≤5 s；±5 cm最大误差≤1 cm"),
        ("AB段动态", "A点出发并通过B点，球稳定在O", "AB时间≤8 s；误差绝对值≤1 cm"),
        ("整圈中心", "顺时针一圈通过A点，球稳定在O", "总时间≤30 s；误差绝对值≤1 cm"),
        ("整圈任意点", "球从任意指定位置开始并保持", "总时间≤30 s；误差绝对值≤1 cm"),
        ("机械约束", "25 cm PPR凹槽，左端铰接", "h≥5 cm；球直径约1 cm；摆杆不超车身"),
    ]
    add_table(doc, ["项目", "任务", "关键指标"], requirements, [1450, 4720, 2902], caption="表1-1  赛题指标分解")

    add_heading(doc, "总体方案", 2, heading_id)
    add_body(
        doc,
        "底盘采用前置8路红外阵列检测黑线横向位置，MSPM0G3507根据归一化后的加权偏差生成左右轮速度目标，AT8236驱动两台MG513X电机实现差速转向，万向轮提供第三点支撑。"
        "滚球部分由固定在摆杆上方的MaixCAM Pro检测钢球纵向坐标，经位置外环和速度内环计算目标摆角，再由连杆标定关系换算为42步进电机绝对角度。两个控制分支并行运行，按键统一触发计时和任务状态机。",
    )
    add_picture(doc, assets["system"], 15.7, "图1-1  系统总体结构框图", "车载平衡滚球系统总体结构框图")

    add_heading(doc, "方案论证与选择", 2, heading_id)
    alternatives = [
        ("底盘结构", "四轮/履带", "双驱动轮+万向轮", "结构轻、差速控制简单，降低滚球机构的载荷与装配复杂度"),
        ("循迹传感", "摄像头识别赛道", "8路红外模块", "满足赛题仅允许红外光电循迹的约束，实时性高"),
        ("摆杆执行", "舵机直接驱动", "42步进电机+连杆", "角度分辨率和保持力更高，可通过连杆放大分辨率"),
        ("球位检测", "阈值连通域/圆检测", "灰度投影+连续跟踪", "摆杆ROI固定时计算量小，对60 fps控制更友好"),
        ("滚球控制", "单位置PID", "位置-速度级联PI", "内环抑制速度过冲，外环实现目标位置跟踪"),
    ]
    add_table(doc, ["对象", "备选方案", "选用方案", "选择理由"], alternatives, [1400, 1950, 2200, 3522], caption="表1-2  主要方案比较")
    add_callout(doc, "当前边界", "图传发送/接收方式、步进驱动器确切型号和最终供电电压尚未定稿，正文采用接口级描述并在硬件表中保留待确认项。", fill=LIGHT_GRAY, color=GRAY)


def chapter_2(doc, heading_id, assets):
    add_page_break(doc)
    add_heading(doc, "理论分析与控制方法", 1, heading_id)
    add_heading(doc, "差速底盘运动学", 2, heading_id)
    add_body(doc, "设左右轮线速度分别为vL、vR，驱动轮轮距为B，则车体中心线速度v和偏航角速度ω可近似表示为：")
    add_equation(doc, "v = (vR + vL) / 2，    ω = (vR - vL) / B")
    add_body(doc, "循迹控制器输出差速修正量Δv，以基础速度v0为前馈，形成vL*=v0-Δv、vR*=v0+Δv。若启用编码器闭环，两侧轮速控制器分别将速度误差转换为PWM占空比；转向正负号须在首次架车测试中确认。")

    add_heading(doc, "八路红外归一化与循迹误差", 2, heading_id)
    add_body(doc, "每一路传感器先分别采集白色背景值Wi和黑线值Bi，当前采样值为Ri。归一化结果Gi限定在0～1000：")
    add_equation(doc, "Gi = sat[ (Ri - Wi) × 1000 / (Bi - Wi) ]，  i = 0, 1, …, 7")
    add_body(doc, "当前程序采用从左到右[7，5，3，1，-1，-3，-5，-7]的权值。线位置误差按加权平均计算：")
    add_equation(doc, "e = 100 × Σ(wi Gi) / [7 × ΣGi]")
    add_body(doc, "控制器对e进行低通滤波后执行PID运算，并设置积分限幅、输出限幅、断线保持和盲转恢复。最终提交时应依据传感器实际安装方向确认权值符号，并记录白/黑标定跨度。")

    add_heading(doc, "摆杆-钢球简化模型", 2, heading_id)
    add_body(doc, "忽略空气阻力并假设钢球无滑动滚动，摆杆小角度倾斜θ时，钢球沿杆方向的加速度可近似为：")
    add_equation(doc, "x¨ ≈ (5/7) g sinθ ≈ (5/7) g θ")
    add_body(doc, "实际系统还包含凹槽曲率、滚动摩擦、车体加速度、机构间隙和摄像头延迟，因此上述模型只用于确定控制方向和初始量级。最终参数以静态阶跃和移动扰动实验为准。")

    add_heading(doc, "球位视觉测量与速度估计", 2, heading_id)
    add_body(doc, "MaixCAM Pro在摆杆内侧设置固定ROI，将RGB图像转为灰度并沿竖直方向求平均，得到一维灰度曲线。通过大尺度背景估计与局部平滑获得暗度峰值，在上一帧附近限制搜索窗口，再利用峰值邻域加权质心求钢球中心像素xpx。")
    add_equation(doc, "x = (xpx - x0) / kpx，    v = LS-slope{(tj, xj)}")
    add_body(doc, "式中x0为中心点O对应像素，kpx为每厘米像素数；速度v由约150 ms时间窗内的位置样本最小二乘斜率得到，再进行低通滤波。启动时将钢球放在O点可自动修正x0，kpx必须用已知刻度实测。")

    add_heading(doc, "位置-速度级联控制", 2, heading_id)
    add_body(doc, "外环根据目标位置xr与测量位置x的误差生成目标速度vr，内环根据目标速度与估计速度v的误差生成目标摆角θr：")
    add_equation(doc, "ex = xr - x，    vr = sat(Kpx ex + Kix ∫ex dt)")
    add_equation(doc, "ev = vr - v，    θr = sat(Kpv ev + Kiv ∫ev dt)")
    add_body(doc, "控制输出还经过摆角限幅和角度变化率限制，再按连杆标定曲线换算为电机角度。当前程序在检测持续丢失超过150 ms时清空积分并缓慢回到水平位置，避免视觉失效后机构继续偏转。")
    add_picture(doc, assets["control"], 15.7, "图2-1  滚球位置-速度级联控制框图", "滚球位置速度级联控制框图")

    add_heading(doc, "连杆非线性补偿", 2, heading_id)
    add_body(doc, "由于电机轴角与摆杆倾角并非严格线性，调试时在若干已知摆杆角度θj下记录电机绝对角度αj，并采用分段线性插值获得命令α(θ)。这种方法不依赖连杆理想尺寸，可同时补偿装配偏差和轻微非线性。")


def chapter_3(doc, heading_id):
    add_page_break(doc)
    add_heading(doc, "硬件电路设计", 1, heading_id)
    add_heading(doc, "硬件组成", 2, heading_id)
    rows = [
        ("底盘主控", "MSPM0G3507", "读取红外、按键与编码器，执行循迹/轮速控制和计时显示", "已确定"),
        ("电机驱动", "AT8236", "双路H桥，20 kHz PWM驱动两台直流减速电机", "已确定"),
        ("底盘电机", "MG513X ×2", "左右轮差速驱动；编码器参数[待确认]", "已确定"),
        ("支撑", "万向轮", "构成三点支撑并减小转向阻力", "已确定"),
        ("循迹", "8路红外光电模块", "采集黑线横向分布，只用于循迹", "已确定"),
        ("视觉主控", "MaixCAM Pro", "钢球检测、速度估计、级联控制及画面显示", "已确定"),
        ("执行机构", "42步进电机+驱动器", "通过连杆调节摆杆活动端高度", "驱动器型号[待确认]"),
        ("显示与按键", "≤2英寸显示屏+启动键", "显示计时与状态，满足赛题约束", "型号[待确认]"),
        ("电源", "车载电池+多路稳压", "动力、步进、主控/传感器分路供电并共地", "电压/容量[待实测]"),
    ]
    add_table(doc, ["模块", "器件", "作用", "状态"], rows, [1350, 2050, 3972, 1700], caption="表3-1  主要硬件组成")

    add_heading(doc, "底盘驱动与接口", 2, heading_id)
    add_body(doc, "AT8236的两路输入分别由MSPM0G3507定时器PWM通道控制。当前底盘程序使用20 kHz边沿对齐PWM；换向前先将同桥两路关断并留出约2 μs死区，以降低桥臂直通风险。两台MG513X分别连接左右驱动轮，万向轮布置在车体纵向中心线上。")
    add_body(doc, "红外模块通过多路选择线S0～S2和一路ADC采样8个通道；每次切换后等待约5 μs，丢弃首次ADC结果并对后续4次采样取平均。实际接线应以最终PCB和主控资源表为准。")

    interfaces = [
        ("8路红外", "S0/S1/S2 + ADC", "MSPM0G3507", "5 V或3.3 V兼容性[待确认]"),
        ("左/右AT8236", "各2路PWM", "MSPM0G3507", "20 kHz；方向翻转留死区"),
        ("编码器", "A/B相或单边计数", "MSPM0G3507", "每米脉冲数[待实测]"),
        ("MaixCAM Pro→步进驱动", "UART", "115200 bit/s初值", "当前程序/dev/ttyS0，地址0x01"),
        ("启动按键", "GPIO输入", "MSPM0G3507", "低电平/高电平有效[待确认]"),
        ("计时显示", "I²C/SPI", "MSPM0G3507", "屏幕尺寸不大于2英寸"),
    ]
    add_table(doc, ["接口对象", "信号", "连接端", "当前说明"], interfaces, [1600, 2000, 2200, 3272], caption="表3-2  建议接口分配表")

    add_heading(doc, "电源与抗干扰设计", 2, heading_id)
    add_body(doc, "动力电机、步进电机和数字控制部分应分路供电，在电池入口设置总开关、保险或自恢复保护。AT8236与步进驱动器附近配置大容量电解电容和高频去耦，MaixCAM Pro与MSPM0G3507使用独立稳压支路。所有控制地必须可靠共地，电机线与摄像头/ADC线分开走线，必要时在传感器电源端增加LC或RC滤波。")
    add_callout(doc, "上电检查", "先断开电机负载，逐路核对电压和极性；确认摆杆机械居中后再使能步进电机。驱动器额定电流、限流值和电池容量必须在联调完成后写入表3-1。")


def chapter_4(doc, heading_id, assets):
    add_page_break(doc)
    add_heading(doc, "机械结构设计", 1, heading_id)
    add_heading(doc, "总体布局", 2, heading_id)
    add_body(doc, "机械方案以“26年H题车载水平小球”文件夹中的H42装配体为基础。装配体包含42步进电机/NEMA17 17HS4401、625ZZ轴承、Axis转轴、Top_body、Bottom、Sheet_metal、PCB、连接盖和水管铰链等部件。左端通过铰链固定在距车体平板不小于5 cm的位置，右端由步进电机通过连杆抬升或下降。")
    add_picture(doc, assets["mechanical"], 15.7, "图4-1  H42摆杆与连杆机构示意", "H42摆杆、步进电机、连杆和摄像头的侧视结构示意")
    add_body(doc, "摆杆使用赛题规定的25 cm、4分PPR水管改造凹槽，内壁保持原状且光滑。摄像头固定在摆杆上方，视场覆盖完整刻度和两端极限位置；连杆运动范围内不得与摄像头、车架、线缆和钢球防滚落挡片干涉。")

    add_heading(doc, "结构参数与待确认尺寸", 2, heading_id)
    rows = [
        ("PPR摆杆长度", "25 cm", "赛题规定", "已确定"),
        ("摆杆外径/壁厚", "约2 cm / 0.34 cm", "赛题规定", "已确定"),
        ("铰接点高度h", "≥5 cm", "赛题规定", "[待实测]"),
        ("车体外廓", "≤35 cm×25 cm", "赛题规定", "[待实测]"),
        ("电机到摆杆连杆长度", "[待实测]", "H42装配体/实物", "影响角度映射"),
        ("摆杆最大安全倾角", "建议初调±2°", "当前程序限幅", "[待实测]"),
        ("摄像头高度与俯角", "[待实测]", "完整覆盖25 cm摆杆", "[待确认]"),
    ]
    add_table(doc, ["参数", "设计值", "依据", "状态"], rows, [2200, 1900, 2800, 2172], caption="表4-1  机械结构参数")

    add_heading(doc, "连杆标定方法", 2, heading_id)
    add_body(doc, "将摆杆调至机械水平并将该位置定义为电机零点。在活动端附近放置数字倾角仪，分别命令电机到若干正负角度，待机构稳定后记录实际摆杆角度。建议至少采集-2°、-1°、0°、+1°、+2°五点，并从正反两个方向重复，以评估回差。")
    rows = [
        ("-2.0", "[待实测]", "[待实测]", "[待实测]"),
        ("-1.0", "[待实测]", "[待实测]", "[待实测]"),
        ("0.0", "0.0", "0.0", "水平基准"),
        ("+1.0", "[待实测]", "[待实测]", "[待实测]"),
        ("+2.0", "[待实测]", "[待实测]", "[待实测]"),
    ]
    add_table(doc, ["摆杆角度θ/(°)", "正向电机角α/(°)", "反向电机角α/(°)", "回差/备注"], rows, [1900, 2400, 2400, 2372], caption="表4-2  连杆角度标定记录")


def chapter_5(doc, heading_id, assets):
    add_page_break(doc)
    add_heading(doc, "软件设计", 1, heading_id)
    add_heading(doc, "软件任务划分", 2, heading_id)
    add_body(doc, "系统软件分为底盘固件和MaixCAM Pro视觉控制程序。底盘固件负责ADC/PWM/编码器/按键/显示等实时任务；视觉程序负责图像采集、球位检测、速度估计、任务状态机和步进电机命令。两部分分别设置安全状态，避免任一控制分支失效时造成整机持续运动。")
    add_picture(doc, assets["flow"], 15.2, "图5-1  系统软件主流程", "底盘和滚球控制并行执行的软件流程图")

    add_heading(doc, "底盘循迹流程", 2, heading_id)
    add_body(doc, "启动前短按按键依次采集白色背景和黑线参考值。启动后以固定周期轮询8路ADC，完成每通道归一化和加权误差计算；误差低通滤波后进入PID，输出转换为左右轮差速量。若各通道归一化和低于可见阈值，则根据最近转向方向短时保持或盲转；超过恢复时限仍未重新识别黑线时立即停车并显示故障。")

    add_heading(doc, "滚球控制流程", 2, heading_id)
    add_body(doc, "当前MaixCAM Pro程序按640×480、60 fps采集图像，在固定ROI内识别钢球。控制周期初值为30 Hz，显示刷新为10 Hz。检测有效时先将像素坐标换算为厘米并滤波，再用150 ms滑动时间窗估计速度；级联控制器生成目标摆角，经分段标定映射后以115200 bit/s串口发送绝对位置命令。")
    add_body(doc, "任务3状态机依次执行“前往+5 cm→稳定确认→前往-5 cm→稳定确认→保持”。当前稳定判据初值为位置误差≤0.5 cm且速度绝对值≤1.0 cm/s，连续保持500 ms；这些阈值须根据最终评分判据和实测噪声调整。")

    add_heading(doc, "当前程序参数", 2, heading_id)
    params = [
        ("图像分辨率/帧率", "640×480 / 60 fps", "当前代码", "确认现场光照下稳定帧率"),
        ("ROI", "x=20～620，y=200～280", "当前代码", "[待复核]"),
        ("像素比例", "22.4 px/cm", "当前初值", "必须用刻度重新标定"),
        ("滚球控制频率", "30 Hz", "当前代码", "记录实际平均周期"),
        ("位置环Kp/Ki", "0.50 / 0.00", "保守初值", "[待整定]"),
        ("速度环Kp/Ki", "0.10 / 0.00", "保守初值", "[待整定]"),
        ("目标速度限幅", "±5.0 cm/s", "当前代码", "[待复核]"),
        ("摆杆角度限幅", "±2.0°", "当前代码", "先小角度安全调试"),
        ("步进电机速度/加速度", "100 rpm / 80", "当前代码", "单位和驱动器含义[待确认]"),
        ("电机绝对角限制", "±12°", "当前代码", "与机械极限联调"),
    ]
    add_table(doc, ["参数", "当前值", "属性", "后续处理"], params, [2000, 1900, 1700, 3472], caption="表5-1  视觉与滚球程序初始参数")

    add_heading(doc, "通信与安全处理", 2, heading_id)
    add_body(doc, "视觉程序上电后先查询步进电机连接状态，再执行使能与零点设置。每次输出前进行角度限幅和命令死区判断；程序退出时将摆杆返回水平并失能。正式系统还应加入机械限位或软件双重限位、串口应答超时、驱动故障码记录和急停处理。")
    add_callout(doc, "调试顺序", "先断开钢球并验证机构方向与极限，再放球整定速度环，最后整定位置环；移动底盘测试应在静态滚球测试稳定后进行。")


def chapter_6(doc, heading_id):
    add_page_break(doc)
    add_heading(doc, "测试方案与测试结果", 1, heading_id)
    add_heading(doc, "测试仪器与公共条件", 2, heading_id)
    add_body(doc, "测试场地按赛题尺寸制作：黑线宽1.8±0.2 cm，AB、CD直线段各1.5 m，BC、DA为半径0.5 m的半圆弧。所有动态测试记录电池电压、赛道光照、球位目标、整圈时间、最大误差和是否脱线；图传接收端同时保存完整视频。")
    instruments = [
        ("数字万用表", "[待填型号]", "电源电压、稳压输出"),
        ("数字倾角仪", "[待填型号]", "摆杆水平与角度标定"),
        ("钢尺/游标卡尺", "[待填型号]", "停车偏差、球位刻度、机构尺寸"),
        ("摄像头录像/上位机", "MaixCAM Pro + [待填接收端]", "球位轨迹、测试视频回放"),
        ("秒表/车载计时", "[待填]", "AB段和整圈时间核对"),
    ]
    add_table(doc, ["仪器", "型号", "用途"], instruments, [2200, 2800, 4072], caption="表6-1  测试仪器")

    add_heading(doc, "基础标定测试", 2, heading_id)
    add_heading(doc, "红外黑白标定", 3, heading_id)
    rows = []
    for ch in range(8):
        rows.append((f"G{ch}", "[待实测]", "[待实测]", "[待实测]", "通过/不通过[待填]"))
    add_table(doc, ["通道", "白色ADC", "黑线ADC", "跨度绝对值", "判定"], rows, [1050, 1700, 1700, 1900, 2722], caption="表6-2  八路红外标定数据")

    add_heading(doc, "底盘电机启动与一致性", 3, heading_id)
    motor_rows = [
        ("左轮正转", "[待实测]%", "[待实测]", "[待实测]"),
        ("左轮反转", "[待实测]%", "[待实测]", "[待实测]"),
        ("右轮正转", "[待实测]%", "[待实测]", "[待实测]"),
        ("右轮反转", "[待实测]%", "[待实测]", "[待实测]"),
    ]
    add_table(doc, ["方向", "最小稳定启动占空比", "空载速度/编码器值", "备注"], motor_rows, [1600, 2500, 2500, 2472], caption="表6-3  底盘电机启动阈值")

    add_heading(doc, "视觉与机构标定", 3, heading_id)
    calib_rows = [
        ("O点像素x0", "[待实测] px", "钢球放中心，采集20帧均值"),
        ("像素比例kpx", "[待实测] px/cm", "使用-10 cm与+10 cm刻度标定"),
        ("位置静态噪声", "[待实测] cm峰峰值", "钢球固定，记录10 s"),
        ("处理帧率", "[待实测] fps", "记录平均值与最低值"),
        ("控制周期", "[待实测] ms", "记录平均值与最大值"),
        ("机械回差", "[待实测]°", "同一摆角正反向到达比较"),
    ]
    add_table(doc, ["项目", "结果", "方法"], calib_rows, [2300, 2600, 4172], caption="表6-4  视觉与机械标定结果")

    add_heading(doc, "赛题功能测试", 2, heading_id)
    add_heading(doc, "静态O→+5 cm→-5 cm测试", 3, heading_id)
    static_rows = []
    for idx in range(1, 6):
        static_rows.append((str(idx), "[待实测]", "[待实测]", "[待实测]", "[待实测]", "[待填]"))
    add_table(doc, ["次数", "总时间/s", "+5 cm最大误差/cm", "-5 cm最大误差/cm", "稳态波动/cm", "判定"], static_rows, [800, 1250, 1900, 1900, 1600, 1622], caption="表6-5  静态滚球往返测试")

    add_heading(doc, "底盘循迹与停车测试", 3, heading_id)
    car_rows = []
    for idx in range(1, 6):
        car_rows.append((str(idx), "[待实测]", "[待实测]", "是/否[待填]", "[待填]"))
    add_table(doc, ["次数", "整圈时间/s", "停车偏差/cm", "是否全程在线", "备注"], car_rows, [850, 1600, 1800, 1900, 2922], caption="表6-6  底盘整圈循迹与停车测试")

    add_heading(doc, "移动滚球综合测试", 3, heading_id)
    combined_rows = [
        ("AB段中心保持", "≤8 s", "≤1 cm", "[待实测]", "[待实测]", "[待填]"),
        ("整圈中心保持", "≤30 s", "≤1 cm", "[待实测]", "[待实测]", "[待填]"),
        ("整圈任意目标-5 cm", "≤30 s", "≤1 cm", "[待实测]", "[待实测]", "[待填]"),
        ("整圈任意目标+5 cm", "≤30 s", "≤1 cm", "[待实测]", "[待实测]", "[待填]"),
        ("整圈任意目标[待填] cm", "≤30 s", "≤1 cm", "[待实测]", "[待实测]", "[待填]"),
    ]
    add_table(doc, ["测试项目", "时间要求", "误差要求", "实测时间/s", "最大误差/cm", "判定"], combined_rows, [2100, 1250, 1250, 1500, 1600, 1372], caption="表6-7  移动滚球综合测试")

    add_heading(doc, "测试结果分析模板", 2, heading_id)
    add_body(doc, "底盘方面，系统在[待填]光照和[待实测] V电池电压下完成[待填]/[待填]次整圈测试，平均用时为[待实测] s，最大停车偏差为[待实测] cm。主要误差来源为[待填：传感器高度/轮胎打滑/左右轮差异/转弯速度等]。")
    add_body(doc, "滚球方面，静态往返总时间为[待实测] s，+5 cm与-5 cm最大误差分别为[待实测] cm和[待实测] cm；移动整圈时最大误差为[待实测] cm。主要误差来源为[待填：视觉延迟/反光/机构回差/车体纵向加速度等]。")
    add_body(doc, "针对上述问题，采取[待填：降低转弯基础速度、分段调度PID、重新标定像素比例、减小连杆间隙、增加光照遮罩等]措施后，指标变化为[待实测]。应在最终报告中给出调整前后的对比数据或曲线。")


def chapter_7(doc, heading_id):
    add_page_break(doc)
    add_heading(doc, "结论与改进方向", 1, heading_id)
    add_body(doc, "本设计形成了由MSPM0G3507底盘控制器和MaixCAM Pro视觉控制器协同工作的车载平衡滚球系统。底盘采用8路红外检测和差速驱动，滚球装置采用42步进电机与连杆机构，视觉端使用球位/速度估计和位置-速度级联控制。现阶段已完成总体方案、H42机械装配体和主要软件框架，尚需通过完整联调补齐性能结论。")
    add_body(doc, "后续工作重点包括：完成红外与电机一致性标定；确定步进驱动器、电源和接口；测量连杆分段映射与回差；先整定滚球速度环再整定位置环；最后在AB段和整圈工况下联合调整底盘速度曲线与滚球控制参数。最终结论应逐项引用第6章数据，明确是否达到20 s/30 s、2 cm停车偏差和1 cm球位误差要求。")
    add_callout(doc, "最终提交前", "删除初稿使用说明和所有橙色占位符，补充实物图、关键电路图、测试曲线、图传记录说明及完整参考文献。", fill=LIGHT_ORANGE, color=ORANGE)

    add_heading(doc, "参考文献", 1, heading_id, numbered=False)
    refs = [
        "[1] 2026年全国大学生电子设计竞赛赛区赛（TI杯）H题：车载平衡滚球运动控制系统，2026。",
        "[2] Texas Instruments. MSPM0G3507 Datasheet / Technical Reference Manual，[待补版本与访问日期]。",
        "[3] AT8236双通道直流电机驱动器数据手册，[待补厂商与版本]。",
        "[4] Sipeed. MaixCAM Pro硬件与MaixPy开发文档，[待补版本与访问日期]。",
        "[5] 42步进电机及所用驱动器通信协议，[待确认型号后补充]。",
    ]
    for ref in refs:
        add_body(doc, ref, first_line=0, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)


def appendix(doc, heading_id):
    add_page_break(doc)
    add_heading(doc, "附录A  调试记录与提交检查表", 1, heading_id, numbered=False)
    add_heading(doc, "参数版本记录", 2, heading_id, numbered=False)
    rows = [
        ("[待填日期]", "底盘循迹", "[待填]", "[待填]", "[待填]"),
        ("[待填日期]", "滚球速度环", "[待填]", "[待填]", "[待填]"),
        ("[待填日期]", "滚球位置环", "[待填]", "[待填]", "[待填]"),
        ("[待填日期]", "移动综合", "[待填]", "[待填]", "[待填]"),
    ]
    add_table(doc, ["日期", "模块", "修改参数", "现象/数据", "结论"], rows, [1400, 1700, 2100, 2300, 1572], caption="表A-1  调试参数版本记录")

    add_heading(doc, "提交检查表", 2, heading_id, numbered=False)
    checks = [
        ("封面队伍信息、日期已填写", "□"),
        ("所有[待填]/[待实测]/[待确认]已处理", "□"),
        ("目录、页码、图表编号已更新", "□"),
        ("实物图、CAD图、电路图清晰且有题注", "□"),
        ("所有测试至少重复3次并保留视频", "□"),
        ("摘要与结论中的指标和测试表一致", "□"),
        ("图传接收和存储设备已准备封存", "□"),
        ("整车外廓、屏幕尺寸和唯一测试点符合要求", "□"),
    ]
    add_table(doc, ["检查项", "完成"], checks, [7572, 1500], caption="表A-2  最终提交检查表")


def build_document(assets):
    doc = Document()
    for section in doc.sections:
        set_section_layout(section)
        setup_header_footer(section)
    style_setup(doc)
    heading_id = create_numbering(doc, "headings")
    bullet_id = create_numbering(doc, "bullet")

    doc.core_properties.title = "2026年H题 车载平衡滚球运动控制系统设计报告"
    doc.core_properties.subject = "全国大学生电子设计竞赛H题初步调试版报告模板"
    doc.core_properties.author = "参赛队（待填）"
    doc.core_properties.keywords = "MSPM0G3507; MaixCAM Pro; 红外循迹; 平衡滚球; 级联PI"
    doc.core_properties.comments = "依据赛题、H42装配体和当前程序生成的可编辑初稿。"

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    cover_page(doc)
    add_instruction_page(doc, bullet_id)
    add_abstract(doc)
    add_toc(doc)
    chapter_1(doc, heading_id, assets)
    chapter_2(doc, heading_id, assets)
    chapter_3(doc, heading_id)
    chapter_4(doc, heading_id, assets)
    chapter_5(doc, heading_id, assets)
    chapter_6(doc, heading_id)
    chapter_7(doc, heading_id)
    appendix(doc, heading_id)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main():
    assets = make_assets()
    output = build_document(assets)
    print(output)


if __name__ == "__main__":
    main()
