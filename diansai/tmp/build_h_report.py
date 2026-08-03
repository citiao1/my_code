# -*- coding: utf-8 -*-
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\my_code\my_code\diansai")
OUT = ROOT / "H题模板" / "2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_排版校正版.docx"
ASSET_DIR = ROOT / "tmp" / "h_report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_LATIN = "Calibri"
COLOR_INK = "1F2933"
COLOR_BLUE = "245A73"
COLOR_LIGHT_BLUE = "EAF2F6"
COLOR_LIGHT_GRAY = "F3F5F7"
COLOR_BORDER = "AAB7C2"
COLOR_WARN = "FFF3CD"


def font_file(preferred="msyh.ttc"):
    windows_fonts = Path(r"C:\Windows\Fonts")
    for name in (preferred, "msyh.ttc", "simhei.ttf", "simsun.ttc"):
        candidate = windows_fonts / name
        if candidate.exists():
            return str(candidate)
    return None


def pil_font(size, bold=False):
    preferred = "msyhbd.ttc" if bold else "msyh.ttc"
    path = font_file(preferred)
    return ImageFont.truetype(path, size) if path else ImageFont.load_default()


def draw_arrow(draw, start, end, color="#476575", width=5):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1.0)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    tip = (x2, y2)
    left = (x2 - 18 * ux + 9 * px, y2 - 18 * uy + 9 * py)
    right = (x2 - 18 * ux - 9 * px, y2 - 18 * uy - 9 * py)
    draw.polygon([tip, left, right], fill=color)


def draw_box(draw, xy, title, detail, fill, outline="#476575"):
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=3)
    x0, y0, x1, y1 = xy
    title_font = pil_font(30, bold=True)
    detail_font = pil_font(22)
    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(
        ((x0 + x1 - (title_bbox[2] - title_bbox[0])) / 2, y0 + 20),
        title,
        font=title_font,
        fill="#173A4A",
    )
    lines = detail.split("\n")
    y = y0 + 72
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=detail_font)
        draw.text(
            ((x0 + x1 - (bbox[2] - bbox[0])) / 2, y),
            line,
            font=detail_font,
            fill="#354A54",
        )
        y += 32


def make_system_diagram(path):
    img = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(38, bold=True)
    draw.text((55, 35), "底盘与球杆执行机构的分工关系", font=title_font, fill="#173A4A")

    draw_box(draw, (70, 150, 360, 330), "8路红外阵列", "GPIO数字量\n黑线状态 b1...b8", "#EAF2F6")
    draw_box(draw, (490, 125, 820, 355), "MSPM0G3507", "循迹偏差PD\n偏航角速度PI\n左右轮速度PID", "#DDEFE8")
    draw_box(draw, (980, 150, 1290, 330), "差速底盘", "双电机 + 编码器\n完成循迹、计时与停车", "#F4ECE2")
    draw_arrow(draw, (360, 240), (490, 240))
    draw_arrow(draw, (820, 240), (980, 240))
    draw_arrow(draw, (980, 305), (820, 305), color="#7A5A00", width=4)

    draw_box(draw, (70, 520, 360, 700), "MaixCAM", "视觉队友输出球位 x\n位置误差与变化率", "#EAF2F6")
    draw_box(draw, (490, 495, 820, 725), "位置PD与安全约束", "生成目标球杆角度\n限幅、斜率限制、丢帧回零", "#DDEFE8")
    draw_box(draw, (980, 495, 1290, 725), "步进电机摇杆机构", "UART绝对位置命令\n旋转转化为杆端升降", "#F4ECE2")
    draw_arrow(draw, (360, 610), (490, 610))
    draw_arrow(draw, (820, 610), (980, 610))

    draw_box(draw, (1360, 300, 1540, 570), "供电", "底盘动力\n逻辑电源\nMaixCAM\n步进驱动\n共地", "#F3F5F7")
    draw_arrow(draw, (1360, 360), (1290, 250), color="#8A98A2", width=3)
    draw_arrow(draw, (1360, 500), (1290, 610), color="#8A98A2", width=3)
    img.save(path)


def make_mechanism_diagram(path):
    img = Image.new("RGB", (1600, 720), "white")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(38, bold=True)
    label_font = pil_font(26)
    small_font = pil_font(22)
    draw.text((55, 35), "单自由度球杆机械机构示意", font=title_font, fill="#173A4A")

    base_y = 590
    draw.line((90, base_y, 1510, base_y), fill="#59666D", width=8)
    draw.text((100, base_y + 18), "车体安装平面", font=label_font, fill="#354A54")

    pivot = (250, 475)
    rod_end = (1210, 390)
    draw.ellipse((pivot[0] - 20, pivot[1] - 20, pivot[0] + 20, pivot[1] + 20), fill="#245A73")
    draw.line((pivot[0], pivot[1], rod_end[0], rod_end[1]), fill="#243B4A", width=20)
    draw.text((155, 420), "固定铰点", font=label_font, fill="#173A4A")
    draw.text((600, 340), "带槽球杆（沿车体纵向布置）", font=label_font, fill="#173A4A")

    motor_center = (1130, 525)
    draw.ellipse((1065, 460, 1195, 590), fill="#EAF2F6", outline="#245A73", width=4)
    draw.text((1015, 610), "步进电机", font=label_font, fill="#173A4A")
    arm_end = (1270, 485)
    draw.line((motor_center[0], motor_center[1], arm_end[0], arm_end[1]), fill="#B26B35", width=13)
    draw.ellipse((arm_end[0] - 12, arm_end[1] - 12, arm_end[0] + 12, arm_end[1] + 12), fill="#B26B35")
    draw.line((arm_end[0], arm_end[1], rod_end[0], rod_end[1]), fill="#507A5C", width=12)
    draw.text((1265, 440), "摇臂", font=label_font, fill="#7A3F19")
    draw.text((1280, 375), "连杆", font=label_font, fill="#315E3A")

    draw_arrow(draw, (1335, 390), (1335, 300), color="#7A5A00", width=5)
    draw_arrow(draw, (1365, 300), (1365, 390), color="#7A5A00", width=5)
    draw.text((1390, 320), "杆端升降 h", font=label_font, fill="#7A5A00")
    draw.arc((170, 390, 430, 620), start=315, end=350, fill="#9B1C1C", width=5)
    draw.text((360, 510), "球杆角 θ", font=label_font, fill="#9B1C1C")

    draw.text((560, 640), "φ：电机绝对角；r：摇臂有效半径；L：铰点至驱动端距离", font=small_font, fill="#52636D")
    img.save(path)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_run_font(run, east_asia=FONT_BODY, latin=FONT_LATIN, size=10.5, bold=None, color=COLOR_INK):
    run.font.name = latin
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_style_font(style, east_asia, latin, size, color, bold=False):
    style.font.name = latin
    style._element.get_or_add_rPr()
    fonts = style._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    set_style_font(normal, FONT_BODY, FONT_LATIN, 10.5, COLOR_INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, FONT_HEADING, FONT_LATIN, 15, COLOR_BLUE, True)
    h1.paragraph_format.space_before = Pt(15)
    h1.paragraph_format.space_after = Pt(7)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, FONT_HEADING, FONT_LATIN, 12, COLOR_BLUE, True)
    h2.paragraph_format.space_before = Pt(11)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, FONT_HEADING, FONT_LATIN, 11, "1F4D5F", True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    set_style_font(caption, FONT_BODY, FONT_LATIN, 9, "52636D")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("H题 MaixCAM方案 · 底盘与机械机构阶段初稿")
    set_run_font(run, size=8.5, color="6B7780")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=8.5, color="6B7780")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fld_run = OxmlElement("w:r")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_run.append(fld_text)
    fld.append(fld_run)
    footer._p.append(fld)
    run = footer.add_run(" 页")
    set_run_font(run, size=8.5, color="6B7780")


def add_body(doc, text, bold_prefix=None, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, east_asia=FONT_BODY, latin="Cambria Math", size=11, color="173A4A")
    return p


def add_note(doc, label, text, warning=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [8880])
    cell = table.cell(0, 0)
    shade_cell(cell, COLOR_WARN if warning else COLOR_LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "：")
    set_run_font(r, bold=True, color="7A5A00" if warning else COLOR_BLUE)
    r = p.add_run(text)
    set_run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths_dxa, header_fill=COLOR_LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9, bold=True, color="173A4A")
    repeat_table_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=9)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.8))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def build_document():
    system_diagram = ASSET_DIR / "system_architecture.png"
    mechanism_diagram = ASSET_DIR / "rocker_mechanism.png"
    make_system_diagram(system_diagram)
    make_mechanism_diagram(mechanism_diagram)

    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("车载平衡滚球运动控制系统（H题）")
    set_run_font(r, east_asia=FONT_HEADING, size=22, bold=True, color="173A4A")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("MaixCAM方案：底盘控制与机械机构阶段初稿")
    set_run_font(r, east_asia=FONT_HEADING, size=15, bold=True, color=COLOR_BLUE)

    add_note(
        doc,
        "阶段范围",
        "本文只完成底盘循迹、球杆机械机构、位置PD到步进电机的控制链设计。MaixCAM内部的小球识别算法、图像标定结果和实测数据由视觉组后续补充。本文所有文字均按本队方案重新组织，未沿用他人的视觉与执行机构方案表述。",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("摘要：")
    set_run_font(r, bold=True)
    r = p.add_run(
        "系统采用底盘与滚球机构相对独立、统一供电和协同调试的结构。底盘以MSPM0G3507为控制器，"
        "利用8路数字红外阵列获得黑线横向偏差，经位置PD生成目标偏航角速度，再由陀螺仪偏航PI和左右轮编码器速度PID形成差速闭环，实现循迹、分段调速、计时和终点停车。滚球机构由MaixCAM直接取得小球位置，计算目标位置误差及其变化率，位置PD输出目标球杆角度；该角度经摇杆机构标定表换算为步进电机绝对角度，并通过UART发送位置指令。机械结构采用一端固定铰接、另一端由步进电机摇臂和连杆抬升的单自由度方案，使电机旋转转化为球杆端部上下位移。本文给出当前代码对应的控制链、接口与安全约束，并将尚未测量的结构尺寸、传动标定和控制参数明确列为后续实测项。"
    )
    set_run_font(r)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("关键词：")
    set_run_font(r, bold=True)
    r = p.add_run("MaixCAM；8路红外循迹；差速底盘；位置PD；步进电机；摇杆机构")
    set_run_font(r)

    doc.add_heading("1 系统方案与任务分工", level=1)
    add_body(
        doc,
        "整车由循迹底盘、球位视觉、单自由度球杆、步进电机执行机构及电源与人机交互模块组成。底盘控制器负责车辆沿黑线运行、轮速闭环、里程与计时；MaixCAM负责球位测量、滚球位置PD以及步进电机串口控制。这样可避免把高频底盘闭环依赖于视觉帧率，同时又保留MaixCAM直接驱动球杆机构的简洁链路。两部分共用启动状态和安全策略，但各自保持独立的实时任务。",
    )
    add_picture(doc, system_diagram, "图1  底盘与球杆执行机构的总体分工")

    add_table(
        doc,
        ["模块", "本阶段采用方案", "主要输入/输出"],
        [
            ("循迹底盘", "MSPM0G3507 + 8路数字红外 + IMU + 双编码器", "黑线状态、角速度、轮速 -> 左右电机PWM"),
            ("球位视觉", "MaixCAM直接测得小球沿杆方向的位置", "输出x、有效标志和采样时刻；识别细节待视觉组补充"),
            ("滚球控制", "位置PD直接生成目标球杆角度", "目标位置与球位 -> 球杆角度命令"),
            ("机械执行", "步进电机 + 摇臂/连杆 + 单端铰接球杆", "电机绝对角 -> 杆端高度与球杆倾角"),
        ],
        [1700, 3680, 3500],
    )

    doc.add_heading("2 底盘循迹控制设计", level=1)
    doc.add_heading("2.1 8路红外采集与横向偏差", level=2)
    add_body(
        doc,
        "8路红外传感器横向安装在车体前部，当前工程把OUT1至OUT8依次接至PA27、PA26、PA25、PA24、PB25、PB24、PB20和PA22。各通道以GPIO数字量读取，默认低电平表示检测到黑线。控制任务每5 ms读取一次原始电平，并将有效状态记为b_i（检测到黑线时b_i=1，否则为0）。",
    )
    add_body(
        doc,
        "为把离散通道变为连续偏差，按从左到右给各传感器分配位置权重w_i。当前可执行代码中的权重为{-3，-2，-1，-0.1，0.1，1，2，3}，比模板中常见的等间距权重更强调中间两路的小误差区。多路同时压线时取已触发通道权重的平均值，再以3.5归一化并限幅。",
    )
    add_formula(doc, "e_y = sat[-1,1]{ [Σ(w_i b_i) / Σb_i] / 3.5 }")
    add_body(
        doc,
        "当8路原始状态全同，程序保持上一帧识别结果，避免瞬时强光、脱线或全黑横线使偏差突变；若在其他组合下没有有效通道，偏差函数返回0。正式比赛前应在实际黑线宽度、安装高度和环境照度下重新核对有效电平、权重与脱线恢复方向，不能只依据注释确定参数。",
    )

    doc.add_heading("2.2 位置、偏航与轮速三级闭环", level=2)
    add_body(
        doc,
        "底盘控制按“黑线横向误差 -> 目标偏航角速度 -> 左右轮目标速度 -> 电机PWM”逐级计算。最外层位置PD把归一化偏差e_y变为目标偏航角速度ω_d。比例项决定回线力度，微分项抑制高速循迹时的左右摆动；微分量经一阶低通处理，降低数字红外跳变带来的尖峰。",
    )
    add_formula(doc, "ω_d(k) = sat{ K_p,e e_y(k) + K_d,e [e_y(k)-e_y(k-1)]/T_s }")
    add_body(
        doc,
        "偏航环使用IMU的Z轴角速度ω_z作为反馈。启动前完成零偏校准，运行时由PI控制器根据ω_d-ω_z产生差速量Δv，并进行积分与输出限幅。当前代码按下式合成左右轮目标速度，其中0.4为实车调试后的差速混合系数，并非由理想差速模型直接得到。",
    )
    add_formula(doc, "Δv = K_p,ω(ω_d-ω_z) + K_i,ω∫(ω_d-ω_z)dt")
    add_formula(doc, "v_L* = v_0 + 0.4Δv，    v_R* = v_0 - 0.4Δv")
    add_body(
        doc,
        "左右轮分别使用编码器速度闭环。编码器每5 ms更新一次，当前工程以3726计数/m换算轮速，再用新值权重0.35的一阶IIR滤波。增量式PID根据目标轮速与实测轮速计算PWM，单次P、I、D增量和最终占空比均有限幅，避免启动或转弯时输出突变。最终PWM通过左右电机驱动接口输出。",
    )
    add_formula(doc, "Δu(k)=K_p[e(k)-e(k-1)] + K_i e(k)T_s + K_d[e(k)-2e(k-1)+e(k-2)]/T_s")

    add_table(
        doc,
        ["控制层", "当前工程初值", "说明"],
        [
            ("循迹位置PD（K1）", "Kp=57.0，Kd=2.1，|ω_d|≤500", "快车模式现有初值，必须结合实际传感器权重复调"),
            ("偏航PI", "Kp=17.5，Ki=15.0，积分限幅80", "输入为目标/实测Z轴角速度"),
            ("左右轮速度PID", "Kp=1.0，Ki=80.0，Kd=0", "PWM限幅9000，单项增量限幅1000"),
            ("控制周期", "5 ms（200 Hz）", "红外、IMU、编码器和三层控制在同一控制帧内更新"),
        ],
        [2100, 2450, 4330],
    )
    add_note(doc, "参数性质", "上表是当前工程中的起始参数，不是最终比赛参数。更换车重、轮胎、红外高度或球杆机构后都必须重新记录阶跃与循迹波形。", warning=True)

    doc.add_heading("2.3 发车、调速与终点停车", level=2)
    add_body(
        doc,
        "现有工程设置K1和K2两种发车模式。K1用于快速循迹和终点停车：启动时清零里程、使能三层闭环和横线检测；K2用于低速联调：以较低初速起步并逐步提升到巡航速度。球杆安装完成后的首次联调应从低速模式开始，确认车体加速度不会使小球冲向端部，再逐步提高车速。",
    )
    add_body(
        doc,
        "终点识别使用最近20个控制帧（约100 ms）的滑动窗口。若窗口内8路中至少4路曾检测到低电平，程序认为传感器阵列扫过横向黑线，随即调用统一停车函数，清空PID历史、关闭控制使能、将目标速度置零并停止电机，同时锁存运行时间。该判据能适应小车斜着压过终点线，但必须增加起步屏蔽或里程门限，避免在起点附近和大曲率弯道误触发。",
    )
    add_body(
        doc,
        "K1当前在累计里程达到2580 mm后开始按步长降低基准速度。该阈值与场地一圈长度并不等价，只能视为当前样车的阶段参数。正式方案应根据完整场地长度、编码器实测比例和制动距离重新确定“允许识别终点线”的里程窗口，并通过连续多圈测试验证停车偏差。",
    )

    doc.add_heading("2.4 底盘程序时序", level=2)
    add_table(
        doc,
        ["执行顺序", "函数/数据", "作用"],
        [
            ("1", "MID_IR_ReadRaw / UpdateBits", "读取8路GPIO并生成黑线状态"),
            ("2", "MID_Line_CalcError", "由位置权重计算归一化横向偏差"),
            ("3", "MID_Chassis_PosStep", "位置PD生成目标偏航角速度"),
            ("4", "MID_Line_CheckEmergency", "滑动窗口判断终点横线并触发停车"),
            ("5", "MID_Chassis_YawStep", "IMU偏航PI生成左右轮差速"),
            ("6", "MID_Encoder_Update / SpeedStep", "编码器速度PID输出左右PWM"),
            ("7", "MID_Chassis_RampStep", "分频执行发车加速或里程减速"),
        ],
        [1050, 3250, 4580],
    )

    doc.add_heading("3 球杆机械机构设计", level=1)
    doc.add_heading("3.1 结构组成与布置", level=2)
    add_body(
        doc,
        "球杆沿小车纵向中心线布置，一端通过支座铰接在车体上，构成固定转动中心；另一端作为驱动端，由步进电机输出轴上的摇臂带动连杆上下运动。摇臂转动后改变连杆连接点的高度，从而抬高或降低球杆驱动端，使球杆绕固定铰点产生小角度转动。该结构只有一个受控自由度，传力路径短，便于在MaixCAM中建立“目标球杆角度—电机绝对角度”的标定表。",
    )
    add_picture(doc, mechanism_diagram, "图2  步进电机摇杆带动球杆单端升降的结构示意")
    add_body(
        doc,
        "固定铰点、驱动端连接点和步进电机轴线应安装在同一纵向工作平面内，减小连杆侧向扭曲。驱动端连接件需允许小范围摆动，避免机构运动时卡死；球杆在水平位置时应处于摇臂有效行程的中部，为正、负两个方向留出相近余量。摄像头支架与球杆支座需要有独立刚度，防止步进电机动作造成视场相对球杆移动。",
    )

    doc.add_heading("3.2 运动学关系与标定", level=2)
    add_body(
        doc,
        "设步进电机相对水平零点的转角为φ，摇臂有效半径为r，摇臂在水平零点处的安装相位为φ_0，固定铰点到驱动端的有效距离为L。忽略连杆摆角和安装偏置时，驱动端相对升降量与球杆倾角可近似写为：",
    )
    add_formula(doc, "h(φ) ≈ r[sin(φ+φ_0)-sinφ_0]，    θ(φ)=arctan[h(φ)/L]≈h(φ)/L")
    add_body(
        doc,
        "实际机构还包含连杆长度、连接孔偏置和装配间隙，θ与φ通常不是严格线性关系。因此控制程序不直接使用理论比例，而是在球杆上放置角度尺或电子水平仪，分别记录若干组“球杆实测角θ—电机绝对角φ”，再采用分段线性插值反求电机目标角。当前程序中的{-2°对应-6°，0°对应0°，+2°对应+6°}只是占位标定，必须由实物测量替换。",
    )

    doc.add_heading("3.3 强度、回差与安全限位", level=2)
    add_body(
        doc,
        "步进电机选型应同时校核静态保持转矩和动态加速转矩。设驱动端等效竖向载荷为F，摇臂半径为r，传动效率为η，安全系数为S，则电机所需转矩至少满足T_m ≥ SFr/η。F应包括球杆自重、钢球在最不利位置的重力分量、连杆惯性和车辆振动附加载荷。待杆长、质量、摇臂半径与电机型号确定后，再代入实测值完成报告中的数值计算。",
    )
    add_body(
        doc,
        "机械端应设置硬限位，软件端同时设置球杆角度限幅、电机绝对角限幅和角度变化率限制。当前位置程序将正常平衡角限制在±2°，电机机械保护范围限制在±40°，命令死区为0.10°，角度变化率为8°/s；这些是保守初值。调试时应先拆下钢球，以低速、小角度检查正负方向、零位和限位，再装球逐步增加控制幅度。",
    )

    doc.add_heading("4 MaixCAM位置PD与步进电机驱动", level=1)
    doc.add_heading("4.1 视觉模块输入边界", level=2)
    add_body(
        doc,
        "视觉算法由队友实现，本阶段只规定控制接口。MaixCAM每次得到有效检测后，应向控制层提供小球沿球杆方向的位置x（单位cm）、检测有效标志valid和采样时刻t_k。像素到厘米的标定、反光抑制、异常点剔除和丢帧判断由视觉模块负责；控制层不得把无效坐标或过期坐标当作当前球位。",
    )

    doc.add_heading("4.2 位置PD生成球杆角度", level=2)
    add_body(
        doc,
        "设目标球位为x_r，位置误差定义为e=x_r-x。位置PD直接根据当前误差与误差变化率生成目标球杆角度θ_d。由于目标位置在保持阶段不变，误差微分近似等于小球速度的相反数，因此D项为系统提供阻尼，可抑制小球越过目标点后的往复振荡。控制符号必须通过实物验证：若小球位于正方向，输出角度应使其产生朝负方向的加速度。",
    )
    add_formula(doc, "e(k)=x_r-x(k)")
    add_formula(doc, "ė_f(k)=LPF{[e(k)-e(k-1)]/T_k}")
    add_formula(doc, "θ_d(k)=sat[-θ_max,θ_max]{ s[K_p e(k)+K_d ė_f(k)] }")
    add_body(
        doc,
        "其中s为由机构正方向确定的符号（+1或-1），θ_max为允许球杆角度。为避免视觉噪声频繁驱动步进电机，控制输出还应经过小误差死区、角度斜率限制和命令死区；目标位置切换时宜使用斜坡或S形轨迹，避免直接阶跃。视觉短时丢帧可保持上一条命令，超过设定时间后应清除PD历史并缓慢回到水平角。",
    )

    doc.add_heading("4.3 串口绝对位置控制", level=2)
    add_body(
        doc,
        "MaixCAM通过UART直接连接步进电机驱动器，当前程序采用/dev/ttyS0、115200 bit/s、电机地址0x01和校验字节0x6B。上电后先读取当前位置确认通信，再使能电机。只有在机构人工调到真实水平位置时，才允许发送一次保存零点命令；正常上电应执行回零命令并轮询当前位置，确认连续两次落在零位容差内后才进入闭环。",
    )
    add_body(
        doc,
        "每个控制周期先把θ_d通过标定表换算成电机绝对角φ_d，再将角度换算为脉冲数，以绝对位置模式发送0xFD指令。当前换算采用3200脉冲/圈，电机速度初值100 r/min、加速度参数80。若通信超时、回零失败或当前位置越限，系统应立即停止发送闭环命令，执行停止/失能，并提示重新检查共地、串口和机械零点。",
    )
    add_table(
        doc,
        ["功能", "当前协议/参数", "安全要求"],
        [
            ("通信检测", "0x36读取单圈位置", "收到完整帧后才认为电机在线"),
            ("保存水平零点", "0x93，人工水平时仅执行一次", "正常运行禁止每次上电重写零点"),
            ("回到已存零点", "0x9A，最长等待10 s", "角度容差1.5°，连续2次满足才通过"),
            ("绝对位置运动", "0xFD，3200脉冲/圈", "先做球杆角度与电机角度双重限幅"),
            ("停止/失能", "0xFE / 0xF3", "异常退出和回零失败必须执行"),
        ],
        [1900, 3100, 3880],
    )

    doc.add_heading("4.4 MaixCAM控制流程", level=2)
    add_table(
        doc,
        ["步骤", "处理内容"],
        [
            ("初始化", "校验参数 -> 打开UART -> 读取电机位置 -> 使能 -> 回到已保存水平零点"),
            ("接收球位", "取得x、valid、t_k；无效或过期数据不进入PD"),
            ("PD计算", "计算e和滤波微分，得到θ_d并执行角度限幅与斜率限制"),
            ("机构映射", "根据实测标定表将θ_d插值为电机绝对角φ_d"),
            ("命令输出", "角度变化超过死区时发送绝对位置命令"),
            ("故障处理", "丢帧超时则回水平；串口/回零/越限故障则停止并失能"),
        ],
        [1650, 7230],
    )

    doc.add_heading("5 联调顺序与待测参数", level=1)
    add_body(
        doc,
        "联调应按机械零位、步进电机方向、球杆角度标定、静态位置PD、低速底盘循迹、动态滚球控制的顺序进行。每一步只改变一组参数，并保存球位、目标角、实际电机角、8路红外状态、左右轮目标/实测速度和PWM日志。上一环节未通过限位与稳定性检查前，不进入下一环节。",
    )
    add_table(
        doc,
        ["待确认项目", "需要记录的量", "用于修正"],
        [
            ("8路红外安装", "横向位置、离地高度、黑/白电平、环境光", "位置权重、有效极性与脱线策略"),
            ("底盘里程", "实跑1 m的左右编码器计数", "当前3726计数/m的标定值"),
            ("底盘速度", "直线/弯道目标和实测速度、制动距离", "K1/K2速度、减速里程与终点窗口"),
            ("球杆机构", "L、r、连杆长度、零位相位、最大安全行程", "运动学、转矩和软硬限位"),
            ("角度映射", "至少5组θ-φ实测点，正负方向分别测", "替换当前占位标定表"),
            ("位置PD", "Kp、Kd、控制周期、D项滤波、稳态误差", "正式控制参数与到位判据"),
            ("视觉接口", "帧率、延迟、位置重复性、最长丢帧", "PD周期、超时与滤波参数"),
        ],
        [2150, 3550, 3180],
    )

    doc.add_heading("6 代码与文字一致性核对", level=1)
    add_note(
        doc,
        "发现的待统一项",
        "这些差异不影响本阶段方案方向，但在提交报告和最终烧录前必须统一，否则报告参数会与实车程序不一致。",
        warning=True,
    )
    add_table(
        doc,
        ["位置", "当前差异", "建议"],
        [
            ("mid_line.h 与 mid_line.c", "注释写等间距-3.5至+3.5；实际数组为-3、-2、-1、-0.1、0.1、1、2、3", "以实测传感器位置确定一套权重，并同步代码和报告"),
            ("mid_line.c", "注释多处写30帧/150 ms；实际宏为20帧，即100 ms", "统一窗口说明并实测误触发率"),
            ("mid_key.h 与调度器", "头文件宏写50 ms；调度器实际每4个5 ms控制帧调用，即20 ms", "统一速度渐变周期，避免参数理解错误"),
            ("main.py 与本方案", "现程序为位置外环+速度内环的串级PI；本阶段确认方案为位置PD直接输出角度", "由视觉/控制组决定后改成同一算法，再据实更新报告"),
        ],
        [2550, 3650, 2680],
    )

    doc.add_heading("7 阶段结论", level=1)
    add_body(
        doc,
        "本阶段已经形成与本队实物方向一致的两条控制链：底盘由8路红外位置PD、IMU偏航PI和双轮编码器速度PID完成循迹；滚球机构由MaixCAM球位输入、位置PD、机构角度标定和UART步进电机绝对位置控制完成调节。机械结构采用固定铰点与摇臂连杆实现球杆单端升降。下一阶段的重点不是继续扩写视觉描述，而是测量机构尺寸和θ-φ映射、统一代码中的权重与调度参数，并在静态球杆上完成位置PD调参后再进行整车动态联调。",
    )

    doc.add_page_break()
    doc.add_heading("附录A  本阶段对应代码入口", level=1)
    add_table(
        doc,
        ["功能", "文件/符号"],
        [
            ("8路红外与横线检测", "Code/Middle/mid_line.c：MID_IR_ReadRaw、MID_Line_CalcError、MID_Line_CheckEmergency"),
            ("底盘三级闭环", "Code/Middle/mid_chassis.c：MID_Chassis_PosStep、YawStep、SpeedStep"),
            ("PID实现与初值", "Code/Middle/mid_pid.c：MID_PID_Pos_Calc、Yaw_Calc、Speed_Calc"),
            ("5 ms任务调度", "Code/App/app_scheduler.c：s_scheduler_step"),
            ("发车、里程与停车", "Code/Middle/mid_key.c、mid_encoder.c"),
            ("步进电机串口", "main.py：MotorController、pipe_angle_to_motor_angle"),
            ("待改位置PD", "main.py：当前CascadeController需按最终算法统一"),
        ],
        [2250, 6630],
    )
    add_body(
        doc,
        "表中路径均相对于八路传感器底盘工程根目录；MaixCAM步进电机接口位于工作区根目录main.py。附录只列核心入口，最终提交时应按实际烧录版本更新文件名、函数名和控制周期。",
        first_line=False,
    )

    core = doc.core_properties
    core.title = "H题 MaixCAM方案：底盘控制与机械机构阶段初稿"
    core.subject = "8路红外差速底盘、位置PD与步进电机摇杆机构"
    core.author = "参赛队"
    core.comments = "根据本队代码与机械方案独立整理；视觉算法待队友补充。"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
