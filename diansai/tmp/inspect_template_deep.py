from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


TEMPLATE = Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_选手填写通用模板.docx")
CURRENT = Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_封面目录分页版.docx")


def color_value(font):
    try:
        return font.color.rgb or font.color.theme_color
    except Exception:
        return None


def fmt_length(value):
    return None if value is None else round(value.pt, 2)


def print_style(doc, name):
    style = doc.styles[name]
    p = style.paragraph_format
    f = style.font
    print(
        f"STYLE {name!r}: font={f.name!r} size={fmt_length(f.size)} bold={f.bold} "
        f"italic={f.italic} color={color_value(f)} align={p.alignment} "
        f"before={fmt_length(p.space_before)} after={fmt_length(p.space_after)} "
        f"line={p.line_spacing!r} first={fmt_length(p.first_line_indent)} "
        f"left={fmt_length(p.left_indent)} keep_next={p.keep_with_next} page_before={p.page_break_before}"
    )


template = Document(TEMPLATE)
print("=== TEMPLATE STYLES ===")
for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption", "TOC 1", "TOC 2"):
    if style_name in [s.name for s in template.styles]:
        print_style(template, style_name)

print("\n=== TEMPLATE OUTLINE ===")
for index, paragraph in enumerate(template.paragraphs):
    style = paragraph.style.name
    text = paragraph.text.strip().replace("\n", " ")
    if style.startswith("Heading") or (index >= 30 and text and ("【请填写" in text or "【建议" in text)):
        print(f"{index:03d} [{style}] {text[:260]}")

print("\n=== REPRESENTATIVE TEMPLATE PARAGRAPHS ===")
for index in (31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48):
    if index >= len(template.paragraphs):
        continue
    p = template.paragraphs[index]
    print(f"{index:03d} [{p.style.name}] {p.text.strip()[:260]}")
    for run_index, run in enumerate(p.runs[:8]):
        print(
            f"    run{run_index}: {run.text[:80]!r} font={run.font.name!r} size={fmt_length(run.font.size)} "
            f"bold={run.bold} italic={run.italic} color={color_value(run.font)} sub={run.font.subscript} sup={run.font.superscript}"
        )

print("\n=== CURRENT INLINE MATH CANDIDATES ===")
current = Document(CURRENT)
needles = ("_", "ω", "θ", "φ", "Δ", "η", "Kp", "Ki", "Kd", "PWM", "0x")
for index, paragraph in enumerate(current.paragraphs):
    text = paragraph.text.strip()
    if text and any(needle in text for needle in needles) and not paragraph._p.findall(
        ".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"
    ):
        print(f"{index:03d} [{paragraph.style.name}] {text[:300]}")
for table_index, table in enumerate(current.tables):
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            text = cell.text.strip().replace("\n", " / ")
            if text and any(needle in text for needle in needles):
                print(f"T{table_index}R{row_index}C{cell_index}: {text[:240]}")

with ZipFile(TEMPLATE) as archive:
    root = etree.fromstring(archive.read("word/document.xml"))
ns = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}
print(f"\nTEMPLATE_OMATH={len(root.xpath('//m:oMath', namespaces=ns))}")
