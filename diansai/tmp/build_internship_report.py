from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\my_code\my_code\diansai")
OUT = ROOT / "认识实习报告_智能车与电赛备赛实践.docx"
SMART_CAR_IMAGE = Path(r"D:\xwechat_files_1\xwechat_files\wxid_dzglakq35c1w22_4bea\temp\RWTemp\2026-07\b7818ae8e42d027cbc353c78e929def8\f218f15accf2adfd27a9bdcbb49d5a2a.jpg")
ELECTRONIC_CAR_IMAGE = Path(r"D:\xwechat_files_1\xwechat_files\wxid_dzglakq35c1w22_4bea\temp\RWTemp\2026-07\b7818ae8e42d027cbc353c78e929def8\ddd649d9a3d2bbc18e512fc70680bdfa.jpg")


def set_run_font(run, size=12, bold=False, color=None, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(paragraph, before=0, after=0, line=22, first_indent=True):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    if first_indent:
        pf.first_line_indent = Cm(0.74)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, value in kwargs.items():
            element.set(qn("w:" + key), str(value))


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=10)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    pf.line_spacing = Pt(24)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 13, bold=True, color=(31, 78, 121))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(18)
    set_run_font(p.add_run(text), size=10)
    return p


def add_image(doc, image_path, width_cm, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    picture = p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)
    return p


def add_cover(doc):
    for _ in range(7):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(34)
    set_run_font(p.add_run("认识实习报告"), size=26, bold=True, color=(31, 78, 121))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(54)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(26)
    set_run_font(p.add_run("智能车与电子设计竞赛备赛实践"), size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(82)
    set_run_font(p.add_run("以竞赛工程实践深化专业认知"), size=13, color=(89, 89, 89))

    for text in ("学院：信息科学与技术学院", "专业班级：自控2401", "姓名：傅思雄", "学号：2024030448"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        set_run_font(p.add_run(text), size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    set_run_font(p.add_run("2026年7月"), size=12)
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = Pt(22)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("认识实习报告：智能车与电子设计竞赛备赛实践"), size=9, color=(89, 89, 89))

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("第 "), size=10)
    add_page_number(p)
    set_run_font(p.add_run(" 页"), size=10)


def build():
    if not SMART_CAR_IMAGE.exists() or not ELECTRONIC_CAR_IMAGE.exists():
        raise FileNotFoundError("提供的比赛照片未找到。")

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("摘  要"), size=15, bold=True, color=(31, 78, 121))
    add_body(doc, "本次认识实习以智能车竞赛和全国大学生电子设计竞赛备赛为主要实践场景。暑期，我先随队赴邯郸参加智能车省赛，在三天的现场训练、调试和竞赛中获得省赛预赛第七名、决赛第五名；其间还完成了硬件盲盒任务，并协助队伍开展比赛保障。随后，我持续在实验室围绕电赛小车开展硬件连接、控制程序调试和整车联调。通过把课堂中的模拟电子技术、数字电路、单片机、传感器与自动控制知识落实到真实设备上，我对工程项目的分工协作、问题定位和迭代验证形成了更加具体的认识。")
    p = doc.add_paragraph()
    set_paragraph_format(p, after=8)
    set_run_font(p.add_run("关键词："), bold=True)
    set_run_font(p.add_run("认识实习；智能车；电子设计竞赛；NE555；闭环控制"))

    add_heading(doc, "一、实习背景与目标")
    add_body(doc, "认识实习不仅是了解行业和专业岗位的过程，也应当与自身的专业学习建立联系。由于暑期集中参加智能车与电子设计竞赛备赛，我将竞赛训练作为本次认识实习的主要内容：一方面在真实比赛节奏中学习车辆调试和团队协作，另一方面在实验室中继续完成电赛小车的软硬件联调。实习目标是把零散的理论知识转化为可测试、可复现的工程能力，并在反复调试中理解电子信息类项目从模块到系统的实施方法。")
    add_body(doc, "整个过程以“发现问题、分析原因、修改方案、再次验证”为主线。相较于只完成单个实验，竞赛备赛更强调时间约束、现场环境和系统可靠性，也让我更直观地认识到硬件连接、程序逻辑、传感器数据与机械结构之间的相互影响。")

    add_heading(doc, "二、智能车省赛实践")
    add_heading(doc, "（一）邯郸赛场的三天训练与比赛", level=2)
    add_body(doc, "我随队在邯郸参加智能车省赛，现场共持续三天。备赛和比赛期间，队伍需要根据赛道状态检查车模、供电、传感器和电机响应，并在有限时间内完成多轮试跑和参数调整。我主要参与硬件检查、线路整理、上电测试和比赛协助，配合队伍确认车辆在起步、转向和连续运行时的状态。最终，队伍获得省赛预赛第七名、决赛第五名的成绩。")
    add_body(doc, "这段经历让我认识到，智能车不是若干模块的简单拼接。任何一个看似细小的接触不良、供电波动或参数变化，都可能在高速运行时被放大。因此，比赛现场的工作必须形成清晰的检查顺序，并通过多次试跑确认修改是否真正有效。")
    add_image(doc, SMART_CAR_IMAGE, 12.2, "智能车车模内部控制板和电机连接")
    add_caption(doc, "图1  智能车车模及控制板调试现场")

    add_heading(doc, "（二）硬件盲盒与比赛协助", level=2)
    add_body(doc, "比赛期间，我完成了硬件盲盒任务：使用NE555定时器搭建闪烁灯电路。该任务要求根据芯片工作原理完成电阻、电容和发光二极管等元件的连接，使NE555在无稳态工作方式下输出周期性高低电平，从而驱动LED闪烁。搭建过程中，我先按照原理图核对引脚功能和极性，再分段检查供电、输出端电平和LED亮灭状态，最后通过更换阻容参数观察闪烁频率的变化。")
    add_body(doc, "这项小任务强化了我对基础模拟电路的理解：在实际硬件上，元件引脚、接线质量和测量方法同样重要。除完成盲盒任务外，我还协助队伍进行设备搬运、调试准备和比赛过程保障。在协作中，我体会到竞赛并非只依赖某一个人的程序或电路，规范分工、及时沟通和相互复核同样决定了现场效率。")

    add_heading(doc, "三、电子设计竞赛备赛与小车联调")
    add_heading(doc, "（一）从单个模块到整车控制", level=2)
    add_body(doc, "智能车省赛结束后，我继续在实验室投入电子设计竞赛备赛。当前工作围绕小车控制平台展开，包括主控板、驱动与供电连接、编码器测速、IMU姿态信息、灰度传感器反馈和串口通信等模块。与单独验证某一模块相比，整车调试更需要关注接口定义、供电共地、数据刷新周期和执行机构响应是否一致。")
    add_body(doc, "在控制程序中，小车以固定周期更新编码器和IMU数据，并根据灰度传感器反馈完成循迹方向修正。控制部分采用分层闭环思路：航向角环把方向误差转化为目标角速度，角速度环将其转换为左右轮差速修正，左右轮速度环再根据编码器测得的速度独立输出电机控制量。这样的结构使方向、姿态与速度问题能够分层分析，也便于在调试时定位是传感器数据、控制参数还是电机执行环节出现偏差。")
    add_image(doc, ELECTRONIC_CAR_IMAGE, 15.0, "电子设计竞赛备赛中的小车硬件与控制板")
    add_caption(doc, "图2  电子设计竞赛备赛中的小车硬件联调")

    add_heading(doc, "（二）调试方法与工程认识", level=2)
    add_body(doc, "备赛过程中，我逐步形成了由静态到动态的调试顺序：先检查接线、接口和供电是否正确，再分别验证传感器数据与电机动作，最后进入整车闭环运行。对于车辆跑偏、速度不一致或响应不稳定等现象，不能只靠反复修改参数，而要先查看编码器速度、姿态角速度和灰度偏差等反馈量是否合理，再决定修改控制增益、限幅策略或硬件连接。")
    add_body(doc, "实验室中的多次试跑使我认识到，代码并不是孤立存在的。比如电机输出看似由一个控制量决定，实际还受电池电压、轮胎状态、线路接触和机械安装影响；传感器数据虽然能被程序读取，也需要结合实际运动状态判断其可信度。因此，备赛本质上是软硬件协同的系统工程，记录现象、保留可复现条件和分步验证比盲目尝试更有效。")

    add_heading(doc, "四、实习收获与反思")
    add_body(doc, "第一，我加深了对专业基础课的理解。NE555闪烁灯电路将模拟电子技术中的定时与充放电过程落实到实物连接；小车控制则把单片机、传感器、执行机构和自动控制知识联系起来。理论公式和程序结构只有经过上电、测量和试跑，才能真正理解其工程含义。")
    add_body(doc, "第二，我提高了工程调试能力。面对故障时，我开始主动区分硬件问题与软件问题，并按照供电、连接、单模块、数据反馈、整车闭环的顺序进行排查。这个过程也让我养成了记录参数、保存可用版本和每次只改变少量条件的习惯。")
    add_body(doc, "第三，我认识到团队协作的重要性。在智能车比赛现场，车辆调试、设备管理、赛道观察和比赛保障需要成员间快速配合；在电赛备赛中，硬件、控制和测试工作也必须及时同步。今后的学习中，我将继续夯实电路设计、嵌入式开发和控制算法基础，同时提高文档记录与沟通能力。")

    add_heading(doc, "五、结语")
    add_body(doc, "本次认识实习以竞赛备赛为载体，使我从“完成课程实验”进一步走向“解决真实工程问题”。邯郸智能车省赛的现场经历、NE555硬件盲盒的基础训练，以及电子设计竞赛小车的持续联调，共同构成了我对本专业实践方式的初步认识。后续我会继续在实验室中完善小车平台，在实践中检验所学知识，并为今后的课程设计、学科竞赛和工程学习打下更扎实的基础。")

    doc.core_properties.title = "认识实习报告：智能车与电子设计竞赛备赛实践"
    doc.core_properties.subject = "认识实习"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
