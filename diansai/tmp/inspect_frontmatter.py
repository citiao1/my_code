from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


FILES = [
    Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_选手填写通用模板.docx"),
    Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_车载平衡滚球运动控制系统_K230-YOLO版.docx"),
    Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_公式优化版.docx"),
]


def block_text(block):
    if block.tag == qn("w:p"):
        return "".join(block.itertext()).strip()
    if block.tag == qn("w:tbl"):
        cells = block.findall(".//w:tc", block.nsmap)
        return " | ".join("".join(cell.itertext()).strip() for cell in cells)
    return ""


for path in FILES:
    print(f"\n=== {path.name} ===")
    doc = Document(path)
    print(f"sections={len(doc.sections)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    for index, block in enumerate(doc.element.body.iterchildren()):
        if index >= 35:
            break
        text = block_text(block).replace("\n", " ")
        kind = block.tag.rsplit("}", 1)[-1]
        style = ""
        if kind == "p":
            p_style = block.find("./w:pPr/w:pStyle", block.nsmap)
            if p_style is not None:
                style = p_style.get(qn("w:val"), "")
        breaks = len(block.findall(".//w:br[@w:type='page']", block.nsmap))
        sect = len(block.findall(".//w:sectPr", block.nsmap))
        print(f"{index:02d} {kind:4s} style={style!r} pagebreak={breaks} sect={sect}: {text[:180]}")
