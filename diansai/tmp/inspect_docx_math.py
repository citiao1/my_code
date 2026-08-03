from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


DOCX = Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_排版校正版.docx")


def iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield "body", paragraph
    for table_index, table in enumerate(parent.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    yield f"table{table_index}/r{row_index}/c{cell_index}", paragraph


doc = Document(DOCX)
needles = ("=", "Σ", "∑", "β", "ω", "Δ", "sqrt", "PWM", "theta", "alpha", "Kp", "Kd")
for index, (where, paragraph) in enumerate(iter_paragraphs(doc)):
    text = paragraph.text.strip()
    if text and any(needle in text for needle in needles):
        print(f"{index:04d} [{where}] style={paragraph.style.name!r} align={paragraph.alignment!r}: {text}")

with ZipFile(DOCX) as archive:
    xml = etree.fromstring(archive.read("word/document.xml"))
ns = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}
math_nodes = xml.xpath("//m:oMath | //m:oMathPara", namespaces=ns)
print(f"OMML_NODES={len(math_nodes)}")
