# -*- coding: utf-8 -*-
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


PATH = Path(r"D:\my_code\my_code\diansai\H题模板\2026电赛报告H题_MaixCAM方案_底盘与机械机构阶段初稿_排版校正版.docx")


with ZipFile(PATH) as archive:
    bad_member = archive.testzip()
    names = set(archive.namelist())
    media = sorted(name for name in names if name.startswith("word/media/"))
    print(f"zip_test={'PASS' if bad_member is None else 'FAIL:' + bad_member}")
    print(f"media_count={len(media)} media={media}")

doc = Document(PATH)
all_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        all_text += "\n" + " | ".join(cell.text for cell in row.cells)

for forbidden in ("K230", "YOLO", "数字舵机", "选手填写通用模板", "【请填写", "____"):
    print(f"forbidden[{forbidden}]={all_text.count(forbidden)}")

print(
    f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} "
    f"sections={len(doc.sections)} inline_shapes={len(doc.inline_shapes)}"
)

heading_counts = {}
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        heading_counts[paragraph.style.name] = heading_counts.get(paragraph.style.name, 0) + 1
print(f"heading_counts={heading_counts}")

geometry_ok = True
for index, table in enumerate(doc.tables):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
    row_widths_ok = True
    for row in table.rows:
        widths = []
        for cell in row.cells:
            tc_w = cell._tc.tcPr.first_child_found_in("w:tcW")
            widths.append(int(tc_w.get(qn("w:w"))))
        if widths != grid_widths:
            row_widths_ok = False
    current_ok = (
        tbl_w is not None
        and int(tbl_w.get(qn("w:w"))) == sum(grid_widths)
        and tbl_ind is not None
        and int(tbl_ind.get(qn("w:w"))) == 120
        and row_widths_ok
    )
    geometry_ok = geometry_ok and current_ok
    print(
        f"table[{index}] rows={len(table.rows)} cols={len(table.columns)} "
        f"width={sum(grid_widths)} geometry={'PASS' if current_ok else 'FAIL'}"
    )

section = doc.sections[0]
print(
    "page_cm="
    f"{section.page_width.cm:.2f}x{section.page_height.cm:.2f} "
    f"margins={section.top_margin.cm:.2f}/"
    f"{section.right_margin.cm:.2f}/"
    f"{section.bottom_margin.cm:.2f}/"
    f"{section.left_margin.cm:.2f}"
)
print(f"table_geometry={'PASS' if geometry_ok else 'FAIL'}")
print(f"text_chars={len(all_text)}")
