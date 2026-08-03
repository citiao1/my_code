from __future__ import annotations

import copy
import hashlib
import math
from pathlib import Path
from zipfile import ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from lxml import etree


ROOT = Path(r"D:\my_code\my_code\diansai")
WORK = ROOT / ".codex_h_report"
TEMPLATE = WORK / "template.docx"
OUTPUT = WORK / "final.docx"
FIG1 = ROOT / "tmp" / "report_builder" / "assets" / "system_overview.png"
FIG2 = WORK / "control_interfaces.png"
FIG3 = ROOT / "tmp" / "report_builder" / "assets" / "software_flow.png"


def c14n(element) -> bytes:
    return etree.tostring(element, method="c14n")


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def archive_media_hashes(path: Path) -> dict[str, str]:
    with ZipFile(path) as package:
        return {
            name: sha256(package.read(name))
            for name in package.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        }


def section_signature(doc: Document) -> list[tuple]:
    return [
        (
            section.start_type,
            section.page_width,
            section.page_height,
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
            section.header_distance,
            section.footer_distance,
            section.different_first_page_header_footer,
            section.header.is_linked_to_previous,
            section.footer.is_linked_to_previous,
        )
        for section in doc.sections
    ]


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_text_run(paragraph, text: str, bold: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def set_paragraph(paragraph, text: str, label: str | None = None) -> None:
    clear_paragraph(paragraph)
    if label and text.startswith(label):
        add_text_run(paragraph, label, bold=True)
        add_text_run(paragraph, text[len(label) :])
    else:
        add_text_run(paragraph, text)


def set_cell(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    add_text_run(paragraph, text)


def set_table_matrix(table, rows: list[list[str]]) -> None:
    if len(rows) != len(table.rows):
        raise ValueError(f"row count mismatch: {len(rows)} != {len(table.rows)}")
    for r_index, values in enumerate(rows):
        if len(values) != len(table.columns):
            raise ValueError(
                f"column count mismatch at row {r_index}: "
                f"{len(values)} != {len(table.columns)}"
            )
        for c_index, value in enumerate(values):
            set_cell(table.cell(r_index, c_index), value)


def set_image_cell(cell, path: Path, width_inches: float, alt_text: str) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # The template uses exact 18 pt line spacing in these image cells. Word
    # expands inline pictures anyway, while LibreOffice clips them to one line.
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    inline = paragraph._p.xpath(".//wp:inline")[-1]
    doc_pr = inline.find(qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("descr", alt_text)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "msyhbd.ttc" if bold else "msyh.ttc"
    return ImageFont.truetype(str(Path(r"C:\Windows\Fonts") / name), size)


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, width=5, color="#4A4A4A"):
    draw.line((start, end), fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max(math.hypot(dx, dy), 1.0)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    left = (x2 - 18 * ux + 9 * px, y2 - 18 * uy + 9 * py)
    right = (x2 - 18 * ux - 9 * py, y2 - 18 * uy - 9 * px)
    # Recompute the second wing explicitly; this avoids depending on arrow slope.
    right = (x2 - 18 * ux - 9 * px, y2 - 18 * uy - 9 * py)
    draw.polygon((end, left, right), fill=color)


def draw_box(draw, xy, title: str, lines: list[str], fill: str) -> None:
    draw.rounded_rectangle(xy, radius=8, fill=fill, outline="#555555", width=3)
    x0, y0, x1, _ = xy
    title_font = font(29, True)
    detail_font = font(21)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(
        ((x0 + x1 - title_box[2]) / 2, y0 + 16),
        title,
        font=title_font,
        fill="#000000",
    )
    y = y0 + 64
    for line in lines:
        box = draw.textbbox((0, 0), line, font=detail_font)
        draw.text(
            ((x0 + x1 - box[2]) / 2, y),
            line,
            font=detail_font,
            fill="#000000",
        )
        y += 31


def make_interface_figure() -> None:
    image = Image.new("RGB", (1800, 920), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 30), "控制系统接口与信号流", font=font(40, True), fill="#000000")

    draw_box(
        draw,
        (60, 135, 390, 360),
        "8路红外模块",
        ["S0/S1/S2多路选择", "ADC每20 ms轮询", "背景/目标线归一化"],
        "#F2F2F2",
    )
    draw_box(
        draw,
        (510, 110, 900, 385),
        "MSPM0G3507",
        ["加权质心循迹误差", "方向环与偏航角速度环", "左右轮速度环/安全状态机"],
        "#E2F0D9",
    )
    draw_box(
        draw,
        (1025, 135, 1385, 360),
        "AT8236与MG513X",
        ["20 kHz慢衰减PWM", "两台直流减速电机", "编码器/IMU反馈接口"],
        "#FCE4D6",
    )
    draw_arrow(draw, (390, 245), (510, 245))
    draw_arrow(draw, (900, 245), (1025, 245))
    draw_arrow(draw, (1025, 325), (900, 325), width=4)

    draw_box(
        draw,
        (60, 570, 390, 805),
        "MaixCAM Pro",
        ["640x480固定ROI", "灰度剖面检测钢球", "位置滤波与速度估计"],
        "#DDEBF7",
    )
    draw_box(
        draw,
        (510, 545, 900, 830),
        "滚球级联控制",
        ["位置PI生成目标速度", "速度PI生成目标摆角", "摆角限幅/斜率/机构映射"],
        "#E2F0D9",
    )
    draw_box(
        draw,
        (1025, 570, 1385, 805),
        "42步进电机机构",
        ["UART 115200 bit/s", "绝对位置命令", "摇臂连杆抬升摆杆端"],
        "#FCE4D6",
    )
    draw_arrow(draw, (390, 685), (510, 685))
    draw_arrow(draw, (900, 685), (1025, 685))

    draw_box(
        draw,
        (1490, 285, 1740, 660),
        "公共约束",
        ["车载电池分路供电", "逻辑共地", "启动/急停", "计时显示", "图传与录像待接入"],
        "#F2F2F2",
    )
    draw_arrow(draw, (1490, 365), (1385, 275), width=3)
    draw_arrow(draw, (1490, 585), (1385, 690), width=3)
    image.save(FIG2)


PARAGRAPHS = {
    5: (
        "摘要：本阶段面向H题车载平衡滚球系统，按底盘循迹与滚球调节相对独立的思路完成总体方案。底盘以MSPM0G3507为主控，八路红外模块经逐路标定和归一化形成横向偏差，方向环、角速度环和轮速环驱动AT8236及两台MG513X电机。滚球端由MaixCAM Pro采集640x480图像，在固定ROI内通过灰度剖面与加权质心获得钢球位置，采用150 ms窗口估计速度，并以30 Hz位置-速度级联控制生成摆杆角度；角度经连杆标定换算为42型步进电机绝对位置指令。机械部分采用一端铰接、另一端由摇臂连杆抬升的H42装配方案。当前系统仍处于初级调试阶段，滚球程序工作在中心保持模式，循迹终点识别、图传装置及各项性能数据尚待实测，因此本文只给出可继续填写的设计框架，不虚构测试结论。",
        "摘要：",
    ),
    6: (
        "关键词：八路红外循迹；MaixCAM Pro；灰度剖面检测；位置-速度级联；42型步进电机",
        "关键词：",
    ),
    7: (
        "阶段说明：本稿严格沿用选手填写通用模板的章节、表格、图号、分页和封面。正文依据现有MSPM0G3507底盘程序、MaixCAM Pro程序及H42机械装配资料填写；未取得的尺寸、器件规格和试验结果统一标为“待确认”或“待实测”。",
        "阶段说明：",
    ),
    8: (
        "提交前检查：完成实车调试后，应补齐图传、终点停车、供电参数和表B-1至表B-6原始数据，再据此改写摘要、结论和表7。封面按本次要求保持母版原样；最终提交前仅在得到队伍明确确认后再填写日期并更新目录、页码和交叉引用。",
        "提交前检查：",
    ),
    39: (
        "总体思路：系统分为底盘循迹链和滚球闭环链。底盘侧由八路红外模块采集轨迹反射强度，MSPM0G3507完成归一化、横向偏差计算、丢线状态判断及方向/角速度/轮速控制，经AT8236驱动两台MG513X电机；一个万向轮构成三点支撑。滚球侧由MaixCAM Pro直接完成钢球检测、位置滤波、速度估计和位置-速度级联控制，再通过UART向42型步进电机发送绝对位置命令，连杆机构把电机角度转换为摆杆端部升降。两条控制链统一由车载电池供电，通过启动、急停、计时显示和故障记录协调。",
        "总体思路：",
    ),
    40: (
        "机械布局：依据“26年H题车载水平小球”中的H42装配体，25 cm PPR凹槽摆杆沿车体纵向布置，左端通过Axis与625ZZ轴承形成固定铰点，右端由42型步进电机的摇臂和连杆抬升。Top_body、Bottom、Sheet_metal等零件构成支承与安装框架。铰点安装高度按题目保持h不小于5 cm，摄像头固定在摆杆上方并覆盖全槽；步进电机、支架和摄像头尽量靠近车体中心线，最终外廓、重心、连杆孔距和摆角范围以装配实测为准，并保证摆杆不超出35 cm x 25 cm车身投影。",
        "机械布局：",
    ),
    45: (
        "备选方案一：采用少量开关式红外探头，根据最左、居中、最右状态分段转向。其接口简单、程序量小，但黑线处于相邻探头之间时误差不连续，环境光和安装高度变化容易造成阈值漂移，半径0.5 m连续弯道上的速度上限和重复性受限。",
        "备选方案一：",
    ),
    46: (
        "备选方案二：采用八路模拟红外阵列，通过S0/S1/S2选择通道，每路切换后丢弃首个ADC样本并平均后续4次；分别采集背景和目标线参考，将每路映射到0至1000，再用位置权重7、5、3、1、-1、-3、-5、-7计算连续偏差。该方案需要两次标定和更多参数整定，但可获得连续误差，并能依据有效通道数和灰度总量判断窄缝、盲转和丢线。",
        "备选方案二：",
    ),
    47: (
        "方案选择：采用八路归一化方案。可验证依据一是20 ms更新周期内可得到-100%至100%的连续偏差，便于方向环输出连续目标角速度；依据二是程序已经实现150 ms窄缝保持、带350 ms转向记忆的盲转恢复及1200 ms超时停车，适合后续分别记录弯道误差和丢线恢复时间。",
        "方案选择：",
    ),
    49: (
        "成像与识别：滚球端使用MaixCAM Pro，当前相机设置为640x480、60 fps，检测ROI为x=20至620、y=200至280。图像转灰度并做5x5高斯滤波，对ROI按列求均值形成一维灰度剖面，再以大尺度高斯基线减去原剖面得到钢球暗度响应；在上一帧附近或全ROI内搜索峰值，并用局部加权质心求横坐标。对比度门限为8，允许相邻帧最大跳变4 cm，曝光方式和实际帧率待实机固定后记录。",
        "成像与识别：",
    ),
    50: (
        "方案比较：学习型目标检测可通过训练覆盖反光、模糊和背景变化，但需要采集标注数据、完成模型转换与板端推理验证，且当前队伍程序并未使用该类模型。固定ROI灰度剖面法不需要训练，运算链短，便于在初级阶段直接观察阈值、峰值和位置误差；其局限是对槽内阴影、钢球高光和相机姿态变化更敏感。后续若跨光照误检率较高，可将学习型检测作为备选，而不能在本稿中把它写成现用方案。",
        "方案比较：",
    ),
    51: (
        "选择结论：现阶段采用“固定ROI-灰度滤波-列均值-背景基线-暗度峰值-局部质心-时序约束”的摄像头检测流程。球位来自MaixCAM Pro图像而非接触式传感器，满足题目对摄像头检测的约束；最终定位误差、丢帧率和环境光适应性在表B-1中实测。",
        "选择结论：",
    ),
    53: (
        "执行机构比较：数字舵机集成位置闭环、接口简洁，但齿轮回差和可选摇臂几何会直接影响小角度重复性；42型步进电机便于按脉冲或串口绝对位置给定，配合摇臂连杆可以获得较细的杆端位移，但需要验证失步、谐振、驱动电流、回零与机构回差。现有H42装配已包含NEMA17/42型电机、轴承、转轴、支承件和连杆关系，故采用42型步进电机方案，额定转矩、有效传动比和回差留待实测。",
        "执行机构比较：",
    ),
    54: (
        "控制结构比较：单位置环直接输出摆角，视觉噪声和延迟会使输出频繁反向，且缺少速度阻尼；位置-速度级联把位置误差先变成有界目标速度，再由速度误差决定摆角，可分别限制球速和杆角，并在丢球时平滑回水平。当前程序保留PI结构，但初调阶段两个积分增益均为0，实际先以P-P级联验证方向和稳定性，待无静差需求明确后再逐步加入积分。",
        "控制结构比较：",
    ),
    55: (
        "最终方案：对象为凹槽内滚动钢球；MaixCAM Pro提供球位图像反馈；低通滤波和最小二乘窗口产生位置、速度状态；位置PI外环输出目标球速，速度PI内环输出目标摆角；摆角经限幅、斜率限制和分段线性连杆标定换算为步进电机绝对角度；步进电机与连杆改变摆杆坡度，钢球运动后再次由摄像头闭环测量。",
        "最终方案：",
    ),
    59: (
        "偏差定义：八路传感器从车左至车右记为G0至G7，归一化值s_i属于[0,1000]，目标线越明显数值越大，权重w_i依次为7、5、3、1、-1、-3、-5、-7。横向误差取e_y=100*sum(w_i*s_i)/(7*sum(s_i))，并限幅到[-100%,100%]。当sum(s_i)<200时进入窄缝保持或盲转恢复；尚未完成黑白标定或恢复超时则停止电机。",
        "偏差定义：",
    ),
    60: "e_y = 100 * sum(w_i s_i) / [7 * sum(s_i)],  w_i={7,5,3,1,-1,-3,-5,-7}",
    61: (
        "循迹控制律：偏差先以新样本权重0.35低通滤波，再按位置式PID生成方向输出，并映射为目标偏航角速度，正常限幅为正负60 deg/s；偏航角速度环输出左右轮差速修正量，轮速增量式PID再输出电机百分比。当前默认方向参数为Kp=200、Ki=0、Kd=350，底盘主循环10 ms、灰度更新20 ms，AT8236采用20 kHz慢衰减PWM。以上均为代码初值，最终值以实车调参记录为准。",
        "循迹控制律：",
    ),
    62: "v_L* = v - Delta_v,  v_R* = v + Delta_v;  Delta_v = PI(omega_r - omega)",
    63: (
        "终点识别与制动：当前底盘程序已实现按键长按启动、运行中按键立即停车以及丢线超时停机，但尚未完成“一圈后识别A点启停线、计时锁存和停车偏差修正”。后续应在起步阶段屏蔽启停线，结合编码器累计里程和多通道同时见线判据，在接近6.142 m时预减速，过A线后制动并低速修正；对应阈值和制动距离必须由表B-2试验确定。",
        "终点识别与制动：",
    ),
    64: (
        "时间预算：现有本地循迹调试速度为0.20 m/s，按6.142 m计算理想一圈约30.71 s；程序当前循迹上限0.30 m/s时理想值仍约20.47 s，尚未计入起停和弯道，因此不能满足20 s任务。要留出制动裕量，后续整车稳定后需把平均速度提高到大于0.307 m/s，并将直线目标暂定为0.35至0.40 m/s、弯道速度和程序上限据实重定，完成后填入表B-2。",
        "时间预算：",
    ),
    66: (
        "建模约定：以摆杆中心O为x=0，向执行机构一端为x正方向；杆端升高、使正x方向势能增加时摆角theta定义为正。钢球近似为实心球并在光滑凹槽内纯滚动，摆角按小角度处理。车体沿杆方向的加速度记为a_x，摩擦、槽截面约束、连杆回差、步进失步和车体俯仰作为未建模扰动。",
        "建模约定：",
    ),
    70: (
        "参数与推导：题目给出钢球直径约1 cm、摆杆长度25 cm，钢球质量m、实际半径r、铰点至驱动端有效距离和连杆孔距待称量或从最终装配图读取。控制程序目前将正常摆角限制为正负2 deg，电机机械角保护为正负40 deg。依据平动方程m*x_ddot、转动方程I*alpha以及无滑动约束x_dot=r*alpha，可得到上式；摩擦、摆角正弦非线性和连杆非线性将通过像素-位置与摆角-电机角标定补偿。",
        "参数与推导：",
    ),
    71: (
        "控制含义：小角度模型近似双积分对象，只反馈位置时阻尼不足，故需要速度估计。小车起步时a_x产生与重力倾角同量纲的惯性扰动，弯道时横向加速度通过机构不对中和槽壁传入，制动时a_x符号突变最容易导致超调。当前阶段先依靠速度环、摆角限幅和缓变约束抑制，车体加速度前馈待动态日志具备后再加入。",
        "控制含义：",
    ),
    73: (
        "坐标标定：当前线性映射采用中心像素u_O=320.0、比例22.4 pixel/cm，即x=(u-u_O)/22.4 cm。启动时在O点连续采集20帧；若横坐标极差不超过0.5 cm，则以均值更新u_O，15 s超时后使用配置值。22.4 pixel/cm仍需利用凹槽刻度的两个以上已知位置重新测量，透视和畸变是否需要校正由端点残差决定。",
        "坐标标定：",
    ),
    74: "x_k = (u_k - u_O) / 22.4  cm  (当前线性初值，提交前用实物标定替换)",
    75: (
        "滤波与估计：位置使用alpha=0.45的一阶低通；速度由最近150 ms带时间戳的位置样本作最小二乘直线拟合，再以alpha=0.35低通。至少3个样本后速度才有效，速度绝对值限100 cm/s。跟踪阶段搜索范围由上一帧位置和4 cm最大跳变限制；250 ms未重获目标则恢复全ROI搜索。若有效测量中断达到150 ms，控制器清积分并使摆杆以斜率限制返回水平。",
        "滤波与估计：",
    ),
    76: (
        "分辨率核算：600 pixel有效横向ROI按22.4 pixel/cm换算约覆盖26.79 cm，理论像素分辨率约0.0446 cm/pixel，数值上小于1 cm指标。但最终误差还包含灰度质心抖动、相机视角、钢球反光、刻度标定和控制动态误差，必须在中心、正负5 cm及靠近端部位置重复测量后才能判断是否达标。",
        "分辨率核算：",
    ),
    78: (
        "误差与控制律：目标位置为x_r，测量位置为x，位置误差e_x=x_r-x；目标速度v_r由位置PI产生并限幅，速度误差e_v=v_r-v；目标摆角theta_r由速度PI产生，再乘控制方向符号。积分按实际dt离散累加并分别限幅，输出进入饱和且误差仍推动更深饱和时冻结对应积分。",
        "误差与控制律：",
    ),
    79: "v_r = sat(K_px e_x + K_ix sum(e_x dt), +/-5 cm/s)\ntheta_r = sat[K_pv(v_r-v) + K_iv sum((v_r-v)dt), +/-2 deg]",
    80: (
        "参数整定：当前初值K_px=0.50 1/s、K_ix=0，K_pv=0.10 deg/(cm/s)、K_iv=0；控制周期约33 ms，摆角限幅正负2 deg，摆角变化率8 deg/s，电机命令死区0.10 deg。整定顺序为先固定x_r=0验证方向和速度环，再增大速度环增益至具有足够阻尼，随后调位置环，最后再视静差逐步加入积分。每次改动只调整一组参数并保留日志。",
        "参数整定：",
    ),
    81: (
        "任务逻辑：现有程序TASK_MODE=0，仅保持TARGET_CM=0.0 cm；代码已包含任务3状态机，可在中心保持稳定后切换为O到+5 cm、判稳、再到-5 cm并保持。判稳初值为位置误差不超过0.5 cm、速度不超过1 cm/s并持续500 ms。该状态机目前未启用，也未形成5 s实测结果；任意位置设置、球位端部软限位以及与底盘任务4至6的同步逻辑仍待补充。",
        "任务逻辑：",
    ),
    90: (
        "电源预算：当前只确认系统由车载电池供电，底盘电机、42型步进电机、MaixCAM Pro、MSPM0G3507及传感器应分支稳压并单点共地。各支路输入电压、峰值电流、电池容量、稳压器额定值和纹波尚未提供，表3和表6保留待确认项。联调时应分别测量静止、双轮堵转边缘、步进加速和相机运行工况，按峰值电流留不小于30%的余量，并检查欠压复位、反接和电机回生尖峰。",
        "电源预算：",
    ),
    91: (
        "关键接口：八路红外通过S0/S1/S2和单路ADC复用采集，灰度周期20 ms；AT8236四个桥臂由两组定时器产生20 kHz PWM，换向前关断2 us。底盘程序预留左右编码器和LSM6DSR角速度反馈。MaixCAM Pro当前使用/dev/ttyS0、115200 bit/s向步进驱动发送带0x6B校验的绝对位置帧，默认速度100 r/min、加速度参数80。最终串口引脚、电平、共地、驱动使能和急停接线须按实物复核。",
        "关键接口：",
    ),
    93: (
        "软件架构：底盘和滚球分别在MSPM0G3507与MaixCAM Pro上运行。底盘SysTick只累加毫秒时基，主循环每10 ms调度编码器、IMU、串级控制和安全检查，灰度每20 ms更新；模式切换会清除对应PID历史，通信看门狗和本地按键可停机。MaixCAM Pro以60 fps请求图像，检测每帧运行，控制器每约33 ms更新，显示10 Hz、遥测5 Hz。两侧不交换球位数据，滚球闭环由MaixCAM Pro直接闭合并发送步进指令。",
        "软件架构：",
    ),
    97: (
        "任务模式：任务2所需一圈计时、A点识别与停车状态尚待在底盘程序中加入；任务3代码入口已具备但当前关闭；任务4至6需要把底盘启动事件和滚球目标模式同时置位，并在通过B点或A点时锁存时间而不中断滚球控制。正式程序应统一定义IDLE、CALIBRATE、READY、RUN_T2至RUN_T6、FINISH和FAULT状态，显示模式、目标球位、当前球位、计时及故障码。",
        "任务模式：",
    ),
    102: (
        "联调顺序：第一步悬空验证两台MG513X正反转、编码器方向和急停；第二步完成八路背景/目标线标定并低速检查e_y符号；第三步只闭合轮速和角速度环，再启用循迹方向环；第四步将摆杆机械置水平并保存步进原点，逐点标定摆角-电机角；第五步固定车辆调中心保持，确认丢球回水平；第六步启用正负5 cm状态机；最后才进行底盘与滚球联合动态测试。每一步均以遥测、图像和重复试验通过为进入下一步的依据。",
        "联调顺序：",
    ),
    107: (
        "测试前检查：记录整车长宽、唯一测试位置、摆杆是否超出车身、铰点高度、25 cm PPR管内壁和刻度、钢球直径、摄像头全杆覆盖、车载供电电压以及图传录像完整性。当前机械装配和控制链仍在调试，所有尺寸、供电和图传项目均须在最终结构冻结后实测填写。",
        "测试前检查：",
    ),
    109: (
        "统一流程：每项试验先完成灰度背景/目标线标定、步进回原点和球心零点标定，检查电池电压并选择任务；将车辆与钢球放到题目规定初始位置，按同一启动键开始计时，记录车载显示、接收端视频和串口日志。计时终点按通过B点、通过A点或停车事件区分。每项至少连续测试5次，脱轨、掉球、丢图或未完成的数据原样保留。",
        "统一流程：",
    ),
    111: (
        "数据处理约定：摆杆坐标向执行机构端为正，位置误差取x-x_r的绝对值；最大误差为有效测试区间内max|x-x_r|，平均绝对误差为所有有效采样的|x-x_r|均值。停车偏差取车辆唯一测试位置与A点基准线的绝对距离。异常值不人工删除，只在日志中标明原因；时间保留0.01 s、位置保留0.1 cm，并同时保存原始值。",
        "数据处理约定：",
    ),
    127: (
        "指定位置说明：5次目标位置应覆盖中心两侧且远离杆端安全区，例如由裁判或队员在允许范围内逐次指定；最终位置范围和0.1 cm设定分辨率待程序实现后确认。每次测试前应回到已保存的机械原点，并检查u_O和像素比例；只有相机或支架发生位移时才重新做完整标定。",
        "指定位置说明：",
    ),
    128: (
        "证据索引：当前尚无任务6录像和控制日志。后续按“T6_日期_次数_目标位置”统一命名视频与CSV日志，并记录开始、再次通过A点和最大误差所在时间戳；表B-6每一行必须能追溯到对应文件，不允许只保留统计值。",
        "证据索引：",
    ),
    133: (
        "误差来源：一是钢球高光、槽壁阴影和相机安装变化会移动灰度暗度峰值，影响像素位置；应通过固定曝光、遮光、端点标定和误检日志验证。二是60 fps采集、30 Hz控制和150 ms估速窗口引入相位延迟，增益过高时会造成振荡；应测端到端延迟并按速度环先行的顺序整定。三是步进电机、摇臂、连杆和铰点的回差使同一指令对应不同摆角；应从同一方向逼近标定并测往返差。四是底盘加减速、轮胎打滑和电池压降改变扰动与执行能力；应同步记录轮速、电压和球位并分段限速。",
        "误差来源：",
    ),
    134: (
        "失败记录：当前处于初级调试阶段，尚未形成可用于评分判定的5次连续测试样本。调试中出现的丢线、球位无效、步进回零失败、串口无响应、超时回水平和未完成任务均应记录发生条件、遥测值、视频时间段、处理措施及复测结果；在表B-1至表B-6填满前不得写“全部成功”。",
        "失败记录：",
    ),
    137: (
        "结论：本阶段已完成基于MSPM0G3507、八路红外、AT8236和两台MG513X的底盘控制程序框架，以及基于MaixCAM Pro灰度视觉、位置/速度估计、级联控制和42型步进电机连杆机构的滚球控制链。底盘工程最新构建记录为0错误、0警告，滚球程序已具备中心保持、丢球回水平和未启用的正负5 cm状态机。由于循迹一圈停车、图传装置和表B-1至表B-6尚未实测，本稿不能判定任务2至任务6是否达标；最终结论必须由连续测试数据替换。",
        "结论：",
    ),
    138: (
        "不足与改进：当前最明确的瓶颈是循迹软件上限0.30 m/s对应理想一圈约20.47 s，尚无20 s任务裕量；应在低速稳定后提高直线速度并加入A点预减速和制动。滚球端需要实测22.4 pixel/cm比例、连杆标定表、步进回差和视觉延迟，再决定是否加入积分或模型前馈。图传发送/接收与录像链尚未接入，应优先完成，避免控制完成后仍无法满足任务1。",
        "不足与改进：",
    ),
    140: "以下资料用于当前设计依据，最终提交时应补齐版本号、发布日期和访问日期。",
    141: "[1] 2026年全国大学生电子设计竞赛组委会. 车载平衡滚球运动控制系统（H题）[Z]. 2026.",
    142: "[2] Texas Instruments. MSPM0G350x Mixed-Signal Microcontrollers Data Sheet[EB/OL].",
    143: "[3] AT8236双通道直流电机驱动芯片数据手册[EB/OL].",
    144: "[4] Sipeed. MaixCAM Pro与MaixPy官方文档[EB/OL].",
    145: "[5] 张大头闭环步进驱动器. X42S串口通信协议与使用说明[EB/OL].",
    146: "[6] 胡寿松. 自动控制原理[M]. 北京: 科学出版社.",
    151: (
        "核心代码说明：底盘核心入口为MSPM0G3507_Diansai_Test/Code/vehicle_line_control.c中的VehicleLine_Update()，每20 ms读取八路归一化值并输出目标偏航角速度，随后由vehicle_cascade_control.c生成左右轮目标与PWM；电机桥控制位于vehicle_motor.c。滚球核心入口为main.py中的GrayProfileBallDetector.detect()、VelocityEstimator.update()和CascadeController.update()，约33 ms计算一次目标摆角，再由pipe_angle_to_motor_angle()和MotorController.goto_motor_angle()发送绝对位置帧。对应正文公式见第二章第1、3、4节。",
        "核心代码说明：",
    ),
    152: (
        "最终自检：正文已将母版提示替换为实际方案或“待确认/待实测”，并保留所有原表格和章节位置。封面按本次要求原样保留，因此仍含母版日期提示和“填写通用模板”字样；最终参赛提交前是否填写日期或删除提示，必须由队伍另行确认。完成测试后还需搜索“待实测”“待确认”，核对图表编号和单位，并在Word中按Ctrl+A、F9更新目录与页码。",
        "最终自检：",
    ),
}


TABLES = {
    1: [
        ["题目任务", "约束/指标", "本队设计对策", "验证位置"],
        ["图传与回放", "覆盖全摆杆、连续记录", "独立图传链待选型并接入；控制相机与图传职责分开", "表B-1/视频索引"],
        ["循迹一圈并停车", "t<=20 s，偏差<=2 cm", "八路归一化循迹；里程+A线识别、预减速和停车逻辑待完成", "表B-2"],
        ["静态O到+5到-5 cm", "t<=5 s，误差<=1 cm", "MaixCAM Pro位置/速度级联；任务3状态机已编写但当前关闭", "表B-3"],
        ["A-B段中心稳定", "t<=8 s，误差<=1 cm", "底盘分段限速；滚球闭环持续运行", "表B-4"],
        ["一圈中心/指定位置", "t<=30 s，误差<=1 cm", "目标位置可配置接口待补充；动态抗扰待联调", "表B-5/B-6"],
    ],
    2: [
        ["类别", "参数", "设计/实测值", "单位", "来源或选择依据"],
        ["底盘", "轮距b", "待实测", "m", "最终装配测量"],
        ["底盘", "直线/弯道速度", "当前调试0.20；目标待定", "m/s", "时间预算与表B-2"],
        ["视觉", "分辨率/帧率/延迟", "640x480；设定60；延迟待实测", "pixel、Hz、ms", "main.py与表B-1"],
        ["执行", "额定转矩/最大摆角", "转矩待确认；正常正负2", "N*m、deg", "电机资料与程序限幅"],
        ["控制", "位置环采样周期", "约33", "ms", "CONTROL_HZ=30"],
        ["供电", "电池电压/容量", "待确认", "V、mAh", "实物规格与续航测试"],
    ],
    3: [
        ["功能模块", "型号/关键规格", "与主控接口", "选择依据", "实物数量"],
        ["主控制器", "MSPM0G3507", "-", "定时器、ADC、串口资源；现有工程已编译", "1"],
        ["循迹传感器", "8路模拟红外模块", "S0/S1/S2+ADC", "逐路标定、连续加权偏差", "1套"],
        ["电机与编码器", "MG513X直流减速电机", "PWM；编码器接口按实物", "两轮差速，转速/转矩待确认", "2"],
        ["电机驱动", "AT8236双路驱动模块", "4路20 kHz PWM", "慢衰减驱动与2 us换向保护", "1套"],
        ["球位摄像头", "MaixCAM Pro", "板载Camera/Display", "640x480 ROI处理与本地控制", "1"],
        ["摆杆执行机构", "42型步进电机+摇臂连杆", "UART 115200", "绝对位置命令；H42机械装配", "1套"],
        ["图传/显示/存储", "发送、接收和录像装置待选型", "独立链路", "必须连续覆盖全摆杆；尚未接入", "待确认"],
        ["电源模块", "车载电池+分路稳压", "-", "电压、电流、纹波和余量待实测", "1套"],
    ],
    5: [
        ["任务", "周期/触发条件", "主要输入", "主要输出", "最坏执行时间"],
        ["系统时基/编码器采样", "1 ms时基；10 ms控制调度", "编码器计数、SysTick", "轮速、时间戳", "待实测 us"],
        ["轮速与循迹控制", "10 ms；灰度20 ms", "八路灰度、角速度、轮速", "左右轮PWM", "待实测 us"],
        ["视觉接收/球位估计", "相机设定60 fps", "640x480 ROI图像", "球位、有效标志、速度", "待实测 ms"],
        ["滚球控制输出", "约33 ms", "目标、球位、球速", "摆角和步进绝对位置", "待实测 ms"],
        ["任务状态机/计时显示", "事件触发；显示10 Hz", "按键、模式、过点事件", "状态与计时", "待实测 us"],
        ["故障监测/日志记录", "每控制周期；遥测200 ms", "丢线、丢球、通信、电压", "停机/回水平/日志", "待实测 us"],
    ],
    7: [
        ["对象", "标定/检测方法", "保存参数", "允许范围", "异常动作"],
        ["循迹阵列", "K1采背景、K2采目标线；逐路归一化", "8路背景/目标线值", "每路跨度>=32 ADC count", "拒绝循迹并报警"],
        ["左右轮/编码器", "悬空方向检查与定长测速", "方向、比例、PID", "待实测", "停机并清控制历史"],
        ["摆杆水平与角度", "水平保存原点；分点测摆角-电机角", "分段线性标定表", "杆角正负2 deg；机械角正负40 deg", "停止并回安全角"],
        ["视觉中心与比例", "O点20帧均值；刻度点拟合", "u_O、pixel/cm", "零点极差<=0.5 cm", "超时用配置值并告警"],
        ["球位软限位", "按杆端安全距离设定", "位置上下限", "待最终装配测量", "冻结积分并回水平"],
        ["电池/通信/丢帧", "电压、UART响应和检测时间戳", "阈值与超时", "视觉长丢失150 ms；其余待测", "底盘停车或摆杆回水平"],
    ],
    8: [
        ["项目", "型号/设定或实测值", "精度/分辨率", "用途", "备注"],
        ["场地直线段/圆弧半径", "1.5 m / 0.5 m，待复核", "待记录", "场地复核", "按题图"],
        ["黑线宽度/材质/照度", "1.8+/-0.2 cm；材质/照度待测", "待记录", "环境记录", "至少三种照度"],
        ["直尺/卡尺", "待填写型号", "待填写", "位置与尺寸", "测停车和机构"],
        ["秒表/视频时间戳", "车载计时+接收端视频", "待校验", "运行时间", "双通道核对"],
        ["示波器/万用表", "待填写型号", "待填写", "PWM、UART和供电", "记录电池压降"],
        ["图传接收与录像设备", "待选型", "待实测", "显示与回放", "尚未接入"],
        ["电池状态", "电压待测，使用时间待记录", "待填写", "一致性控制", "每轮测试前记录"],
    ],
    15: [
        ["任务", "题目限值", "本队最不利结果", "是否达成", "证据位置"],
        ["图传", "按题目要求", "待实测", "未判定", "视频待建立/表B-1"],
        ["循迹停车", "<=20 s；<=2 cm", "待实测", "未判定", "表B-2"],
        ["静态往返", "<=5 s；<=1 cm", "待实测", "未判定", "表B-3"],
        ["A-B中心稳定", "<=8 s；<=1 cm", "待实测", "未判定", "表B-4"],
        ["一圈中心稳定", "<=30 s；<=1 cm", "待实测", "未判定", "表B-5"],
        ["任意位置稳定", "<=30 s；<=1 cm", "待实测", "未判定", "表B-6"],
    ],
    16: [
        ["材料", "文件名/路径", "对应任务", "说明"],
        ["主控程序", "MSPM0G3507_Diansai_Test", "循迹、底盘安全", "Keil MDK-ARM 6.19；v24构建0错误0警告"],
        ["视觉程序", "main.py", "球位识别与滚球控制", "MaixCAM Pro；灰度剖面与30 Hz级联"],
        ["电路图/PCB", "待整理", "底盘与电源", "版本待确认"],
        ["机械图/尺寸图", "26年H题车载水平小球/H42装配.stp", "摆杆与执行机构", "SolidWorks装配与STEP；尺寸待冻结"],
        ["原始测试数据", "待建立data目录", "B-1至B-6", "必须保留逐次原始值"],
        ["录像与日志", "待建立video/log目录", "任务1至6", "时间戳对应关系待填写"],
    ],
}


def fill_test_table(table, statistics_word: str = "待统计") -> None:
    for r_index, row in enumerate(table.rows):
        if r_index == 0:
            continue
        for c_index, cell in enumerate(row.cells):
            if c_index == 0:
                continue
            set_cell(cell, statistics_word if c_index == len(row.cells) - 1 else "待实测")


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def main() -> None:
    make_interface_figure()
    source = Document(TEMPLATE)
    cover_xml = [etree.tostring(p._p, with_tail=False) for p in source.paragraphs[:5]]
    cover_ppr = [
        etree.tostring(p._p.pPr, with_tail=False) if p._p.pPr is not None else b""
        for p in source.paragraphs[:5]
    ]
    cover_text = [p.text for p in source.paragraphs[:5]]
    toc_text = [p.text for p in source.paragraphs[10:33]]
    headings = [p.text for p in source.paragraphs if p.style.name.startswith("Heading")]
    table_shapes = [(len(t.rows), len(t.columns)) for t in source.tables]
    sections = section_signature(source)
    media_before = archive_media_hashes(TEMPLATE)

    for index, value in PARAGRAPHS.items():
        if isinstance(value, tuple):
            text, label = value
            set_paragraph(source.paragraphs[index], text, label)
        else:
            set_paragraph(source.paragraphs[index], value)

    for table_index, rows in TABLES.items():
        set_table_matrix(source.tables[table_index], rows)

    set_image_cell(source.tables[0].cell(0, 0), FIG1, 5.55, "系统总体结构框图")
    set_image_cell(source.tables[4].cell(0, 0), FIG2, 5.55, "控制系统接口与信号流")
    set_image_cell(source.tables[6].cell(0, 0), FIG3, 5.35, "系统软件主流程")

    for table_index in range(9, 14):
        fill_test_table(source.tables[table_index])

    # Table B-6 has a different orientation and a statistics row.
    table_b6 = source.tables[14]
    for r_index in range(1, 6):
        for c_index in range(1, len(table_b6.columns)):
            set_cell(table_b6.cell(r_index, c_index), "待实测")
    for c_index in range(1, len(table_b6.columns)):
        set_cell(table_b6.cell(6, c_index), "待统计" if c_index != 1 else "-")

    set_update_fields_on_open(source)
    source.core_properties.title = "车载平衡滚球运动控制系统（H题）阶段初稿"
    source.core_properties.subject = "严格沿用2026 H题选手填写通用模板"
    source.save(OUTPUT)

    result = Document(OUTPUT)
    assert len(result.paragraphs) == 153
    assert len(result.tables) == 17
    assert len(result.sections) == 3
    assert [p.text for p in result.paragraphs[:5]] == cover_text
    assert [
        etree.tostring(p._p, with_tail=False) for p in result.paragraphs[:5]
    ] == cover_xml
    assert [
        etree.tostring(p._p.pPr, with_tail=False) if p._p.pPr is not None else b""
        for p in result.paragraphs[:5]
    ] == cover_ppr
    assert [p.text for p in result.paragraphs[10:33]] == toc_text
    assert [p.text for p in result.paragraphs if p.style.name.startswith("Heading")] == headings
    assert [(len(t.rows), len(t.columns)) for t in result.tables] == table_shapes
    assert section_signature(result) == sections
    media_after = archive_media_hashes(OUTPUT)
    for name, digest in media_before.items():
        assert media_after.get(name) == digest, f"template media changed: {name}"

    body_text = "\n".join(p.text for p in result.paragraphs[5:])
    body_text += "\n" + "\n".join(
        cell.text for table in result.tables for row in table.rows for cell in row.cells
    )
    forbidden = ["K230", "YOLOv8", "TB6612", "七路红外"]
    leaked = [token for token in forbidden if token in body_text]
    assert not leaked, f"reference-plan terms leaked: {leaked}"
    assert "【请填写" not in body_text
    assert "____" not in body_text

    print(f"created={OUTPUT}")
    print(f"paragraphs={len(result.paragraphs)} tables={len(result.tables)} sections={len(result.sections)}")
    print(f"inline_shapes={len(result.inline_shapes)} media={len(media_after)}")
    print("cover_c14n=IDENTICAL")
    print("toc_text=IDENTICAL")
    print("headings=IDENTICAL")
    print("sections=IDENTICAL")
    print("template_media=IDENTICAL")


if __name__ == "__main__":
    main()
