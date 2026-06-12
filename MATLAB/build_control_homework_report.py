from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE = Path(r"C:\Users\28097\Desktop\my_code\my_code\MATLAB")
RES = BASE / "results_control_homework"
OUT = BASE / "control_homework_report.docx"
CODE_PATH = BASE / "control_homework_image_tasks.m"


def set_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def para(doc, text="", bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, bold=bold)
    if align is not None:
        p.alignment = align
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_font(r, "黑体", 14 if level == 1 else 12, True)


def equation(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "  ")
    set_font(r1, bold=True)
    r2 = p.add_run(text)
    set_font(r2, "Cambria Math", 11)


def picture(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        set_font(r, size=9)


def table_rows(doc, rows, caption):
    para(doc, caption, bold=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9, bold=(i == 0))


def code_block(doc, title, code):
    heading(doc, title, 2)
    p = doc.add_paragraph()
    r = p.add_run(code)
    set_font(r, "Consolas", 8)


doc = Document()
doc.styles["Normal"].font.name = "宋体"
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
doc.styles["Normal"].font.size = Pt(10.5)

sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MATLAB 控制系统作业报告")
set_font(r, "黑体", 18, True)
para(doc, "姓名：__________    学号：__________    班级：__________", align=WD_ALIGN_PARAGRAPH.CENTER)

heading(doc, "一、微分方程通解", 1)
para(doc, "题目要求求解一阶线性微分方程：")
equation(doc, "微分方程：", "dy/dx + 2xy = x e^{-x^2}")
para(doc, "使用积分因子法或 MATLAB 符号求解可得通解：")
equation(doc, "通解：", "y = (C + x^2/2)e^{-x^2}")
para(doc, "MATLAB 输出形式为 C1*exp(-x^2) + (x^2*exp(-x^2))/2，与上式等价。")

heading(doc, "二、单位负反馈系统分析", 1)
para(doc, "开环传递函数为：")
equation(doc, "开环传递函数：", "G(s)=25/[s(s+5)]")
para(doc, "单位负反馈闭环传递函数为：")
equation(doc, "闭环传递函数：", "Φ(s)=G(s)/(1+G(s))=25/(s^2+5s+25)")

heading(doc, "三、单位阶跃响应与峰值时间", 1)
picture(doc, RES / "closed_loop_step_response.png", "图 1 闭环系统单位阶跃响应")
table_rows(
    doc,
    [
        ["指标", "数值"],
        ["峰值时间 Tp", "0.7184 s"],
        ["超调量", "16.2929 %"],
        ["调节时间", "1.6152 s"],
    ],
    "表 1 闭环系统阶跃响应指标",
)
para(doc, "由阶跃响应可见，系统为欠阻尼二阶系统，存在一定超调，响应最终稳定到 1。")

heading(doc, "四、根轨迹分析", 1)
picture(doc, RES / "root_locus_comparison.png", "图 2 原系统与增加极点后的根轨迹对比", 6.0)
para(doc, "原系统开环极点位于 s=0 和 s=-5。增加一个极点 s=-10 后，系统开环极点数量增加，根轨迹分支和渐近线位置发生变化，闭环极点随增益变化的趋势也随之改变。")
para(doc, "从图中可以观察到，附加极点会改变根轨迹在复平面中的走向，从而影响系统稳定性和动态性能。")

heading(doc, "五、伯德图分析", 1)
picture(doc, RES / "bode_open_loop.png", "图 3 开环系统 G(s) 的伯德图")
para(doc, "伯德图反映系统幅频特性和相频特性，可用于分析系统的频域稳定裕度、截止频率以及不同频率输入下的响应能力。")

heading(doc, "六、结论", 1)
para(doc, "本实验使用 MATLAB 完成了微分方程符号求解、闭环传递函数建立、阶跃响应绘制、峰值时间计算、根轨迹分析和伯德图绘制。")
para(doc, "闭环系统 Φ(s)=25/(s^2+5s+25) 为稳定欠阻尼二阶系统，峰值时间约为 0.7184 s，超调量约为 16.29%。")
para(doc, "根轨迹结果表明，增加开环极点会明显改变根轨迹形状，因此会影响闭环极点位置和系统动态性能。")

doc.add_page_break()
heading(doc, "附录：MATLAB 完整代码", 1)
code_block(doc, "control_homework_image_tasks.m", CODE_PATH.read_text(encoding="utf-8"))

doc.save(OUT)
print(OUT)
