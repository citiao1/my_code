from pathlib import Path
import csv

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


BASE = Path(r"C:\Users\28097\Desktop\my_code\my_code\MATLAB")
RES = BASE / "results_lab4_5_6"
ASSETS = BASE / "report_assets"
OUT = BASE / "auto_control_report_lab4_5_6.docx"


def set_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def setup(doc):
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def p(doc, text="", bold=False, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, bold=bold)
    if align is not None:
        para.alignment = align
    return para


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_font(run, "黑体", 14 if level == 1 else 12, True)


def pic(doc, path, caption, width=5.8):
    path = Path(path)
    if not path.exists():
        p(doc, f"【缺少图片：{path.name}】")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_font(run, size=9)


def rows_from_csv(path):
    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.reader(f))


def fmt(v):
    try:
        x = float(v)
    except Exception:
        return v
    if abs(x) < 1e-10:
        return "0"
    if abs(x) >= 100:
        return f"{x:.2f}"
    return f"{x:.4f}"


HEADER_MAP = {
    "Case": "工况",
    "Stable": "是否稳定",
    "Overshoot_percent": "超调量/%",
    "PeakTime_s": "峰值时间/s",
    "SettlingTime_s": "调节时间/s",
    "FinalValue": "稳态值",
    "SteadyStateError": "稳态误差",
    "MaxOvershootValue": "最大超调值",
    "Range": "稳定区间",
    "Kc_min": "Kc最小值",
    "Kc_max": "Kc最大值",
    "Kmin": "K最小值",
    "Kmax": "K最大值",
    "error_open": "开环稳态误差",
    "error_close_K1": "闭环误差(Kc=1)",
    "error_close_K05": "闭环误差(Kc=0.5)",
    "ratio_K1": "误差比(Kc=1)",
    "ratio_K05": "误差比(Kc=0.5)",
    "CharacteristicEquation": "闭环特征方程",
    "RouthConclusion": "劳斯判据结论",
    "ErrorAt10s_deg": "10s后误差/deg",
}

VALUE_MAP = {
    "standard second order": "标准二阶系统",
    "added pole far": "附加极点远离虚轴",
    "added pole comparable": "附加极点接近原极点",
    "added pole dominant": "附加极点成为主导极点",
    "first order reference": "一阶系统对照",
    "added zero far": "附加零点远离虚轴",
    "added zero comparable": "附加零点接近原极点",
    "added zero dominant": "附加零点靠近虚轴",
    "stable range 1": "稳定区间1",
    "no added zero": "未增加开环零点",
    "added zero z=-4": "附加零点 z=-4",
    "added zero z=-1": "附加零点 z=-1",
    "unstable for K>0": "K>0时始终不稳定",
    "can be stable for K>0": "K>0时可稳定",
}


def translate(value, header=False):
    if header:
        return HEADER_MAP.get(value, value)
    return VALUE_MAP.get(value, value)


def csv_table(doc, path, caption):
    rows = rows_from_csv(path)
    p(doc, caption, bold=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = translate(val, True) if i == 0 else fmt(translate(val))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=9, bold=(i == 0))


def code_block(doc, title, code):
    p(doc, title, bold=True)
    para = doc.add_paragraph()
    run = para.add_run(code.strip())
    set_font(run, "Consolas", 8)


def eq(doc, label, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = para.add_run(label + "  ")
    set_font(r1, bold=True)
    r2 = para.add_run(text)
    set_font(r2, "Cambria Math", 11)


doc = Document()
setup(doc)
sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("自动控制原理实验报告")
set_font(r, "黑体", 18, True)
p(doc, "实验四：传递函数零极点对系统过渡过程的影响", align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "实验五：PID 调节规律对系统调节质量的影响", align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "实验六：根轨迹实验", align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "姓名：__________    学号：__________    班级：__________", align=WD_ALIGN_PARAGRAPH.CENTER)

h(doc, "实验四 传递函数零极点对系统过渡过程的影响", 1)
h(doc, "一、实验目的", 2)
p(doc, "研究传递函数的闭环零点、闭环极点位置变化对系统单位阶跃响应和动态性能指标的影响，并理解高阶系统中闭环主导极点的作用。")
h(doc, "二、实验内容与 MATLAB 实现", 2)
eq(doc, "标准二阶系统：", "G0(s)=1/(s^2+s+1)")
eq(doc, "附加极点：", "T=0.2、2、5，对应 -5、-0.5、-0.2")
eq(doc, "附加零点：", "τ=0.2、2、5，对应 -5、-0.5、-0.2")
code_block(doc, "实验四核心代码：", """
s = tf('s'); zeta = 0.5; wn = 1; t = 0:0.01:30;
G0 = wn^2/(s^2+2*zeta*wn*s+wn^2);
T = [0.2 2 5]; tau = [0.2 2 5];
Gp1 = G0/(T(1)*s+1); Gp2 = G0/(T(2)*s+1); Gp3 = G0/(T(3)*s+1);
Gz1 = (tau(1)*s+1)*G0; Gz2 = (tau(2)*s+1)*G0; Gz3 = (tau(3)*s+1)*G0;
step(G0,Gp1,Gp2,Gp3,t); grid on;
figure; pzmap(G0,Gp1,Gp2,Gp3); grid on;
figure; step(G0,Gz1,Gz2,Gz3,t); grid on;
figure; pzmap(G0,Gz1,Gz2,Gz3); grid on;
""")
h(doc, "三、实验结果", 2)
pic(doc, RES / "lab4_added_poles_step.png", "图 1 增加闭环极点时的单位阶跃响应")
pic(doc, RES / "lab4_added_poles_pzmap.png", "图 2 增加闭环极点时的零极点分布")
csv_table(doc, RES / "lab4_added_poles_metrics.csv", "表 1 增加闭环极点时的动态指标")
pic(doc, RES / "lab4_added_zeros_step.png", "图 3 增加闭环零点时的单位阶跃响应")
pic(doc, RES / "lab4_added_zeros_pzmap.png", "图 4 增加闭环零点时的零极点分布")
csv_table(doc, RES / "lab4_added_zeros_metrics.csv", "表 2 增加闭环零点时的动态指标")
p(doc, "标准二阶系统的闭环极点为 -0.5±0.8660j，超调量约为 16.30%，调节时间约为 8.08 s。")
p(doc, "当附加极点为 -5 时，该极点远离虚轴，响应曲线与标准二阶系统接近，影响较小。")
p(doc, "当附加极点靠近虚轴，特别是 -0.2 时，它成为主导极点，系统响应明显变慢，调节时间增大。")
p(doc, "附加零点离虚轴较远时影响较小；附加零点靠近虚轴时会明显加快初始响应，但超调显著增大。")

h(doc, "实验五 PID 调节规律对系统调节质量的影响", 1)
h(doc, "一、实验目的", 2)
p(doc, "通过改变 P、PI、PID 调节规律及参数，观察飞机自动驾驶仪航向角过渡过程，比较峰值时间、过渡时间、超调量、余差和衰减比等指标。")
h(doc, "二、实验模型与参数", 2)
eq(doc, "等效对象：", "G0(s)=10(s+5)/[(s+10)(s^2+3.5s+6)]")
pic(doc, ASSETS / "2d3c22b567f5e97efefc202d0849d9e7.png", "图 5 PID 控制仿真结构")
pic(doc, ASSETS / "6db688e3a347e1f1467ae738bb65a8c4.png", "图 6 PID 控制器参数：P=6，I=11，D=2，N=100")
p(doc, "实验指导书中 PID 形式为 Kc(1+1/(Ti s)+Td s)，仿真软件并行 PID 形式为 P+I/s+D s，两者换算关系为 P=Kc，I=Kc/Ti，D=KcTd。")
h(doc, "三、阶跃响应结果", 2)
pic(doc, ASSETS / "3f7bf66d4bc0b33e128e8f7dab31c97c.png", "图 7 PID 控制单位阶跃响应曲线", 6.0)
csv_table(doc, RES / "lab5_all_step_metrics.csv", "表 3 不同调节规律的阶跃响应指标")
p(doc, "你给出的最终 PID 结果为：峰值时间 2.8570 s、过渡时间 1.8246 s、最大偏差 0.0037、超调量 0.37%、稳态余差约 0、衰减比 0.98。")
h(doc, "四、斜坡输入结果", 2)
pic(doc, ASSETS / "ada74221e0c4110f4913e87961173072.png", "图 8 斜坡输入下航向角响应曲线")
csv_table(doc, RES / "lab5_ramp_errors.csv", "表 4 斜坡输入 10 s 后误差")
p(doc, "输入为 θd(t)=0.5t，比例控制器 Kc=2 时，10 s 后航向角误差约为 1.9316 或 0.0598，具体以你最终确认的仿真版本为准。")
h(doc, "五、结果分析", 2)
p(doc, "比例作用能够提高响应速度、减小余差，但比例增益过大会增大超调和振荡。")
p(doc, "积分作用能够消除阶跃输入下的稳态余差，但积分过强会降低稳定裕度。")
p(doc, "微分作用根据误差变化趋势提前修正，可改善动态品质并减小超调。")

h(doc, "实验六 根轨迹实验", 1)
h(doc, "内容 1：二阶水槽系统根轨迹", 2)
eq(doc, "参数：", "A1=1000, A2=800, R1=0.005, R2=0.005, Kv=1250, KB=1, TB=0.5")
eq(doc, "环节：", "G1=1/(A1R1s+1),  G2=R2/(A2R2s+1),  GB=KB/(TBs+1)")
eq(doc, "等效开环：", "L0(s)=KvG1G2GB")
pic(doc, ASSETS / "9ae66d961515a5bc9a9429cc390e0c10.png", "图 9 二阶水槽系统资料与教材模型")
code_block(doc, "实验六内容 1 核心代码：", """
s = tf('s');
A1=1000; A2=800; R1=0.005; R2=0.005; Kv=1250; KB=1; TB=0.5;
G1 = 1/(A1*R1*s+1); G2 = R2/(A2*R2*s+1); GB = KB/(TB*s+1);
L0 = minreal(Kv*G1*G2*GB);
rlocus(L0); grid on;
Hopen = G2;
Hclose_K1 = feedback(G2,1*Kv*G1*GB);
Hclose_K05 = feedback(G2,0.5*Kv*G1*GB);
step(Hopen,'k--',Hclose_K1,'b-',Hclose_K05,'r-',0:0.01:40); grid on;
""")
pic(doc, RES / "lab6_water_tank_root_locus.png", "图 10 二阶水槽系统根轨迹")
csv_table(doc, RES / "lab6_water_tank_stable_range.csv", "表 5 二阶水槽系统比例增益稳定范围")
pic(doc, RES / "lab6_water_tank_disturbance_response.png", "图 11 单位阶跃扰动下的液位响应")
csv_table(doc, RES / "lab6_water_tank_disturbance_metrics.csv", "表 6 二阶水槽系统扰动稳态误差")
p(doc, "由根轨迹扫描可得系统稳定的比例增益范围约为 0≤Kc≤3.56。")
p(doc, "扰动响应中，开环稳态误差为 0.005；Kc=1 时闭环误差约为 0.00069；Kc=0.5 时闭环误差约为 0.00121。")

h(doc, "内容 2：增加开环零点对稳定性的影响", 2)
eq(doc, "开环对象：", "G(s)=K/[s^2(s+a)]，本实验取 a=2")
pic(doc, ASSETS / "5de0e4de88534ff30c4d3530ffe7f3c7.png", "图 12 增加开环零点对根轨迹和稳定性的影响")
csv_table(doc, RES / "lab6_content2_routh_conclusions.csv", "表 7 内容 2 Routh 判据结论")
p(doc, "未增加零点时，闭环特征方程为 s^3+a s^2+K=0；增加负开环零点 z=-b 后，闭环特征方程变为 s^3+a s^2+K s+Kb=0。")
p(doc, "Routh 判据要求 b<a 才可能稳定，因此 b=4>a=2 时系统始终不稳定，b=1<a=2 时可通过调节 K 使系统稳定。")

h(doc, "总 结", 1)
p(doc, "实验四表明，靠近虚轴的闭环极点会成为主导极点并显著减慢响应；靠近虚轴的闭环零点会明显改变过渡过程并可能增大超调。")
p(doc, "实验五表明，P、I、D 三种调节作用分别影响响应速度、稳态误差和动态阻尼；合理整定 PID 参数后系统可满足给定动态指标。")
p(doc, "实验六表明，根轨迹能够直观判断比例增益变化对闭环极点和稳定性的影响；对于二阶水槽系统可得到稳定增益范围，对于附加零点系统可由 Routh 判据判断 b<a 是稳定的重要条件。")

doc.save(OUT)
print(OUT)
