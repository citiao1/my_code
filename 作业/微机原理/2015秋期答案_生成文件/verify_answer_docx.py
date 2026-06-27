from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "必修-CSE22500E《微机原理及应用》电科2015秋期_答案解析.docx"

doc = Document(DOCX)
headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
text = "\n".join(p.text for p in doc.paragraphs)
text += "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

required = [
    "填空题",
    "选择题",
    "编程和读程题",
    "综合题",
    "AND BX,0FEF7H",
    "XOR BYTE PTR [100H],03H",
    "CMP AL,BL",
    "8255 查询输入完整程序",
    "秒表系统连接",
    "INT3 PROC FAR",
    "OUT 80H,AL",
]

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
print(f"headings={headings}")
for item in required:
    print(f"required[{item}]={item in text}")
