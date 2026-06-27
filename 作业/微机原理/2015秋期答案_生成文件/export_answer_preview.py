from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "必修-CSE22500E《微机原理及应用》电科2015秋期_答案解析.docx"
OUT = Path(__file__).resolve().parent / "answer_preview.txt"

doc = Document(DOCX)
with OUT.open("w", encoding="utf-8") as f:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            f.write(paragraph.text + "\n")
    for idx, table in enumerate(doc.tables, 1):
        f.write(f"\n[TABLE {idx}]\n")
        for row in table.rows:
            f.write(" | ".join(cell.text.replace("\n", " ") for cell in row.cells) + "\n")

print(OUT)
