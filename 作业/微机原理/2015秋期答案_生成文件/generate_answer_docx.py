from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "必修-CSE22500E《微机原理及应用》电科2015秋期_答案解析.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def set_table_borders(table, color="9AA6B2", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 20, BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "北京化工大学《微机原理及应用》2015 秋期答案解析"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x66, 0x77, 0x88)


def add_title(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("《微机原理及应用》2015 秋期期末试卷答案解析")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("适用于 CSE22500E 电科试卷；包含解析、完整程序和连接图")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x77, 0x88)


def add_para(doc, text, style=None, bold_label=False):
    p = doc.add_paragraph(style=style)
    if bold_label and "：" in text:
        label, rest = text.split("：", 1)
        p.add_run(label + "：").bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.12)
    r = p.add_run(code.strip())
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(8.5)


def add_answer_table(doc, rows, widths=(1.2, 5.1), header=None):
    table = doc.add_table(rows=0, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if header:
        row = table.add_row()
        for i, text in enumerate(header):
            row.cells[i].width = Inches(widths[i])
            set_cell_text(row.cells[i], text, bold=True)
            set_cell_shading(row.cells[i], LIGHT_BLUE)
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for cells in rows:
        row = table.add_row()
        for i, text in enumerate(cells):
            row.cells[i].width = Inches(widths[i])
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_text(row.cells[i], text, bold=True)
                set_cell_shading(row.cells[i], LIGHT_GRAY)
            else:
                row.cells[i].text = ""
                p = row.cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                run.font.name = "Calibri"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_memory_diagram(doc):
    doc.add_heading("存储器连接图（128Kx8，由 16Kx1 芯片构成）", level=3)
    rows = [
        ["CPU", "A0-A13", "接所有 RAM 芯片的片内地址线，共 14 根"],
        ["CPU", "A14-A16", "接 74LS138 的 A、B、C 输入，产生 8 组选通信号"],
        ["CPU", "D0-D7", "每组 8 片 16Kx1 分别接 D0-D7，位扩展成 16Kx8"],
        ["CPU", "RD/WR", "分别接各芯片 OE/WE 或读写控制端"],
        ["74LS138", "Y0-Y7", "分别接第 0-7 组 RAM 的 CS，地址从低到高排列"],
    ]
    add_answer_table(doc, rows, widths=(1.1, 1.15, 4.25), header=["部件", "信号", "连接关系"])
    add_para(doc, "地址范围：若高位 A17-A19 未参与译码，则存在地址覆盖。第一组（A16-A14=000）对应 00000H-03FFFH、20000H-23FFFH、40000H-43FFFH、60000H-63FFFH、80000H-83FFFH、A0000H-A3FFFH、C0000H-C3FFFH、E0000H-E3FFFH。")


def add_stopwatch_diagram(doc):
    doc.add_heading("秒表电路连接图与程序结构", level=3)
    rows = [
        ["8259A", "IR3", "接 100Hz 脉冲输入，产生类型号 13H 的中断请求"],
        ["8259A", "INT/INTA", "与 CPU 的 INTR/INTA 相连，命令端口 80H，数据端口 81H"],
        ["8255A", "PA0-PA7", "接数码管段码 a-g,dp，输出共阴极段码"],
        ["8255A", "PC0", "接数码管位码，PC0=0 点亮"],
        ["8255A", "控制口", "地址 283H，方式 0，A 口和 C 口输出，控制字 80H"],
        ["程序", "向量表", "0000:004C 写 INT3 的 IP，0000:004E 写 INT3 的 CS"],
    ]
    add_answer_table(doc, rows, widths=(1.1, 1.15, 4.25), header=["模块", "信号/端口", "说明"])


CODE_MIN = """
; 比较 AL、BL、CL 中无符号数，将最小值存入 DS:0100H
MOV AL,81H
MOV BL,22H
MOV CL,33H
MOV SI,0100H
CMP AL,BL
JB  LOW1
XCHG AL,BL
LOW1:
CMP AL,CL
JB  LOW2
XCHG AL,CL
LOW2:
MOV [SI],AL
"""


CODE_8255_INPUT = """
; 8255: PA 输入，PC3 状态输入，方式 0；端口 80H-83H
MOV AL,91H          ; PA 输入，PC 高 4 位输入，PB/PC 低 4 位输出
OUT 83H,AL          ; 写控制口
MOV CX,20
MOV DI,1000H
XML:
    IN  AL,82H      ; 读 PC 口状态
    TEST AL,08H     ; 检查 PC3
    JZ  XML         ; READY=0 继续等待
    IN  AL,80H      ; READY=1，从 PA 读数据
    MOV [DI],AL
    INC DI
    LOOP XML
"""


CODE_STOPWATCH = """
DATA SEGMENT
TAB    DB 3FH,06H,5BH,4FH,66H,6DH,7DH,07H,7FH,6FH
Second DB 0
Tick   DB 0
DATA ENDS

CODE SEGMENT
ASSUME DS:DATA,CS:CODE

INT3 PROC FAR
    PUSH AX
    PUSH DS
    MOV AX,DATA
    MOV DS,AX
    INC Tick
    CMP Tick,100
    JB  INT_DONE
    MOV Tick,0
    INC Second
    CMP Second,10
    JB  INT_DONE
    MOV Second,0
INT_DONE:
    MOV AL,20H
    OUT 80H,AL      ; 给 8259 发普通 EOI
    POP DS
    POP AX
    IRET
INT3 ENDP

DISP PROC FAR
    PUSH AX
    PUSH BX
    PUSH DX
    MOV BL,Second
    MOV BH,0
    MOV AL,TAB[BX]
    MOV DX,280H
    OUT DX,AL       ; PA 输出段码
    MOV AL,0FEH
    MOV DX,282H
    OUT DX,AL       ; PC0=0 点亮
    POP DX
    POP BX
    POP AX
    RET
DISP ENDP

START:
    MOV AX,DATA
    MOV DS,AX
    CLI
    MOV AL,80H      ; 8255: A 口、C 口方式 0 输出
    MOV DX,283H
    OUT DX,AL

    MOV AX,0
    MOV ES,AX
    MOV DI,13H*4
    MOV AX,OFFSET INT3
    MOV ES:[DI],AX
    MOV AX,SEG INT3
    MOV ES:[DI+2],AX

    ; 按题目给出的端口补全 8259 初始化/屏蔽字
    MOV AL,13H
    MOV DX,80H
    OUT DX,AL
    MOV AL,10H
    MOV DX,81H
    OUT DX,AL
    MOV AL,09H      ; 普通 EOI 方式
    OUT DX,AL
    MOV AL,0F7H     ; 允许 IR3，屏蔽其它位
    OUT DX,AL
    STI
BEGIN:
    CALL DISP
    JMP BEGIN
CODE ENDS
END START
"""


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("一、填空题", level=1)
    add_answer_table(doc, [
        ["1", "CPU、存储器、I/O 接口"],
        ["2", "[x+y]补 = 11101111B；[y-x]补 = 11000011B。解析：22 + (-39) = -17，8 位补码 EFH；-39 - 22 = -61，8 位补码 C3H。"],
        ["3", "非屏蔽中断；INTR。"],
        ["4", "控制总线。"],
        ["5", "中断服务程序入口地址；00000H；003FFH；00028H。解析：每个向量 4 字节，n=10 时 4n=28H。"],
        ["6", "总线周期。"],
        ["7", "总线接口部件；执行部件；指令。"],
        ["8", "DRAM。"],
        ["9", "地址锁存允许信号。"],
        ["10", "无条件传送、查询传送、中断传送、DMA 传送。"],
        ["10(1)", "ADD BYTE PTR[BP+SI],10H：基址变址寻址，默认段 SS。物理地址 = 1000H*10H + 00A0H + 0050H = 100F0H。"],
        ["10(2)", "MOV [BX],AX：寄存器间接寻址，默认段 DS。物理地址 = 1500H*10H + 0500H = 15500H。"],
        ["11", "清 D3、D8：AND BX,0FEF7H。低 2 位取反：XOR BYTE PTR [100H],03H。"],
    ], header=["题号", "答案与解析"])

    doc.add_heading("二、选择题", level=1)
    add_answer_table(doc, [
        ["答案", "1 C；2 D；3 A；4 B；5 B；6 G、F；7 B、F；8 D；9 D；10 B；11 C；12 B；13 C；14 C。"],
    ], header=["题号", "选项"])

    doc.add_heading("三、编程和读程题", level=1)
    add_para(doc, "1. 33H - 4FH = E4H。AL=E4H；无符号减法发生借位，CF=1；两个正数相减结果在有符号范围内，OF=0。")
    add_para(doc, "2. CNT 从 STR1 到当前位置，包含 STR1 两字节和 STR2 十字节，所以 CX=000CH。STR1 DW 'AB' 按小端存储为 41H、42H，MOV AX,STR1 后 AX=4241H。STR0+2 为 56H，所以 BL=56H。")
    add_para(doc, "3. 内存示意图：ORG 3H 后从 DS:0003 开始分配。")
    add_answer_table(doc, [
        ["DS:0003", "A DB 10H -> 10H"],
        ["DS:0004-0005", "B DW 'EF' -> 45H,46H"],
        ["DS:0006-0007", "B DW 'CD' -> 43H,44H"],
        ["DS:0008-0009", "C DB 2 DUP(3) -> 03H,03H"],
        ["DS:000A-000D", "SUM DD 80906050H -> 50H,60H,90H,80H"],
    ], widths=(1.4, 5.0), header=["地址", "内容"])
    add_para(doc, "4. 完整程序如下：")
    add_code(doc, CODE_MIN)

    doc.add_heading("四、综合题", level=1)
    add_para(doc, "1. 存储器扩展：128Kx8 / 16Kx1 = 64 片；每 8 片构成一组 16Kx8，共 8 组。16K 深度需要 14 根片内地址线；8 组至少需要 3 根片选地址线。")
    add_memory_diagram(doc)
    add_para(doc, "2. 8255 查询输入完整程序：")
    add_code(doc, CODE_8255_INPUT)
    add_para(doc, "3. 秒表系统连接和完整参考程序：")
    add_stopwatch_diagram(doc)
    add_code(doc, CODE_STOPWATCH)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
