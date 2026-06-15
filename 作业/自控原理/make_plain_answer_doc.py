from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "自控原理_第5章_直观版答案.docx"


def set_run_font(run, name="Arial", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2 + 0.25 * level)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=(90, 90, 90))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=15 if level == 1 else 12, bold=True)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=18, bold=True)
    return p


def add_image(doc, path, width=Inches(6.6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(path), width=width)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(11)

add_title(doc, "自控原理 第5章 直观版答案")
add_note(doc, "这份是给你“照着抄”的版本：每题先看思路，再看关键结果，公式尽量少写。")

add_heading(doc, "5-8", 1)
add_bullet(doc, "先看分母里的三个一阶环节，对应拐点是 0.5、2、5 rad/s。")
add_bullet(doc, "低频增益是 10/(2×0.5)=10，也就是 20 dB。")
add_bullet(doc, "幅频图怎么画：0.5 之后斜率先变成 -20 dB/dec，2 之后变成 -40 dB/dec，5 之后变成 -60 dB/dec。")
add_bullet(doc, "相频图怎么画：三项反正切加起来，最后从 0° 慢慢掉到 -270°。")
add_bullet(doc, "结果：相位交叉附近的相角裕度约 20.25°，幅值裕度约 1.925 倍。")
add_image(doc, ROOT / "fig_5_8_reference.png")

add_heading(doc, "5-9", 1)
add_bullet(doc, "这个系统有两个积分环节，所以低频幅值会先很陡地下去。")
add_bullet(doc, "K=1 时，增益交叉频率约 1.01 rad/s，相角裕度约 10.26°。")
add_bullet(doc, "如果要求相角裕度 45°，先把相位条件写成“两个反正切之差=45°”。")
add_bullet(doc, "实际常取低频那组解：ωc≈6.49 rad/s，对应 K≈25.9，作业里通常写成 K≈25。")
add_bullet(doc, "如果老师不排斥高频解，还可以补一句：还有一组 ωc≈38.51 rad/s，对应 K≈241。")
add_image(doc, ROOT / "fig_5_9_reference.png")

add_heading(doc, "5-13", 1)
add_bullet(doc, "(1) 直接把闭环特征方程展开：s²+(0.1K-1)s+K=0。")
add_bullet(doc, "二阶系统稳定的直观条件就是两个系数都为正，所以 K>10。")
add_bullet(doc, "(3) 展开后是 0.025s³+0.35s²+s+K=0。")
add_bullet(doc, "Routh 判据最后得到 0<K<14 稳定。")
add_image(doc, ROOT / "fig_5_13_reference.png")

add_heading(doc, "5-14", 1)
add_bullet(doc, "这题看的是 Nyquist 图，不是硬算代数式。先记住题图给的是 K=500。")
add_bullet(doc, "如果换成别的 K，整条曲线按比例缩放。")
add_bullet(doc, "从原图读临界点，可以得到三个分界增益：10、25、10000。")
add_bullet(doc, "所以稳定区间是 0<K<10 或 25<K<10000。")
add_bullet(doc, "不稳定区间就是 10<K<25 或 K>10000。")
add_image(doc, ROOT / "fig_5_14_reference.png", width=Inches(6.8))

add_heading(doc, "5-15", 1)
add_bullet(doc, "开环相角就是 -90° 再减去 atan(ωT)。")
add_bullet(doc, "给定相角裕度 36°，可以推出 atan(ωcT)=54°。")
add_bullet(doc, "最后得到 T≈0.0234 s。")
add_bullet(doc, "闭环是标准二阶型，谐振峰值约 Mr≈1.62。")
add_image(doc, ROOT / "fig_5_15_reference.png")

add_note(doc, "如果你想要，我还能再给你压缩成“纯手写抄题版”，每题只留 3 到 5 行。")

doc.save(OUT)
print(OUT)
