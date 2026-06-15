from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "自控原理_第5章_完整过程版答案.docx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text="", size=11, bold=False, color=None, space_after=3, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color)
    return p


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    font(r, size=18, bold=True, name="Microsoft YaHei")


def h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    font(r, size=15, bold=True, color=(0, 70, 140), name="Microsoft YaHei")


def eq(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    font(r, size=10.5, name="Consolas")
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    return p


def ans(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("结论：" + text)
    font(r, size=11.5, bold=True, color=(180, 0, 0), name="Microsoft YaHei")
    return p


def image(doc, filename, width=6.45):
    path = ROOT / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(width))


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
normal.font.size = Pt(11)

title(doc, "自控原理第5章作业完整过程版")
para(doc, "说明：这版尽量按你给的旧答案图来写，保留完整计算过程；公式用普通文本排版，方便照着抄。", size=10.5, color=(90, 90, 90))

h(doc, "5-8")
para(doc, "题目：G(s)=10/[(0.2s+1)(s+2)(s+0.5)]，画频率特性曲线和对数幅频渐近线，并估算裕度。")
para(doc, "解：先把分母写成标准的一阶惯性环节。")
eq(doc, "s+2 = 2(s/2+1)，s+0.5 = 0.5(s/0.5+1)")
eq(doc, "G(s)=10/[(0.2s+1)(s/2+1)(s/0.5+1)]")
para(doc, "所以低频增益为 10，即 20lg10=20 dB。三个转折频率为：")
eq(doc, "ω1=0.5，ω2=2，ω3=5  rad/s")
para(doc, "对数幅频渐近线画法：")
eq(doc, "ω<0.5：L(ω)=20 dB")
eq(doc, "0.5<ω<2：斜率为 -20 dB/dec")
eq(doc, "2<ω<5：斜率为 -40 dB/dec")
eq(doc, "ω>5：斜率为 -60 dB/dec")
para(doc, "精确幅值为：")
eq(doc, "|G(jω)|=10/sqrt[(1+0.04ω^2)(1+0.25ω^2)(1+4ω^2)]")
para(doc, "令 |G(jωc)|=1，解得增益交叉频率：")
eq(doc, "ωc≈2.63 rad/s")
para(doc, "相角为：")
eq(doc, "φ(ω)= -arctan(0.2ω) - arctan(0.5ω) - arctan(2ω)")
para(doc, "代入 ωc≈2.63：")
eq(doc, "φ(ωc)≈-159.75°")
eq(doc, "相角裕度 γ = 180° + φ(ωc) ≈ 20.25°")
para(doc, "再求幅值裕度。令相角等于 -180°：")
eq(doc, "arctan(0.2ω) + arctan(0.5ω) + arctan(2ω) = 180°")
para(doc, "由三角关系可得：")
eq(doc, "0.2ω^3 = 2.7ω")
eq(doc, "ωg = sqrt(13.5) ≈ 3.67 rad/s")
para(doc, "代入幅值公式：")
eq(doc, "|G(jωg)|≈0.519")
eq(doc, "幅值裕度 Kg = 1/|G(jωg)| ≈ 1.925")
eq(doc, "20lgKg≈5.69 dB")
ans(doc, "相角裕度约 20.25°；幅值裕度约 1.925 倍，即 5.69 dB。")
para(doc, "图按下面这样画，红虚线是渐近线，蓝线是精确曲线：", size=10.5)
image(doc, "fig_5_8_reference.png")

h(doc, "5-9")
para(doc, "题目：G(s)=K(0.2s+1)/[s^2(0.02s+1)]。")
para(doc, "（1）当 K=1 时，求相角裕度。")
eq(doc, "G(s)=(0.2s+1)/[s^2(0.02s+1)]")
para(doc, "幅值为：")
eq(doc, "|G(jω)|=sqrt(1+0.04ω^2)/(ω^2 sqrt(1+0.0004ω^2))")
para(doc, "令 |G(jωc)|=1，解得：")
eq(doc, "ωc≈1.01 rad/s")
para(doc, "相角为：")
eq(doc, "φ(ω)= -180° + arctan(0.2ω) - arctan(0.02ω)")
para(doc, "代入 ωc≈1.01：")
eq(doc, "φ(ωc)≈-169.74°")
eq(doc, "γ=180°+φ(ωc)≈10.26°")
ans(doc, "K=1 时，相角裕度约为 10.26°。")
para(doc, "（2）要求相角裕度为 45°，求 K。")
para(doc, "相角裕度 γ=45°，说明增益交叉频率处相角应为：")
eq(doc, "φ(ωc)=-135°")
para(doc, "代入相角表达式：")
eq(doc, "-180° + arctan(0.2ωc) - arctan(0.02ωc) = -135°")
eq(doc, "arctan(0.2ωc) - arctan(0.02ωc) = 45°")
para(doc, "对两边取正切：")
eq(doc, "tan[A-B] = (0.2ωc-0.02ωc)/(1+0.2ωc·0.02ωc) = 1")
eq(doc, "0.18ωc/(1+0.004ωc^2)=1")
eq(doc, "0.004ωc^2 - 0.18ωc + 1 = 0")
para(doc, "解得两个频率。和旧答案一致，取低频交叉频率：")
eq(doc, "ωc≈6.49 rad/s")
para(doc, "再由增益交叉条件 |G(jωc)|=1 求 K：")
eq(doc, "K·sqrt(1+0.04ωc^2)/(ωc^2 sqrt(1+0.0004ωc^2)) = 1")
eq(doc, "K = ωc^2 sqrt(1+0.0004ωc^2)/sqrt(1+0.04ωc^2)")
eq(doc, "K≈25.9")
ans(doc, "取旧答案的低频解，K≈25，较精确为 25.9。")
para(doc, "注：方程还有高频解 ωc≈38.51，对应 K≈241；旧答案通常取低频解。", size=10.5, color=(90, 90, 90))
image(doc, "fig_5_9_reference.png")

h(doc, "5-13")
para(doc, "题目要求用奈奎斯特稳定判据确定闭环稳定的 K 范围。")
para(doc, "（1）G(s)=K(0.1s+1)/[s(s-1)]。")
para(doc, "开环有一个右半平面极点 s=1，因此 P=1。先求 Nyquist 曲线与实轴的交点。")
eq(doc, "G(jω)=K(1+0.1jω)/[jω(jω-1)]")
para(doc, "化成实部和虚部：")
eq(doc, "G(jω)= -1.1K/(ω^2+1) + j·K(1-0.1ω^2)/[ω(ω^2+1)]")
para(doc, "令虚部为 0：")
eq(doc, "1-0.1ω^2=0")
eq(doc, "ω^2=10")
para(doc, "代回实部：")
eq(doc, "Re[G(jω)] = -K/10")
para(doc, "由于 P=1，要使闭环稳定，Nyquist 曲线应按要求包围 (-1,j0) 一次。分界点为：")
eq(doc, "-K/10 = -1")
eq(doc, "K=10")
ans(doc, "第（1）小题闭环稳定条件：K>10。")
para(doc, "（3）G(s)=K/[s(0.1s+1)(0.25s+1)]。")
para(doc, "先把一阶环节换成整数形式：")
eq(doc, "(0.1s+1)=(s+10)/10，(0.25s+1)=(s+4)/4")
eq(doc, "G(s)=40K/[s(s+10)(s+4)]")
para(doc, "开环除原点极点外，没有右半平面极点，按旧答案记 P=0。")
eq(doc, "G(jω)=40K/[jω(jω+10)(jω+4)]")
para(doc, "展开分母：")
eq(doc, "jω(jω+10)(jω+4) = -14ω^2 + jω(40-ω^2)")
para(doc, "所以：")
eq(doc, "G(jω)=40K[-14ω^2 - jω(40-ω^2)]/[196ω^4+ω^2(40-ω^2)^2]")
para(doc, "令虚部为 0：")
eq(doc, "40-ω^2=0")
eq(doc, "ω^2=40")
para(doc, "代回实部：")
eq(doc, "Re[G(jω)] = -K/14")
para(doc, "临界时通过 (-1,j0)：")
eq(doc, "-K/14 = -1")
eq(doc, "K=14")
ans(doc, "第（3）小题闭环稳定条件：0<K<14。")
image(doc, "fig_5_13_reference.png")

h(doc, "5-14")
para(doc, "题目给的是 K=500 时的 Nyquist 图。开环极点在右半平面的个数 p=0。")
para(doc, "这题的关键是：改变 K 时，Nyquist 图的形状不变，只是整体按比例缩放。")
eq(doc, "一般 K 的曲线 = 题图曲线 × (K/500)")
para(doc, "题图中负实轴交点大约为 -50、-20、-0.05。把它们缩放后分别令其等于 -1，就得到临界 K。")
para(doc, "第一处交点：")
eq(doc, "-50 × K/500 = -1")
eq(doc, "K=10")
para(doc, "第二处交点：")
eq(doc, "-20 × K/500 = -1")
eq(doc, "K=25")
para(doc, "第三处交点：")
eq(doc, "-0.05 × K/500 = -1")
eq(doc, "K=10000")
para(doc, "所以三个临界增益是：")
eq(doc, "K=10，K=25，K=10000")
para(doc, "按 Nyquist 判据，因为 P=0，闭环稳定要求曲线不包围 (-1,j0)。根据题图包围情况分段判断：")
eq(doc, "0<K<10：不包围，稳定")
eq(doc, "10<K<25：包围，不稳定")
eq(doc, "25<K<10000：不包围，稳定")
eq(doc, "K>10000：包围，不稳定")
ans(doc, "闭环稳定区间：0<K<10 或 25<K<10000；不稳定区间：10<K<25 或 K>10000。")
image(doc, "fig_5_14_reference.png", width=6.6)

h(doc, "5-15")
para(doc, "题目：G(s)=100/[s(Ts+1)]，已知相角裕度 γ=36°，求 T 和闭环谐振峰值 Mr。")
para(doc, "开环频率特性：")
eq(doc, "G(jω)=100/[jω(jωT+1)]")
para(doc, "幅值和相角分别为：")
eq(doc, "|G(jω)|=100/[ω sqrt(1+ω^2T^2)]")
eq(doc, "φ(ω)= -90° - arctan(ωT)")
para(doc, "相角裕度定义：")
eq(doc, "γ = 180° + φ(ωc)")
para(doc, "代入 γ=36°：")
eq(doc, "36° = 180° - 90° - arctan(ωcT)")
eq(doc, "arctan(ωcT)=54°")
eq(doc, "ωcT=tan54°≈1.376")
para(doc, "又因为 ωc 是增益交叉频率，所以 |G(jωc)|=1：")
eq(doc, "100/[ωc sqrt(1+(ωcT)^2)] = 1")
eq(doc, "ωc = 100/sqrt(1+tan^2 54°)")
eq(doc, "ωc≈58.78 rad/s")
para(doc, "于是：")
eq(doc, "T = tan54°/ωc ≈ 1.376/58.78")
eq(doc, "T≈0.0234 s")
para(doc, "闭环传递函数为：")
eq(doc, "Φ(s)=G(s)/(1+G(s)) = 100/(Ts^2+s+100)")
para(doc, "化成标准二阶系统形式：")
eq(doc, "Φ(s)=ωn^2/(s^2+2ζωn s+ωn^2)")
para(doc, "由 Ts^2+s+100 除以 T 得：")
eq(doc, "s^2 + (1/T)s + 100/T")
eq(doc, "ωn^2=100/T，2ζωn=1/T")
eq(doc, "ζ=1/(20sqrt(T))")
para(doc, "代入 T≈0.0234：")
eq(doc, "ζ≈0.327")
para(doc, "二阶系统谐振峰值公式：")
eq(doc, "Mr=1/[2ζsqrt(1-ζ^2)]")
eq(doc, "Mr≈1.62")
ans(doc, "T≈0.0234 s，闭环谐振峰值 Mr≈1.62。")
image(doc, "fig_5_15_reference.png")

doc.save(OUT)
print(OUT)
